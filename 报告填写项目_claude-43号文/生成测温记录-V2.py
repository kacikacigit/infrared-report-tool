# -*- coding: utf-8 -*-
"""
测温记录自动生成脚本 V2
自动找最新版本的测温报告，输出自动编号的测温记录，不覆盖旧文件
"""

import os
import re
import glob

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================== 配置 =====================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))
REPORT_BASE = os.path.join(BASE_DIR, "测温报告_已生成")   # V2: 自动找最新
TEMPLATE_PATH = os.path.join(BASE_DIR, "测温记录.docx")
OUTPUT_BASE = os.path.join(BASE_DIR, "测温记录_已生成")   # V2: 自动编号

CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"

PHASE_PATTERN = re.compile(r'^(.*)(A相|B相|C相)$')

# ===================== 字体工具 =====================


def apply_font(run):
    run.font.name = WESTERN_FONT
    run.font.size = Pt(8.5)
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rFonts.set(qn('w:ascii'), WESTERN_FONT)
    rFonts.set(qn('w:hAnsi'), WESTERN_FONT)


def apply_font_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    apply_font(run)


# ===================== 读取测温报告 =====================


def read_report(report_path):
    """从测温报告中提取所有需要的数据"""
    doc = Document(report_path)
    table = doc.tables[0]

    # 基础信息（0-indexed row/col）
    date_raw = table.rows[2].cells[4].text.strip()          # 第3行第5列
    env_temp = table.rows[4].cells[4].text.strip()          # 第5行第5列
    humidity = table.rows[4].cells[7].text.strip()          # 第5行第8列
    person = table.rows[2].cells[7].text.strip()            # 第3行第8列

    # 结论：第20行第4列
    conclusion = table.rows[19].cells[3].text.strip()

    # 数据行：Rows 7-16（0-indexed），10条记录
    data_rows = []
    for i in range(7, 17):
        row = table.rows[i]
        data_rows.append({
            'device_name': row.cells[2].text.strip(),   # C列 设备名称
            'defect_part': row.cells[3].text.strip(),   # D列 缺陷部位
            'surface_temp': row.cells[4].text.strip(),  # E列 表面温度
        })

    return {
        'date_raw': date_raw,
        'env_temp': env_temp,
        'humidity': humidity,
        'person': person,
        'conclusion': conclusion,
        'data_rows': data_rows,
    }


# ===================== 数据处理 =====================


def format_date(date_str):
    """2025年09月15日 → 2025.09.15"""
    m = re.match(r'(\d{4})年(\d{2})月(\d{2})日', date_str)
    if m:
        return f"{m.group(1)}.{m.group(2)}.{m.group(3)}"
    return date_str


def parse_heating_devices(conclusion_text):
    """
    从结论文本中解析出发热设备名列表。
    "11201刀闸B相、11231刀闸A相、11231刀闸C相发热,..." → ["11201刀闸B相", "11231刀闸A相", "11231刀闸C相"]
    """
    if '发热' not in conclusion_text:
        return []
    before_fever = conclusion_text.split('发热')[0]
    devices = [d.strip() for d in before_fever.split('、') if d.strip()]
    return devices


def group_by_base_device(device_names):
    """
    按设备基名分组，保持首次出现顺序。
    "11231刀闸A相", "11231刀闸C相" → {"11231刀闸": ["A", "C"]}
    """
    groups = {}  # base_name -> [phase_letter, ...]
    order = []   # 保持基名出现顺序
    for name in device_names:
        m = PHASE_PATTERN.match(name)
        if m:
            base = m.group(1)
            phase_letter = m.group(2).replace('相', '')
            if base not in groups:
                groups[base] = []
                order.append(base)
            if phase_letter not in groups[base]:
                groups[base].append(phase_letter)
    return [(base, groups[base]) for base in order]


def get_defect_for_base(base_name, data_rows):
    """从数据行中查找该基名设备对应的缺陷部位"""
    if not data_rows:
        return ''
    for row in data_rows:
        if row['device_name'].startswith(base_name) and row['defect_part'] != '无':
            return row['defect_part']
    return ''


def build_problems_text(groups, data_rows):
    """生成'发现问题'的填充文本"""
    lines = []
    for base_name, phases in groups:
        defect = get_defect_for_base(base_name, data_rows)
        phase_str = '、'.join(phases)
        line = (
            f"{base_name}{phase_str}相{defect}发热，"
            f"根据《国家电网公司变电检测管理规定》，温差超过15K，"
            f"未达到严重缺陷的要求,达到一般缺陷标准"
        )
        lines.append(line)
    return '\n'.join(lines)


def build_heating_parts_text(groups, data_rows):
    """生成'发热部分'的填充文本，每个基名设备列出ABC三相温度"""
    lines = []
    for base_name, _heating_phases in groups:
        temps = {}
        for row in data_rows:
            name = row['device_name']
            m = PHASE_PATTERN.match(name)
            if m and m.group(1) == base_name:
                letter = m.group(2).replace('相', '')
                temps[letter] = row['surface_temp']

        parts = []
        for p in ['A', 'B', 'C']:
            t = temps.get(p, '?')
            parts.append(f"{p}相{t}℃")

        line = f"{base_name} " + '  '.join(parts)
        lines.append(line)
    return '\n'.join(lines)


# ===================== 写入模板 =====================


def set_cell_text(cell, text):
    """清空单元格所有内容，写入新文本（保留原第一段的格式作为参考）"""
    # 清空所有段落的所有runs
    for p in cell.paragraphs:
        for run in list(p.runs):
            run.text = ""

    # 将换行文本拆分到多个段落：第一段放第一行，后续每段放一行
    parts = text.split('\n')

    # 第一段：写入第一部分
    first_p = cell.paragraphs[0]
    if first_p.runs:
        first_p.runs[0].text = parts[0]
        apply_font(first_p.runs[0])
    else:
        run = first_p.add_run(parts[0])
        apply_font(run)

    # 剩余部分：追加到第一段后面（用换行分隔），保持像人工填写版本一样的样式
    for extra in parts[1:]:
        if extra:
            run = first_p.add_run('\n' + extra)
        else:
            run = first_p.add_run('\n')
        apply_font(run)

    # 清空多余的段落
    for extra_p in cell.paragraphs[1:]:
        for run in list(extra_p.runs):
            run.text = ""


def fill_basic_info(table, report_data):
    """填写基础信息到测温记录模板"""
    # 测温时间: Row 0, Col 4
    set_cell_text(table.rows[0].cells[4], report_data['date_raw'])

    # 环境温度: Row 2, Col 4
    set_cell_text(table.rows[2].cells[4], report_data['env_temp'])

    # 湿度: Row 3, Col 1 和 Col 2
    set_cell_text(table.rows[3].cells[1], report_data['humidity'])
    set_cell_text(table.rows[3].cells[2], report_data['humidity'])

    # 测温人: Row 5, Col 1 和 Col 2
    set_cell_text(table.rows[5].cells[1], report_data['person'])
    set_cell_text(table.rows[5].cells[2], report_data['person'])


def fill_problems(table, problems_text):
    """填写'发现问题'到 Row 8 的所有单元格"""
    content = '发现问题\n' + problems_text
    for cell in table.rows[8].cells:
        set_cell_text(cell, content)


def fill_heating_parts(table, heating_text):
    """填写'发热部分'到 Row 9 的所有单元格"""
    content = '发热部分\n' + heating_text
    for cell in table.rows[9].cells:
        set_cell_text(cell, content)


# ===================== 主流程 =====================


def natural_sort_key(s):
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def find_latest_report():
    """找编号最大的 测温报告_已生成_*.docx"""
    pattern = f"{REPORT_BASE}_*.docx"
    files = glob.glob(pattern)
    if not files:
        return None
    return sorted(files, key=lambda p: natural_sort_key(p))[-1]


def resolve_output_path():
    """生成不冲突的输出路径：测温记录_已生成_1.docx, _2.docx ..."""
    counter = 1
    while True:
        path = f"{OUTPUT_BASE}_{counter}.docx"
        if not os.path.exists(path):
            return path
        counter += 1


def main():
    print("=" * 60)
    print("测温记录自动生成 V2")
    print("=" * 60)

    # 1. 找最新测温报告
    report_path = find_latest_report()
    if report_path is None:
        print(f"\n错误：找不到任何 测温报告_已生成_*.docx 文件。")
        print("请先运行 43号文报告自动生成 的 Step 2 生成测温报告。")
        return

    print(f"\n[1/3] 自动读取最新测温报告: {os.path.basename(report_path)}")
    report_data = read_report(report_path)
    print(f"  - 测温时间: {report_data['date_raw']}")
    print(f"  - 环境温度: {report_data['env_temp']}")
    print(f"  - 湿度: {report_data['humidity']}")
    print(f"  - 测温人: {report_data['person']}")
    print(f"  - 数据行数: {len(report_data['data_rows'])}")

    # 2. 处理数据
    print(f"\n[2/3] 处理结论数据...")
    device_names = parse_heating_devices(report_data['conclusion'])

    if device_names:
        groups = group_by_base_device(device_names)
        print(f"  - 发现发热设备: {len(device_names)} 个")
        for base, phases in groups:
            print(f"    {base}: {', '.join(phases)}相")
        problems_text = build_problems_text(groups, report_data['data_rows'])
        heating_text = build_heating_parts_text(groups, report_data['data_rows'])
    else:
        print("  - 无发热设备")
        problems_text = '无'
        heating_text = '无'

    # 3. 填写模板
    output_path = resolve_output_path()
    print(f"\n[3/3] 填写模板 → 输出: {os.path.basename(output_path)}")
    if not os.path.isfile(TEMPLATE_PATH):
        print(f"错误：找不到模板文件: {TEMPLATE_PATH}")
        return

    doc = Document(TEMPLATE_PATH)
    table = doc.tables[0]

    fill_basic_info(table, report_data)
    fill_problems(table, problems_text)
    fill_heating_parts(table, heating_text)

    apply_font_to_table(table)

    doc.save(output_path)
    print(f"\n{'=' * 60}")
    print(f"测温记录已生成: {output_path}")
    print(f"{'=' * 60}")


if __name__ == "__main__":
    main()
