# -*- coding: utf-8 -*-
"""
红外测温报告自动生成程序
======================================================

本脚本假设以下三样东西和本脚本放在【同一个文件夹】里（不需要修改任何路径配置）：
    ./图谱/                图片文件夹（红外测温图片）
    ./必填信息.xlsx         数据表（工作表名："普测"）
    ./测温报告.docx         Word报告模板

功能：
【第一步】表面温度识别（写入必填信息.xlsx 的 E 列）
    - 交互框选：用第一张图片弹窗，鼠标拖拽框选出显示"表面温度数值"的区域
      （区域按相对比例保存，不同图片尺寸也能自动定位，无需每张图都框一次）。
    - 图片与Excel行的匹配方式：
        优先按文件名匹配：图片"图1.jpg"对应Excel"图谱编号"列内容"图1"
        如果图片命名不是这个规律，自动改为"按文件名顺序"和"Excel行顺序"一一对应。
    - 识别失败的行，E列填"无法识别"。

【第二步】正常温度计算（写入 F 列）
    - 按"设备名称"（C列）分组：
        如果某个前缀（去掉末尾"A相/B相/C相"）恰好凑齐3个（A/B/C三相都有），
        这3行为一组，正常温度 = 这3行表面温度的最小值，3行都填一样的值。
        除此以外的所有情况（不含A/B/C相后缀的设备、或凑不齐3个的），
        每行单独成组，正常温度 = 该行自己的表面温度。

【第三步】温差计算（写入 K 列，新增列）
    - 温差 = 表面温度(E) - 正常温度(F)

【第四步】生成测温报告.docx
    - 把必填信息.xlsx中A~J列数据（从第2行开始）填入测温报告.docx第一个表格
      （模板行数不够会自动增加）。
    - 温差 > 15 的行视为"发热对象"，自动生成结论文字填入第一个表格倒数第二行
      （结论一栏，横跨4~10列）。若没有温差>15的行，结论填"无异常"。
    - 把图谱文件夹里的图片按顺序插入第二个表格（图片列，尺寸 12cm × 9cm），
      模板行数不够会自动增加。

依赖安装（在你的虚拟环境中执行）：
    pip install opencv-python numpy pytesseract openpyxl pillow python-docx

还需要系统安装 Tesseract OCR 引擎本体（pip装不了这个）：
    brew install tesseract   或   conda install -c conda-forge tesseract

使用方法：
    python3 红外测温报告自动生成.py
"""

import os
import re
import sys
import json
import glob
import copy

import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================== 配置区域 =====================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

IMAGE_DIR = os.path.join(BASE_DIR, "图谱")
EXCEL_PATH = os.path.join(BASE_DIR, "必填信息.xlsx")
EXCEL_OUTPUT_PATH = os.path.join(BASE_DIR, "必填信息_生成.xlsx")
EXCEL_SHEET_NAME = "普测"

DOCX_TEMPLATE_PATH = os.path.join(BASE_DIR, "测温报告.docx")
DOCX_OUTPUT_PATH = os.path.join(BASE_DIR, "测温报告_已生成.docx")

ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_infrared.json")
SAVE_DEBUG_CROPS = True
DEBUG_DIR = os.path.join(BASE_DIR, "debug_crops")
VERBOSE = True

NOT_RECOGNIZED_TEXT = "无法识别"

# 发热判定阈值：温差 > 该值 视为发热对象
HOT_THRESHOLD = 15

# 结论模板
CONCLUSION_TEMPLATE = "{names}发热,达到一般缺陷标准,未达到严重缺陷标准（发热缺陷较前期无明显发展趋势），无其他异常。"
CONCLUSION_NORMAL = "无异常"

# 图片插入尺寸
IMAGE_WIDTH_CM = 12
IMAGE_HEIGHT_CM = 9

# =================================================================


# ===================== 工具函数：自然排序 =====================

def natural_sort_key(s):
    """自然排序：让"图2"排在"图10"前面，而不是按字符串排序排到后面。"""
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def collect_image_files(image_dir):
    patterns = ["*.jpg", "*.JPG", "*.jpeg", "*.JPEG", "*.png", "*.PNG"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(image_dir, p)))
    files = sorted(set(files), key=lambda p: natural_sort_key(os.path.basename(p)))
    return files


# ===================== 第一步：OCR识别表面温度 =====================

def select_roi_interactively(sample_image_path):
    img = cv2.imread(sample_image_path)
    if img is None:
        raise FileNotFoundError(f"无法读取示例图片: {sample_image_path}")
    img_h, img_w = img.shape[:2]

    print("\n请在弹出的窗口中用鼠标左键拖拽，框选出【表面温度数值】所在的区域。")
    print("框选完成后按【空格键】或【回车键】确认；按 c 键可取消重新框选。\n")

    window_name = "请框选表面温度数值区域 (Space/Enter确认, C取消)"
    roi = cv2.selectROI(window_name, img, showCrosshair=True, fromCenter=False)
    cv2.destroyAllWindows()

    x, y, w, h = roi
    if w == 0 or h == 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    rx, ry, rw, rh = x / img_w, y / img_h, w / img_w, h / img_h
    print(f"框选区域（像素）: x={x}, y={y}, w={w}, h={h}  (示例图片尺寸: {img_w}x{img_h})")
    print(f"已换算为相对比例: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
    return rx, ry, rw, rh


def get_roi(sample_image_path):
    if os.path.exists(ROI_CONFIG_FILE):
        with open(ROI_CONFIG_FILE, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        rx, ry, rw, rh = cfg["rx"], cfg["ry"], cfg["rw"], cfg["rh"]
        print(f"检测到已保存的识别区域配置（相对比例）: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
        choice = input("是否复用该区域？(y=复用 / n=重新框选) [y]: ").strip().lower()
        if choice in ("", "y", "yes"):
            return rx, ry, rw, rh

    rx, ry, rw, rh = select_roi_interactively(sample_image_path)
    with open(ROI_CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump({"rx": rx, "ry": ry, "rw": rw, "rh": rh}, f, ensure_ascii=False, indent=2)
    print(f"识别区域已保存到: {ROI_CONFIG_FILE}")
    return rx, ry, rw, rh


def crop_by_relative_roi(img, roi_ratio):
    rx, ry, rw, rh = roi_ratio
    img_h, img_w = img.shape[:2]
    x = int(round(rx * img_w))
    y = int(round(ry * img_h))
    w = int(round(rw * img_w))
    h = int(round(rh * img_h))
    x2, y2 = min(x + w, img_w), min(y + h, img_h)
    x, y = max(x, 0), max(y, 0)
    return img[y:y2, x:x2]


def preprocess_variants(crop):
    scale = 4
    crop_big = cv2.resize(crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(crop_big, cv2.COLOR_BGR2GRAY)

    variants = []

    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh1 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh1)

    inv = cv2.bitwise_not(blur)
    _, thresh2 = cv2.threshold(inv, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh2)

    thresh3 = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
    )
    variants.append(thresh3)

    return variants


def extract_number_signed(text):
    """提取数字（保留正负号，温度不取绝对值），返回字符串或None。"""
    if not text:
        return None
    text = text.strip()
    match = re.search(r'-?\d+\.?\d*', text)
    if not match:
        return None
    try:
        num = float(match.group())
        if num == int(num):
            return str(int(num))
        return str(num)
    except ValueError:
        return None


def ocr_extract_temperature(image_path, roi_ratio):
    img = cv2.imread(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]

    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, ["[裁剪区域为空，请检查ROI设置]"]

    variants = preprocess_variants(crop)
    configs = [
        r'--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 11 -c tessedit_char_whitelist=0123456789.-',
    ]

    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            value = extract_number_signed(text)
            if value is not None:
                return value, variant, raw_texts

    return None, variants[0], raw_texts


def match_images_to_rows(image_files, atlas_codes):
    """
    将图片文件和Excel行的"图谱编号"进行匹配。
    优先按文件名（不含扩展名）精确匹配；如果匹配率过低，改为按顺序位置匹配。
    返回：{行索引(0开始): 图片路径 或 None}
    """
    name_to_path = {os.path.splitext(os.path.basename(p))[0]: p for p in image_files}

    name_matched = 0
    for code in atlas_codes:
        if code is not None and str(code).strip() in name_to_path:
            name_matched += 1

    use_name_matching = (len(atlas_codes) > 0 and name_matched >= len(atlas_codes) * 0.8)

    result = {}
    if use_name_matching:
        print("图片匹配方式：按文件名与Excel'图谱编号'精确匹配。")
        for i, code in enumerate(atlas_codes):
            code_str = str(code).strip() if code is not None else ""
            result[i] = name_to_path.get(code_str)
    else:
        print("图片匹配方式：文件名与'图谱编号'对不上，改为按【图片文件名顺序】与【Excel行顺序】依次对应。")
        for i, code in enumerate(atlas_codes):
            result[i] = image_files[i] if i < len(image_files) else None

    return result


def step1_ocr_fill_surface_temperature(ws, header_row_idx, data_rows_info):
    """
    识别表面温度并写入E列。
    data_rows_info: list of dict，每个元素至少包含 {'row': excel行号, 'atlas_code': 图谱编号}
    返回：{行号: 表面温度值字符串 或 None}
    """
    image_files = collect_image_files(IMAGE_DIR)
    if not image_files:
        print(f"警告：图谱文件夹中未找到任何图片: {IMAGE_DIR}")
        return {info['row']: None for info in data_rows_info}

    print(f"图谱文件夹共找到 {len(image_files)} 张图片。")

    atlas_codes = [info['atlas_code'] for info in data_rows_info]
    row_to_image = match_images_to_rows(image_files, atlas_codes)

    # 用第一个有效图片作为框选样本
    sample_path = next((p for p in row_to_image.values() if p), image_files[0])
    roi_ratio = get_roi(sample_path)

    if SAVE_DEBUG_CROPS:
        os.makedirs(DEBUG_DIR, exist_ok=True)

    surface_temp = {}
    print("\n开始批量识别表面温度，请稍候...\n")
    for i, info in enumerate(data_rows_info):
        row = info['row']
        img_path = row_to_image.get(i)
        if img_path is None:
            surface_temp[row] = None
            print(f"[第{row}行] 图谱编号={info['atlas_code']}  ->  未找到对应图片，跳过")
            continue

        value, processed, raw_texts = ocr_extract_temperature(img_path, roi_ratio)
        surface_temp[row] = value

        status = value if value is not None else NOT_RECOGNIZED_TEXT
        print(f"[第{row}行] {os.path.basename(img_path)}  ->  {status}")

        if VERBOSE and value is None:
            cleaned = [t if t else "(空)" for t in raw_texts]
            print(f"      OCR原始识别文本（各模式）: {cleaned}")

        if SAVE_DEBUG_CROPS and processed is not None:
            debug_name = f"row{row}_{status}.png"
            debug_name = re.sub(r'[\\/:*?"<>|]', "_", debug_name)
            cv2.imwrite(os.path.join(DEBUG_DIR, debug_name), processed)

    return surface_temp


# ===================== 第二、三步：正常温度 & 温差计算 =====================

PHASE_SUFFIX_RE = re.compile(r'^(.*)(A相|B相|C相)$')


def compute_normal_temperature_and_diff(data_rows_info, surface_temp):
    """
    根据分组规则计算正常温度(F)和温差(K)。
    返回：{行号: (正常温度值 或 None 或 '无法识别', 温差值 或 None 或 '无法识别')}
    """
    # 按前缀分组（仅针对以 A相/B相/C相 结尾的设备名称）
    prefix_groups = {}
    for info in data_rows_info:
        name = str(info['device_name']).strip() if info['device_name'] else ""
        m = PHASE_SUFFIX_RE.match(name)
        if m:
            prefix = m.group(1)
            prefix_groups.setdefault(prefix, []).append(info['row'])

    # 只有恰好凑齐3个的前缀才算一组，其余的都当作单独一行处理
    valid_groups = {prefix: rows for prefix, rows in prefix_groups.items() if len(rows) == 3}
    grouped_rows = set(r for rows in valid_groups.values() for r in rows)

    result = {}

    # 先处理分组（3个一组）
    for prefix, rows in valid_groups.items():
        values = []
        ok = True
        for r in rows:
            v = surface_temp.get(r)
            if v is None:
                ok = False
                break
            try:
                values.append(float(v))
            except ValueError:
                ok = False
                break
        if ok:
            normal_temp = min(values)
            for r in rows:
                e_val = float(surface_temp[r])
                diff = round(e_val - normal_temp, 2)
                result[r] = (format_number(normal_temp), format_number(diff))
        else:
            for r in rows:
                result[r] = (NOT_RECOGNIZED_TEXT, NOT_RECOGNIZED_TEXT)

    # 再处理单独成组的行
    for info in data_rows_info:
        r = info['row']
        if r in grouped_rows:
            continue
        v = surface_temp.get(r)
        if v is None:
            result[r] = (NOT_RECOGNIZED_TEXT, NOT_RECOGNIZED_TEXT)
        else:
            e_val = float(v)
            normal_temp = e_val
            diff = 0.0
            result[r] = (format_number(normal_temp), format_number(diff))

    return result


def format_number(num):
    """把浮点数格式化成字符串：整数不带小数点，否则保留原有精度（最多2位小数）。"""
    if num == int(num):
        return str(int(num))
    return str(round(num, 2))


# ===================== Excel 读写 =====================

def read_excel_data_rows(ws, header_row_idx):
    """从表头下一行开始读取数据，直到设备名称列(C)为空为止。"""
    data_rows_info = []
    row = header_row_idx + 1
    while True:
        c_val = ws.cell(row=row, column=3).value  # C列 设备名称
        if c_val is None or str(c_val).strip() == "":
            break
        info = {
            'row': row,
            'seq': ws.cell(row=row, column=1).value,       # A 序号
            'interval': ws.cell(row=row, column=2).value,   # B 间隔名称
            'device_name': c_val,                             # C 设备名称
            'defect_part': ws.cell(row=row, column=4).value,  # D 缺陷部位
            'env_temp': ws.cell(row=row, column=7).value,     # G 环境温度
            'load_current': ws.cell(row=row, column=8).value, # H 负荷电流
            'atlas_code': ws.cell(row=row, column=9).value,   # I 图谱编号
            'remark': ws.cell(row=row, column=10).value,      # J 备注
        }
        data_rows_info.append(info)
        row += 1
    return data_rows_info


def find_header_row(ws, keyword="序号", max_scan_rows=10):
    for row in range(1, max_scan_rows + 1):
        if str(ws.cell(row=row, column=1).value).strip() == keyword:
            return row
    return 1


def step2_process_excel():
    print(f"正在读取Excel: {EXCEL_PATH}  (工作表: {EXCEL_SHEET_NAME})")
    wb = load_workbook(EXCEL_PATH)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        print(f"错误：Excel中找不到名为'{EXCEL_SHEET_NAME}'的工作表。")
        sys.exit(1)
    ws = wb[EXCEL_SHEET_NAME]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)

    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行，请检查表格内容。")
        sys.exit(1)

    print(f"共读取到 {len(data_rows_info)} 行数据。")

    # 第一步：OCR识别表面温度，写入E列
    surface_temp = step1_ocr_fill_surface_temperature(ws, header_row_idx, data_rows_info)
    for info in data_rows_info:
        r = info['row']
        v = surface_temp.get(r)
        ws.cell(row=r, column=5).value = v if v is not None else NOT_RECOGNIZED_TEXT

    # 第二、三步：计算正常温度(F)和温差(K)
    calc_result = compute_normal_temperature_and_diff(data_rows_info, surface_temp)

    # 确保K列表头存在
    k_header = ws.cell(row=header_row_idx, column=11).value
    if not k_header:
        ws.cell(row=header_row_idx, column=11).value = "温差（℃）"

    for info in data_rows_info:
        r = info['row']
        normal_temp, diff = calc_result[r]
        ws.cell(row=r, column=6).value = normal_temp   # F列 正常温度
        ws.cell(row=r, column=11).value = diff          # K列 温差

    wb.save(EXCEL_OUTPUT_PATH)
    print(f"\nExcel已生成并保存（原文件未改动）: {EXCEL_OUTPUT_PATH}")

    # 组装完整数据（含E/F/K），供后续生成Word报告使用
    full_data = []
    for info in data_rows_info:
        r = info['row']
        row_data = dict(info)
        row_data['surface_temp'] = surface_temp.get(r)
        row_data['normal_temp'], row_data['diff'] = calc_result[r]
        full_data.append(row_data)

    return full_data


# ===================== 第四步：生成 测温报告.docx =====================

CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"


def apply_font(run, east_asian_font=CHINESE_FONT, latin_font=WESTERN_FONT):
    """设置字体：中文用宋体，英文/数字用Times New Roman。"""
    run.font.name = latin_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian_font)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)


def set_cell_text(cell, text, keep_first_run_format=True):
    """清空单元格已有内容，写入新文本，并统一设置字体（中文宋体/英文数字Times New Roman）。"""
    text = "" if text is None else str(text)
    paragraphs = cell.paragraphs
    p = paragraphs[0]
    # 清空该段落已有的runs
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    apply_font(run)
    # 清空多余段落
    for extra_p in paragraphs[1:]:
        extra_p.text = ""


def apply_font_to_paragraph(paragraph):
    for run in paragraph.runs:
        apply_font(run)


def apply_font_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                apply_font_to_paragraph(p)
            # 兼容单元格内嵌套表格的情况
            for nested_table in cell.tables:
                apply_font_to_table(nested_table)


def apply_font_to_whole_document(doc):
    """把整篇文档（包括模板里原有的静态文字）统一设置为：中文宋体，英文/数字Times New Roman。"""
    for p in doc.paragraphs:
        apply_font_to_paragraph(p)
    for t in doc.tables:
        apply_font_to_table(t)


def find_row_by_first_cell_text(table, target_text):
    for row in table.rows:
        if row.cells[0].text.strip() == target_text:
            return row
    return None


def clone_row_before(table, template_row, target_row):
    """复制template_row的格式，插入到target_row之前，返回新行对象。"""
    new_tr = copy.deepcopy(template_row._tr)
    target_row._tr.addprevious(new_tr)
    # 重新从table.rows里找到刚插入的这一行（对应新tr）
    for row in table.rows:
        if row._tr is new_tr:
            # 清空新行内容（避免复制了模板行的示例文字）
            for cell in row.cells:
                set_cell_text(cell, "")
            return row
    raise RuntimeError("插入新行失败")


def ensure_enough_rows(table, header_marker_row, stop_marker_text, need_count):
    """
    确保表格里"数据区"（表头行之后、stop_marker_text所在行之前）至少有need_count行空行。
    header_marker_row: 表头所在的row对象
    stop_marker_text: 数据区结束标志行的第一列文字（比如"检测仪器"）
    返回：数据区起始的那些行对象列表（长度 == need_count）
    """
    rows = list(table.rows)
    header_tr = header_marker_row._tr
    header_idx = next(i for i, row in enumerate(rows) if row._tr is header_tr)

    stop_row = None
    stop_idx = None
    for idx, row in enumerate(rows):
        if idx > header_idx and row.cells[0].text.strip() == stop_marker_text:
            stop_row = row
            stop_idx = idx
            break
    if stop_row is None:
        raise RuntimeError(f"未在表格中找到标志行: {stop_marker_text}")

    existing_data_rows = list(rows)[header_idx + 1: stop_idx]
    current_count = len(existing_data_rows)

    if current_count < need_count:
        template_row = existing_data_rows[-1] if existing_data_rows else None
        if template_row is None:
            raise RuntimeError("模板中没有可复制的空白数据行，无法自动增加行数。")
        for _ in range(need_count - current_count):
            new_row = clone_row_before(table, template_row, stop_row)
            existing_data_rows.append(new_row)

    return existing_data_rows[:need_count]


def build_conclusion_text(full_data):
    hot_names = []
    for row_data in full_data:
        diff = row_data['diff']
        if diff in (None, NOT_RECOGNIZED_TEXT):
            continue
        try:
            diff_val = float(diff)
        except ValueError:
            continue
        if diff_val > HOT_THRESHOLD:
            hot_names.append(str(row_data['device_name']).strip())

    if not hot_names:
        return CONCLUSION_NORMAL

    names_text = "、".join(hot_names)
    return CONCLUSION_TEMPLATE.format(names=names_text)


def step3_fill_word_data_table(doc, full_data):
    table = doc.tables[0]
    header_row = find_row_by_first_cell_text(table, "序号")
    if header_row is None:
        raise RuntimeError("在测温报告.docx第一个表格中未找到表头行（第一列内容为'序号'）。")

    data_rows = ensure_enough_rows(table, header_row, "检测仪器", len(full_data))

    for row_obj, row_data in zip(data_rows, full_data):
        values = [
            row_data['seq'],
            row_data['interval'],
            row_data['device_name'],
            row_data['defect_part'],
            row_data['surface_temp'] if row_data['surface_temp'] is not None else NOT_RECOGNIZED_TEXT,
            row_data['normal_temp'],
            row_data['env_temp'],
            row_data['load_current'],
            row_data['atlas_code'],
            row_data['remark'],
        ]
        for cell, val in zip(row_obj.cells, values):
            set_cell_text(cell, val)

    # 填写结论：table中"结论"所在行的第4~10列（索引3~9）
    conclusion_row = find_row_by_first_cell_text(table, "结论")
    if conclusion_row is None:
        print("警告：未找到'结论'所在行，跳过结论填写。")
    else:
        conclusion_text = build_conclusion_text(full_data)
        set_cell_text(conclusion_row.cells[3], conclusion_text)
        print(f"结论已生成: {conclusion_text}")


def step4_insert_images(doc, full_data, row_to_image):
    table = doc.tables[1]
    # 该表没有"结束标志行"，数据行一直延伸到表格末尾，所以直接在末尾追加新行即可
    rows = list(table.rows)
    current_count = len(rows) - 1  # 除表头外的现有行数
    need_count = len(full_data)

    if current_count < need_count:
        template_row = rows[-1]
        # 在表格末尾追加新行
        for _ in range(need_count - current_count):
            new_tr = copy.deepcopy(template_row._tr)
            table._tbl.append(new_tr)
        rows = list(table.rows)

    data_rows = rows[1:1 + need_count]

    for i, (row_obj, row_data) in enumerate(zip(data_rows, full_data), start=1):
        set_cell_text(row_obj.cells[0], f"图{i}")
        # 清空第二列已有内容
        cell = row_obj.cells[1]
        for p in cell.paragraphs:
            for run in list(p.runs):
                run.text = ""

        img_path = row_to_image.get(i - 1)
        if img_path and os.path.exists(img_path):
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM), height=Cm(IMAGE_HEIGHT_CM))
        else:
            set_cell_text(cell, "（未找到对应图片）")


def step5_generate_docx(full_data):
    print(f"\n正在生成测温报告: {DOCX_TEMPLATE_PATH}")
    doc = Document(DOCX_TEMPLATE_PATH)

    step3_fill_word_data_table(doc, full_data)

    image_files = collect_image_files(IMAGE_DIR)
    atlas_codes = [row['atlas_code'] for row in full_data]
    row_to_image = match_images_to_rows(image_files, atlas_codes)
    step4_insert_images(doc, full_data, row_to_image)

    apply_font_to_whole_document(doc)

    doc.save(DOCX_OUTPUT_PATH)
    print(f"测温报告已生成: {DOCX_OUTPUT_PATH}")


# ===================== 主程序 =====================

def main():
    if not os.path.isdir(IMAGE_DIR):
        print(f"错误：图谱文件夹不存在: {IMAGE_DIR}")
        sys.exit(1)
    if not os.path.isfile(EXCEL_PATH):
        print(f"错误：Excel文件不存在: {EXCEL_PATH}")
        sys.exit(1)
    if not os.path.isfile(DOCX_TEMPLATE_PATH):
        print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
        sys.exit(1)

    full_data = step2_process_excel()
    step5_generate_docx(full_data)

    print("\n全部完成！")
    print(f"  - Excel结果: {EXCEL_OUTPUT_PATH}")
    print(f"  - Word报告: {DOCX_OUTPUT_PATH}")
    if SAVE_DEBUG_CROPS:
        print(f"  - 调试截图: {DEBUG_DIR}")


if __name__ == "__main__":
    main()