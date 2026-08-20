# -*- coding: utf-8 -*-
"""
红外测温报告自动生成程序 V6
======================================================

V6 新增功能：
  1. PyInstaller 打包支持：exe 内嵌便携版 Tesseract OCR，用户无需安装任何依赖

V5 已有功能（全部保留）：
  1. 分步执行：--step1（仅OCR+核对+写Excel）/ --step2（仅生成Word报告）
     PyCharm 直接运行弹出交互式菜单，无需记命令
  2. 输出文件自动递增编号，不再覆盖旧版本：
     - 必填信息_已生成_1.xlsx, 必填信息_已生成_2.xlsx ...
     - 测温报告_已生成_1.docx, 测温报告_已生成_2.docx ...
  3. step2 自动读取编号最大的 Excel 文件

工作流程：
  Step 1: OCR识别表面温度 → 人工核对 → 计算正常温度/温差 → 写入Excel
  Step 2: 读取Excel → 生成测温报告.docx（含数据表、结论、图片）

文件结构（与脚本同目录）：
  ./图谱/                      红外测温图片
  ./必填信息.xlsx               数据表模板（工作表名："普测"）
  ./测温报告.docx                Word报告模板
  ./必填信息_已生成_N.xlsx       step1 输出（自动编号）
  ./测温报告_已生成_N.docx       step2 输出（自动编号）

依赖安装：
  pip install opencv-python numpy pytesseract openpyxl pillow python-docx
  brew install tesseract  或  conda install -c conda-forge tesseract

使用方法：
  python3 红外测温报告自动生成-claude-v5.py              完整运行
  python3 红外测温报告自动生成-claude-v5.py --step1       仅 step1
  python3 红外测温报告自动生成-claude-v5.py --step2       仅 step2
  python3 红外测温报告自动生成-claude-v5.py --cli         命令行核对模式
"""

import os
import re
import sys
import json
import glob
import copy
import time
import platform
import subprocess
import datetime
import argparse

import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook
from openpyxl.styles import PatternFill

from docx import Document
from docx.shared import Cm
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ===================== 配置区域 =====================

BASE_DIR = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))

IMAGE_DIR = os.path.join(BASE_DIR, "图谱")
EXCEL_PATH = os.path.join(BASE_DIR, "必填信息.xlsx")
EXCEL_OUTPUT_BASE = os.path.join(BASE_DIR, "必填信息_已生成")
EXCEL_SHEET_NAME = "普测"

DOCX_TEMPLATE_PATH = os.path.join(BASE_DIR, "测温报告.docx")
DOCX_OUTPUT_BASE = os.path.join(BASE_DIR, "测温报告_已生成")

# ===================== Tesseract 路径自适应 =====================
import sys as _sys
if getattr(_sys, 'frozen', False):
    _tesseract_dir = os.path.join(os.path.dirname(_sys.executable), 'tesseract')
else:
    _tesseract_dir = os.path.join(BASE_DIR, 'tesseract')
_tesseract_exe = os.path.join(_tesseract_dir, 'tesseract.exe')
if os.path.exists(_tesseract_exe):
    if hasattr(os, 'add_dll_directory'):
        os.add_dll_directory(_tesseract_dir)
    os.environ['PATH'] = _tesseract_dir + os.pathsep + os.environ.get('PATH', '')
    os.environ['TESSDATA_PREFIX'] = os.path.join(_tesseract_dir, 'tessdata')
    os.environ['TMP'] = _tesseract_dir
    os.environ['TEMP'] = _tesseract_dir
    pytesseract.pytesseract.tesseract_cmd = _tesseract_exe

ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_infrared.json")
SAVE_DEBUG_CROPS = True
DEBUG_DIR = os.path.join(BASE_DIR, "debug_crops")
VERBOSE = True

NOT_RECOGNIZED_TEXT = "无法识别"

# 发热判定阈值：温差 > 该值 视为发热对象
HOT_THRESHOLD = 15

# 人工核对相关：L列放调试图片超链接；"无法识别"的单元格用红色底色高亮
DEBUG_LINK_COLUMN = 12  # L列
RED_FILL = PatternFill(start_color="FFC7CE", end_color="FFC7CE", fill_type="solid")

# 结论模板
CONCLUSION_TEMPLATE = "{names}发热,达到一般缺陷标准,未达到严重缺陷标准（发热缺陷较前期无明显发展趋势），无其他异常。"
CONCLUSION_NORMAL = "无异常"

# 图片插入尺寸
IMAGE_WIDTH_CM = 12
IMAGE_HEIGHT_CM = 9

# =================================================================


# ===================== 工具函数：自然排序 =====================

def natural_sort_key(s):
    """自然排序：让"图2"排在"图10"前面，而不是按字符串排序排到后面。"""
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def collect_image_files(image_dir):
    patterns = ["*.jpg", "*.JPG", "*.jpeg", "*.JPEG", "*.png", "*.PNG"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(image_dir, p)))
    files = sorted(set(files), key=lambda p: natural_sort_key(os.path.basename(p)))
    return files


def resolve_output_path(base_path, ext, suffix=None):
    """生成不冲突的输出路径（Excel / Word 通用）。"""
    if suffix:
        return f"{base_path}_{suffix}.{ext}"
    counter = 1
    while True:
        path = f"{base_path}_{counter}.{ext}"
        if not os.path.exists(path):
            return path
        counter += 1


def find_latest_output(base_path, ext):
    """找到编号最大的已有输出文件，供 step2 自动读取。"""
    pattern = f"{base_path}_*.{ext}"
    files = sorted(glob.glob(pattern))
    if not files:
        return None
    return sorted(files, key=lambda p: natural_sort_key(p))[-1]


# ===================== 第一步：OCR识别表面温度 =====================

def imread_unicode(path):
    """cv2.imread 的 Unicode 安全替代（Windows 中文路径兼容）。"""
    img_array = np.fromfile(path, dtype=np.uint8)
    return cv2.imdecode(img_array, cv2.IMREAD_COLOR)


def select_roi_interactively(sample_image_path):
    """自定义鼠标框选ROI区域（细线矩形），空格/回车确认，c键取消重选。"""
    img = imread_unicode(sample_image_path)
    if img is None:
        raise FileNotFoundError(f"无法读取示例图片: {sample_image_path}")
    img_h, img_w = img.shape[:2]

    state = {"drawing": False, "x0": -1, "y0": -1, "x1": -1, "y1": -1}

    def on_mouse(event, x, y, flags, param):
        if event == cv2.EVENT_LBUTTONDOWN:
            state["drawing"] = True
            state["x0"], state["y0"] = x, y
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_MOUSEMOVE and state["drawing"]:
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_LBUTTONUP:
            state["drawing"] = False
            state["x1"], state["y1"] = x, y

    window_title = "框选表面温度区域 (Space/Enter确认, C取消)"
    cv2.namedWindow(window_title)
    cv2.setMouseCallback(window_title, on_mouse)
    display = img.copy()

    print(f"\n请在弹出的窗口中用鼠标左键拖拽框选【表面温度数值】区域。")
    print("  空格/回车=确认   c=取消重选   Esc=退出\n")

    while True:
        display = img.copy()
        x0, y0 = state["x0"], state["y0"]
        x1, y1 = state["x1"], state["y1"]
        if x0 >= 0 and y0 >= 0:
            cv2.rectangle(display, (x0, y0), (x1, y1), (0, 255, 0), 1)
        cv2.imshow(window_title, display)
        key = cv2.waitKey(20) & 0xFF
        if key == 13 or key == 32:  # Enter or Space
            break
        elif key == ord('c') or key == ord('C') or key == 27:  # c or Esc
            state["x0"] = state["y0"] = state["x1"] = state["y1"] = -1
            state["drawing"] = False
            if key == 27:
                break

    cv2.destroyAllWindows()
    cv2.waitKey(1)
    time.sleep(0.3)

    x0, y0 = state["x0"], state["y0"]
    x1, y1 = state["x1"], state["y1"]
    if x0 < 0 or y0 < 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    x = min(x0, x1)
    y = min(y0, y1)
    w = abs(x1 - x0)
    h = abs(y1 - y0)
    if w == 0 or h == 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    rx, ry, rw, rh = x / img_w, y / img_h, w / img_w, h / img_h
    print(f"框选区域（像素）: x={x}, y={y}, w={w}, h={h}  (示例图片尺寸: {img_w}x{img_h})")
    return rx, ry, rw, rh


def get_roi(sample_image_path):
    if os.path.exists(ROI_CONFIG_FILE):
        with open(ROI_CONFIG_FILE, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        rx, ry, rw, rh = cfg["rx"], cfg["ry"], cfg["rw"], cfg["rh"]
        print(f"检测到已保存的识别区域配置（相对比例）: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
        choice = input("是否复用该区域？(y=复用 / n=重新框选) [y]: ").strip().lower()
        if choice in ("", "y", "yes"):
            return rx, ry, rw, rh

    rx, ry, rw, rh = select_roi_interactively(sample_image_path)
    with open(ROI_CONFIG_FILE, "w", encoding="utf-8") as f:
        json.dump({"rx": rx, "ry": ry, "rw": rw, "rh": rh}, f, ensure_ascii=False, indent=2)
    print(f"识别区域已保存到: {ROI_CONFIG_FILE}")
    return rx, ry, rw, rh


def crop_by_relative_roi(img, roi_ratio):
    rx, ry, rw, rh = roi_ratio
    img_h, img_w = img.shape[:2]
    x = int(round(rx * img_w))
    y = int(round(ry * img_h))
    w = int(round(rw * img_w))
    h = int(round(rh * img_h))
    x2, y2 = min(x + w, img_w), min(y + h, img_h)
    x, y = max(x, 0), max(y, 0)
    return img[y:y2, x:x2]


def preprocess_variants(crop):
    scale = 4
    crop_big = cv2.resize(crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(crop_big, cv2.COLOR_BGR2GRAY)

    variants = []

    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh1 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh1)

    inv = cv2.bitwise_not(blur)
    _, thresh2 = cv2.threshold(inv, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh2)

    thresh3 = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
    )
    variants.append(thresh3)

    return variants


def extract_number_signed(text):
    """提取数字（保留正负号，温度不取绝对值），返回字符串或None。"""
    if not text:
        return None
    text = text.strip()
    match = re.search(r'-?\d+\.?\d*', text)
    if not match:
        return None
    try:
        num = float(match.group())
        if num == int(num):
            return str(int(num))
        return str(num)
    except ValueError:
        return None


def ocr_extract_temperature(image_path, roi_ratio):
    img = imread_unicode(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"], None

    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, ["[裁剪区域为空，请检查ROI设置]"], None

    variants = preprocess_variants(crop)
    configs = [
        r'--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 11 -c tessedit_char_whitelist=0123456789.-',
    ]

    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            value = extract_number_signed(text)
            if value is not None:
                return value, variant, raw_texts, crop

    return None, variants[0], raw_texts, crop


def match_images_to_rows(image_files, atlas_codes):
    """
    将图片文件和Excel行的"图谱编号"进行匹配。
    优先按文件名（不含扩展名）精确匹配；如果匹配率过低，改为按顺序位置匹配。
    返回：{行索引(0开始): 图片路径 或 None}
    """
    name_to_path = {os.path.splitext(os.path.basename(p))[0]: p for p in image_files}

    name_matched = 0
    for code in atlas_codes:
        if code is not None and str(code).strip() in name_to_path:
            name_matched += 1

    use_name_matching = (len(atlas_codes) > 0 and name_matched >= len(atlas_codes) * 0.8)

    result = {}
    if use_name_matching:
        print("图片匹配方式：按文件名与Excel'图谱编号'精确匹配。")
        for i, code in enumerate(atlas_codes):
            code_str = str(code).strip() if code is not None else ""
            result[i] = name_to_path.get(code_str)
    else:
        print("图片匹配方式：文件名与'图谱编号'对不上，改为按【图片文件名顺序】与【Excel行顺序】依次对应。")
        for i, code in enumerate(atlas_codes):
            result[i] = image_files[i] if i < len(image_files) else None

    return result


def step1_ocr_fill_surface_temperature(ws, header_row_idx, data_rows_info):
    """
    识别表面温度并写入E列。
    data_rows_info: list of dict，每个元素至少包含 {'row': excel行号, 'atlas_code': 图谱编号}
    返回：
        surface_temp: {行号: 表面温度值字符串 或 None}
        debug_file_map: {行号: 调试图片文件名 或 None}（供人工核对用的原始彩色裁剪图）
    """
    image_files = collect_image_files(IMAGE_DIR)
    if not image_files:
        print(f"警告：图谱文件夹中未找到任何图片: {IMAGE_DIR}")
        return {info['row']: None for info in data_rows_info}, {}

    print(f"图谱文件夹共找到 {len(image_files)} 张图片。")

    atlas_codes = [info['atlas_code'] for info in data_rows_info]
    row_to_image = match_images_to_rows(image_files, atlas_codes)

    # 用第一个有效图片作为框选样本
    sample_path = next((p for p in row_to_image.values() if p), image_files[0])
    roi_ratio = get_roi(sample_path)

    if SAVE_DEBUG_CROPS:
        os.makedirs(DEBUG_DIR, exist_ok=True)

    surface_temp = {}
    debug_file_map = {}
    print("\n开始批量识别表面温度，请稍候...\n")
    for i, info in enumerate(data_rows_info):
        row = info['row']
        img_path = row_to_image.get(i)
        if img_path is None:
            surface_temp[row] = None
            debug_file_map[row] = None
            print(f"[第{row}行] 图谱编号={info['atlas_code']}  ->  未找到对应图片，跳过")
            continue

        value, processed, raw_texts, raw_crop = ocr_extract_temperature(img_path, roi_ratio)
        surface_temp[row] = value

        status = value if value is not None else NOT_RECOGNIZED_TEXT
        print(f"[第{row}行] {os.path.basename(img_path)}  ->  {status}")

        if VERBOSE and value is None:
            cleaned = [t if t else "(空)" for t in raw_texts]
            print(f"      OCR原始识别文本（各模式）: {cleaned}")

        debug_file_map[row] = None
        if SAVE_DEBUG_CROPS and raw_crop is not None:
            debug_name = f"row{row}_{status}.png"
            debug_name = re.sub(r'[\\/:*?"<>|]', "_", debug_name)
            # 保存"原始彩色裁剪图"（放大2倍方便肉眼看清），而不是二值化后的调试图，
            # 这样人工核对时看到的和图片本身长得一样，不容易被二值化处理误导。
            save_img = cv2.resize(raw_crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
            cv2.imwrite(os.path.join(DEBUG_DIR, debug_name), save_img)
            debug_file_map[row] = debug_name

    return surface_temp, debug_file_map


# ===================== 第二、三步：正常温度 & 温差计算 =====================

PHASE_SUFFIX_RE = re.compile(r'^(.*)(A相|B相|C相)$')


def compute_normal_temperature_and_diff(data_rows_info, surface_temp):
    """
    根据分组规则计算正常温度(F)和温差(K)。
    返回：{行号: (正常温度值 或 None 或 '无法识别', 温差值 或 None 或 '无法识别')}
    """
    # 按前缀分组（仅针对以 A相/B相/C相 结尾的设备名称）
    prefix_groups = {}
    for info in data_rows_info:
        name = str(info['device_name']).strip() if info['device_name'] else ""
        m = PHASE_SUFFIX_RE.match(name)
        if m:
            prefix = m.group(1)
            prefix_groups.setdefault(prefix, []).append(info['row'])

    # 只有恰好凑齐3个的前缀才算一组，其余的都当作单独一行处理
    valid_groups = {prefix: rows for prefix, rows in prefix_groups.items() if len(rows) == 3}
    grouped_rows = set(r for rows in valid_groups.values() for r in rows)

    result = {}

    # 先处理分组（3个一组）
    for prefix, rows in valid_groups.items():
        values = []
        ok = True
        for r in rows:
            v = surface_temp.get(r)
            if v is None:
                ok = False
                break
            try:
                values.append(float(v))
            except ValueError:
                ok = False
                break
        if ok:
            normal_temp = min(values)
            for r in rows:
                e_val = float(surface_temp[r])
                diff = round(e_val - normal_temp, 2)
                result[r] = (format_number(normal_temp), format_number(diff))
        else:
            for r in rows:
                result[r] = (NOT_RECOGNIZED_TEXT, NOT_RECOGNIZED_TEXT)

    # 再处理单独成组的行
    for info in data_rows_info:
        r = info['row']
        if r in grouped_rows:
            continue
        v = surface_temp.get(r)
        if v is None:
            result[r] = (NOT_RECOGNIZED_TEXT, NOT_RECOGNIZED_TEXT)
        else:
            e_val = float(v)
            normal_temp = e_val
            diff = 0.0
            result[r] = (format_number(normal_temp), format_number(diff))

    return result


def format_number(num):
    """把浮点数格式化成字符串：整数不带小数点，否则保留原有精度（最多2位小数）。"""
    if num == int(num):
        return str(int(num))
    return str(round(num, 2))


# ===================== Excel 读写 =====================

def read_excel_data_rows(ws, header_row_idx):
    """从表头下一行开始读取数据，直到设备名称列(C)为空为止。"""
    data_rows_info = []
    row = header_row_idx + 1
    while True:
        c_val = ws.cell(row=row, column=3).value  # C列 设备名称
        if c_val is None or str(c_val).strip() == "":
            break
        info = {
            'row': row,
            'seq': ws.cell(row=row, column=1).value,       # A 序号
            'interval': ws.cell(row=row, column=2).value,   # B 间隔名称
            'device_name': c_val,                             # C 设备名称
            'defect_part': ws.cell(row=row, column=4).value,  # D 缺陷部位
            'env_temp': ws.cell(row=row, column=7).value,     # G 环境温度
            'load_current': ws.cell(row=row, column=8).value, # H 负荷电流
            'atlas_code': ws.cell(row=row, column=9).value,   # I 图谱编号
            'remark': ws.cell(row=row, column=10).value,      # J 备注
        }
        data_rows_info.append(info)
        row += 1
    return data_rows_info


def find_header_row(ws, keyword="序号", max_scan_rows=10):
    for row in range(1, max_scan_rows + 1):
        if str(ws.cell(row=row, column=1).value).strip() == keyword:
            return row
    return 1


def write_review_helpers(ws, header_row_idx, data_rows_info, surface_temp, debug_file_map):
    """写入L列调试图片超链接，并把仍然'无法识别'的E列单元格标红（供事后复查留档用）。"""
    header_cell = ws.cell(row=header_row_idx, column=DEBUG_LINK_COLUMN)
    if not header_cell.value:
        header_cell.value = "调试图片（点击核对）"

    debug_rel_dir = os.path.relpath(DEBUG_DIR, BASE_DIR)

    for info in data_rows_info:
        r = info['row']
        e_cell = ws.cell(row=r, column=5)

        if surface_temp.get(r) is None:
            e_cell.fill = RED_FILL
        else:
            e_cell.fill = PatternFill(fill_type=None)

        link_cell = ws.cell(row=r, column=DEBUG_LINK_COLUMN)
        debug_name = debug_file_map.get(r)
        if debug_name:
            rel_path = os.path.join(debug_rel_dir, debug_name)
            link_cell.value = "查看裁剪图"
            link_cell.hyperlink = rel_path
            link_cell.style = "Hyperlink"
        else:
            link_cell.value = "（无对应图片）"


def open_image_for_preview(image_path):
    """用系统默认图片查看器打开图片（Mac用'预览'，Windows/Linux也做了兼容）。"""
    try:
        system = platform.system()
        if system == "Darwin":
            subprocess.run(["open", image_path], check=False)
        elif system == "Windows":
            os.startfile(image_path)  # noqa
        else:
            subprocess.run(["xdg-open", image_path], check=False)
    except Exception as e:
        print(f"  （提示：自动打开图片失败：{e}，可手动到以下路径查看：{image_path}）")


def interactive_review_cli(data_rows_info, surface_temp, debug_file_map):
    """
    命令行版人工核对（图形界面不可用时的兜底方案）。
      - 直接回车：确认当前识别值无误，进入下一条
      - 输入数字：修正当前行数值
      - 输入 s（或 skip）：跳过剩余所有行的核对，剩余行保留原有识别结果
    返回：更新后的 surface_temp（可能已被人工修正）
    """
    reviewed = dict(surface_temp)
    total = len(data_rows_info)

    print("\n" + "=" * 60)
    print("进入人工核对环节（命令行模式）：每一行会自动弹出对应的裁剪图片")
    print("  - 数值正确：直接按【回车】确认，自动进入下一条")
    print("  - 数值错误：直接输入正确数值后按【回车】修正")
    print("  - 想跳过剩余所有行的核对：输入 s 后回车（剩余行保留识别结果）")
    print("=" * 60)

    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        current = reviewed.get(r)
        debug_name = debug_file_map.get(r)

        status_display = current if current is not None else NOT_RECOGNIZED_TEXT
        mark = "   ⚠️ 识别失败，需手动输入" if current is None else ""

        print(f"\n[{i}/{total}] 设备: {info['device_name']}  图谱编号: {info['atlas_code']}  "
              f"当前识别值: {status_display}{mark}")

        if debug_name and os.path.exists(os.path.join(DEBUG_DIR, debug_name)):
            open_image_for_preview(os.path.join(DEBUG_DIR, debug_name))
        else:
            print("  （未找到对应裁剪图片，请自行核实）")

        while True:
            user_input = input("  回车确认 / 输入正确数值修正 / 输入 s 跳过剩余核对: ").strip()
            if user_input == "":
                break  # 保留当前值
            if user_input.lower() in ("s", "skip"):
                print("\n已跳过剩余行的核对，剩余行保留原有识别结果。")
                return reviewed
            try:
                float(user_input)
                reviewed[r] = user_input
                break
            except ValueError:
                print("  输入的不是有效数字，请重新输入（或直接回车保留原值，或输入 s 跳过剩余核对）。")

    print("\n人工核对完成。")
    return reviewed


def _load_image_into(label, debug_filename, ref_dict, ref_key, max_w=580, max_h=320):
    """加载裁剪图片到 tkinter Label，适配 43号文的 debug_file_map（仅存文件名）。"""
    try:
        from PIL import Image, ImageTk
    except ImportError:
        return
    img_path = os.path.join(DEBUG_DIR, debug_filename) if debug_filename else None
    if img_path and os.path.exists(img_path):
        pil_img = Image.open(img_path)
        ratio = min(max_w / pil_img.width, max_h / pil_img.height, 1.0)
        new_size = (max(1, int(pil_img.width * ratio)), max(1, int(pil_img.height * ratio)))
        pil_img = pil_img.resize(new_size)
        photo = ImageTk.PhotoImage(pil_img)
        ref_dict[ref_key] = photo
        label.config(image=photo, text="")
    else:
        ref_dict[ref_key] = None
        label.config(image="", text="（未找到对应裁剪图片）")


def interactive_review_gui(data_rows_info, surface_temp, debug_file_map):
    """
    图形界面人工核对：
    - ↑/↓ 键或按钮自由翻阅，每翻页自动保存输入框内容
    - 最后一条按「确认提交（回车）」才最终提交
    - 如果当前环境没有tkinter/Pillow，自动回退到命令行版本
    """
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（提示：图形界面核对工具不可用（{e}），自动改用命令行方式核对。）")
        return interactive_review_cli(data_rows_info, surface_temp, debug_file_map)

    reviewed = dict(surface_temp)
    total = len(data_rows_info)
    state = {'index': 0, 'confirmed': False}
    photo_refs = {}

    root = tk.Tk()
    root.title("红外测温 - 人工核对")
    root.geometry("620x520")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 14), justify="left", wraplength=580)
    info_label.pack(pady=(15, 5))

    img_label = tk.Label(root)
    img_label.pack(pady=8)

    entry_var = tk.StringVar()
    entry = tk.Entry(root, textvariable=entry_var, font=("PingFang SC", 22), justify="center", width=14)
    entry.pack(pady=8)

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(0, 5))

    btn_ref = {}

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']
        current = reviewed.get(r)
        status = current if current is not None else ""
        mark = "   ⚠️ 识别失败，需手动输入" if current is None else ""
        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}{mark}")
        _load_image_into(img_label, debug_file_map.get(r), photo_refs, 'crop')
        entry_var.set(status)
        hint_label.config(text="↑ = 上一条    |    ↓ = 下一条    |    回车 = 确认提交", fg="gray")
        entry.focus_set()
        entry.select_range(0, tk.END)
        update_buttons()

    def update_buttons():
        idx = state['index']
        is_first = (idx == 0)
        is_last = (idx + 1 >= total)
        btn_ref['prev'].config(state=(tk.DISABLED if is_first else tk.NORMAL))
        btn_ref['next'].config(state=(tk.DISABLED if is_last else tk.NORMAL))

    def try_save():
        idx = state['index']
        r = data_rows_info[idx]['row']
        val = entry_var.get().strip()
        if val != "":
            try:
                float(val)
                reviewed[r] = val
            except ValueError:
                hint_label.config(text="⚠️ 不是有效数字，请重新输入！", fg="red")
                return False
        return True

    def go_previous(event=None):
        if not try_save():
            return
        if state['index'] == 0:
            return
        state['index'] -= 1
        load_row(state['index'])

    def go_next(event=None):
        if not try_save():
            return
        if state['index'] + 1 >= total:
            hint_label.config(text="已是最后一条，请按回车提交", fg="orange")
            return
        state['index'] += 1
        load_row(state['index'])

    def confirm_submit(event=None):
        if not try_save():
            return
        state['confirmed'] = True
        root.destroy()

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=15)
    btn_ref['prev'] = tk.Button(button_frame, text="↑ 上一条", command=go_previous, width=10)
    btn_ref['prev'].grid(row=0, column=0, padx=4)
    btn_ref['next'] = tk.Button(button_frame, text="↓ 下一条", command=go_next, width=10)
    btn_ref['next'].grid(row=0, column=1, padx=4)
    btn_ref['submit'] = tk.Button(button_frame, text="确认提交（回车）", command=confirm_submit, width=16, fg="blue")
    btn_ref['submit'].grid(row=0, column=2, padx=4)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=3, padx=4)

    root.bind('<Up>', go_previous)
    root.bind('<Down>', go_next)
    root.bind('<Return>', confirm_submit)

    print("\n已打开图形界面核对窗口（↑上一条 | ↓下一条 | 回车=确认提交）...")
    load_row(0)
    root.lift()
    root.focus_force()
    entry.focus_set()
    root.mainloop()

    if state['confirmed']:
        print("人工核对完成，用户已确认提交。")
    else:
        print("人工核对已跳过。")
    return reviewed


def step1_ocr_and_review(excel_output_path):
    """
    Step 1: OCR识别表面温度 -> 人工核对（可跳过）-> 计算正常温度/温差 -> 保存Excel。
    返回：full_data，供后续生成Word报告使用。
    """
    print(f"\n{'=' * 50}")
    print("  Step 1: OCR识别 + 人工核对 + 写入Excel")
    print(f"{'=' * 50}\n")

    print(f"正在读取Excel模板: {EXCEL_PATH}  (工作表: {EXCEL_SHEET_NAME})")
    wb = load_workbook(EXCEL_PATH)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        print(f"错误：Excel中找不到名为'{EXCEL_SHEET_NAME}'的工作表。")
        sys.exit(1)
    ws = wb[EXCEL_SHEET_NAME]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行，请检查表格内容。")
        sys.exit(1)
    print(f"共读取到 {len(data_rows_info)} 行数据。")

    # 第一步：OCR识别
    surface_temp, debug_file_map = step1_ocr_fill_surface_temperature(ws, header_row_idx, data_rows_info)

    fail_count = sum(1 for v in surface_temp.values() if v is None)

    # 第二步：人工核对（如果全部识别成功，询问是否仍要核对；否则强制核对）
    do_review = True
    if fail_count == 0:
        print(f"\n本次 {len(data_rows_info)} 行全部识别成功。")
        choice = input("是否仍要逐张人工核对确认？(y=核对 / n=跳过，直接使用识别结果) [n]: ").strip().lower()
        do_review = choice in ("y", "yes")
    else:
        print(f"\n共有 {fail_count} / {len(data_rows_info)} 行识别失败，需要人工核对补充数值。")

    if do_review:
        if USE_GUI:
            surface_temp = interactive_review_gui(data_rows_info, surface_temp, debug_file_map)
        else:
            print("\n（已启用 --cli 模式，使用命令行核对）")
            surface_temp = interactive_review_cli(data_rows_info, surface_temp, debug_file_map)

    # 第三步：写入E列
    for info in data_rows_info:
        r = info['row']
        v = surface_temp.get(r)
        ws.cell(row=r, column=5).value = v if v is not None else NOT_RECOGNIZED_TEXT

    # 第四步：从Word模板读取环境温度（第1个表格，第5行第5列）并写入G列
    env_temp_doc = Document(DOCX_TEMPLATE_PATH)
    try:
        env_temp_text = env_temp_doc.tables[0].rows[4].cells[4].text.strip()
        # 保留原始文本，去除可能的"℃"后缀，避免 float 转换后变成 31.0 而非 31
        env_temp_display = env_temp_text.replace("℃", "").strip()
        float(env_temp_display)  # 仅验证是否为有效数字
        env_temp_val = env_temp_display
        print(f"\n从测温报告模板读取环境温度: {env_temp_display}")
    except (ValueError, IndexError, AttributeError) as e:
        print(f"\n警告：从测温报告读取环境温度失败（{e}），G列将留空。")
        env_temp_val = None

    if env_temp_val is not None:
        for info in data_rows_info:
            r = info['row']
            ws.cell(row=r, column=7).value = env_temp_val

    # 第五步：计算正常温度(F)和温差(K)
    calc_result = compute_normal_temperature_and_diff(data_rows_info, surface_temp)

    k_header = ws.cell(row=header_row_idx, column=11).value
    if not k_header:
        ws.cell(row=header_row_idx, column=11).value = "温差（℃）"

    for info in data_rows_info:
        r = info['row']
        normal_temp, diff = calc_result[r]
        ws.cell(row=r, column=6).value = normal_temp
        ws.cell(row=r, column=11).value = diff

    # 留档：L列超链接 + 仍未识别的行标红，方便事后复查
    write_review_helpers(ws, header_row_idx, data_rows_info, surface_temp, debug_file_map)

    wb.save(excel_output_path)
    print(f"\nExcel已生成: {excel_output_path}")

    full_data = []
    for info in data_rows_info:
        r = info['row']
        row_data = dict(info)
        row_data['surface_temp'] = surface_temp.get(r)
        row_data['normal_temp'], row_data['diff'] = calc_result[r]
        row_data['env_temp'] = env_temp_val  # 补到内存中，确保串行模式 step2 也能获取
        full_data.append(row_data)

    return full_data


# ===================== 从已生成Excel读回数据（step2 复用） =====================

def read_excel_output(excel_output_path=None):
    """从 step1 生成的最新 Excel 中读回完整数据，供 step2 使用。"""
    if excel_output_path is None:
        excel_output_path = find_latest_output(EXCEL_OUTPUT_BASE, "xlsx")
    if excel_output_path is None:
        print(f"错误：找不到任何 必填信息_已生成_*.xlsx 文件。")
        print("  请先运行 Step 1 生成识别结果。")
        sys.exit(1)

    print(f"\n自动读取最新 Excel: {os.path.basename(excel_output_path)}")

    wb = load_workbook(excel_output_path)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        print(f"错误：工作表中找不到 '{EXCEL_SHEET_NAME}'。")
        sys.exit(1)
    ws = wb[EXCEL_SHEET_NAME]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行。")
        sys.exit(1)

    full_data = []
    for info in data_rows_info:
        r = info['row']
        e_val = ws.cell(row=r, column=5).value   # 表面温度
        f_val = ws.cell(row=r, column=6).value   # 正常温度
        k_val = ws.cell(row=r, column=11).value  # 温差

        surface_temp = e_val if e_val != NOT_RECOGNIZED_TEXT else None
        # 尝试转为 float
        try:
            surface_temp = float(surface_temp) if surface_temp is not None else None
        except (ValueError, TypeError):
            surface_temp = surface_temp

        row_data = dict(info)
        row_data['surface_temp'] = surface_temp
        row_data['normal_temp'] = f_val
        row_data['diff'] = k_val
        full_data.append(row_data)

    print(f"  共读取 {len(full_data)} 行数据。")
    return full_data


# ===================== 第四步：生成 测温报告.docx =====================

CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"


def apply_font(run, east_asian_font=CHINESE_FONT, latin_font=WESTERN_FONT):
    """设置字体：中文用宋体，英文/数字用Times New Roman。"""
    run.font.name = latin_font
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.append(rFonts)
    rFonts.set(qn('w:eastAsia'), east_asian_font)
    rFonts.set(qn('w:ascii'), latin_font)
    rFonts.set(qn('w:hAnsi'), latin_font)


def set_cell_text(cell, text, keep_first_run_format=True):
    """清空单元格已有内容，写入新文本，并统一设置字体（中文宋体/英文数字Times New Roman）。"""
    text = "" if text is None else str(text)
    paragraphs = cell.paragraphs
    p = paragraphs[0]
    # 清空该段落已有的runs
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    apply_font(run)
    # 清空多余段落
    for extra_p in paragraphs[1:]:
        extra_p.text = ""


def apply_font_to_paragraph(paragraph):
    for run in paragraph.runs:
        apply_font(run)


def apply_font_to_table(table):
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                apply_font_to_paragraph(p)
            # 兼容单元格内嵌套表格的情况
            for nested_table in cell.tables:
                apply_font_to_table(nested_table)


def apply_font_to_whole_document(doc):
    """把整篇文档（包括模板里原有的静态文字）统一设置为：中文宋体，英文/数字Times New Roman。"""
    for p in doc.paragraphs:
        apply_font_to_paragraph(p)
    for t in doc.tables:
        apply_font_to_table(t)


def visual_cells(row):
    """按底层XML元素身份去重（正确识别真实的合并单元格），返回去重后的Cell对象列表。"""
    result = []
    last_tc = None
    for cell in row.cells:
        if cell._tc is not last_tc:
            result.append(cell)
            last_tc = cell._tc
    return result


def find_row_by_first_cell_text(table, target_text):
    for row in table.rows:
        if row.cells[0].text.strip() == target_text:
            return row
    return None


def clone_row_before(table, template_row, target_row):
    """复制template_row的格式，插入到target_row之前，返回新行对象。"""
    new_tr = copy.deepcopy(template_row._tr)
    target_row._tr.addprevious(new_tr)
    # 重新从table.rows里找到刚插入的这一行（对应新tr）
    for row in table.rows:
        if row._tr is new_tr:
            # 清空新行内容（避免复制了模板行的示例文字）
            for cell in row.cells:
                set_cell_text(cell, "")
            return row
    raise RuntimeError("插入新行失败")


def ensure_enough_rows(table, header_marker_row, stop_marker_text, need_count):
    """
    确保表格里"数据区"（表头行之后、stop_marker_text所在行之前）至少有need_count行空行。
    header_marker_row: 表头所在的row对象
    stop_marker_text: 数据区结束标志行的第一列文字（比如"检测仪器"）
    返回：数据区起始的那些行对象列表（长度 == need_count）
    """
    rows = list(table.rows)
    header_tr = header_marker_row._tr
    header_idx = next(i for i, row in enumerate(rows) if row._tr is header_tr)

    stop_row = None
    stop_idx = None
    for idx, row in enumerate(rows):
        if idx > header_idx and row.cells[0].text.strip() == stop_marker_text:
            stop_row = row
            stop_idx = idx
            break
    if stop_row is None:
        raise RuntimeError(f"未在表格中找到标志行: {stop_marker_text}")

    existing_data_rows = list(rows)[header_idx + 1: stop_idx]
    current_count = len(existing_data_rows)

    if current_count < need_count:
        template_row = existing_data_rows[-1] if existing_data_rows else None
        if template_row is None:
            raise RuntimeError("模板中没有可复制的空白数据行，无法自动增加行数。")
        for _ in range(need_count - current_count):
            new_row = clone_row_before(table, template_row, stop_row)
            existing_data_rows.append(new_row)

    return existing_data_rows[:need_count]


def build_conclusion_text(full_data):
    hot_names = []
    for row_data in full_data:
        diff = row_data['diff']
        if diff in (None, NOT_RECOGNIZED_TEXT):
            continue
        try:
            diff_val = float(diff)
        except ValueError:
            continue
        if diff_val > HOT_THRESHOLD:
            hot_names.append(str(row_data['device_name']).strip())

    if not hot_names:
        return CONCLUSION_NORMAL

    names_text = "、".join(hot_names)
    return CONCLUSION_TEMPLATE.format(names=names_text)


def step3_fill_word_data_table(doc, full_data):
    table = doc.tables[0]
    header_row = find_row_by_first_cell_text(table, "序号")
    if header_row is None:
        raise RuntimeError("在测温报告.docx第一个表格中未找到表头行（第一列内容为'序号'）。")

    data_rows = ensure_enough_rows(table, header_row, "检测仪器", len(full_data))

    for row_obj, row_data in zip(data_rows, full_data):
        values = [
            row_data['seq'],
            row_data['interval'],
            row_data['device_name'],
            row_data['defect_part'],
            row_data['surface_temp'] if row_data['surface_temp'] is not None else NOT_RECOGNIZED_TEXT,
            row_data['normal_temp'],
            row_data['env_temp'],
            row_data['load_current'],
            row_data['atlas_code'],
            row_data['remark'],
        ]
        for cell, val in zip(visual_cells(row_obj), values):
            set_cell_text(cell, val)

    # 填写结论：table中"结论"所在行的第4~10列（索引3~9）
    conclusion_row = find_row_by_first_cell_text(table, "结论")
    if conclusion_row is None:
        print("警告：未找到'结论'所在行，跳过结论填写。")
    else:
        conclusion_text = build_conclusion_text(full_data)
        set_cell_text(conclusion_row.cells[3], conclusion_text)
        print(f"结论已生成: {conclusion_text}")


def step4_insert_images(doc, full_data, row_to_image):
    table = doc.tables[1]
    # 该表没有"结束标志行"，数据行一直延伸到表格末尾，所以直接在末尾追加新行即可
    rows = list(table.rows)
    current_count = len(rows) - 1  # 除表头外的现有行数
    need_count = len(full_data)

    if current_count < need_count:
        template_row = rows[-1]
        # 在表格末尾追加新行
        for _ in range(need_count - current_count):
            new_tr = copy.deepcopy(template_row._tr)
            table._tbl.append(new_tr)
        rows = list(table.rows)

    data_rows = rows[1:1 + need_count]

    for i, (row_obj, row_data) in enumerate(zip(data_rows, full_data), start=1):
        set_cell_text(row_obj.cells[0], f"图{i}")
        # 清空第二列已有内容
        cell = row_obj.cells[1]
        for p in cell.paragraphs:
            for run in list(p.runs):
                run.text = ""

        img_path = row_to_image.get(i - 1)
        if img_path and os.path.exists(img_path):
            p = cell.paragraphs[0]
            run = p.add_run()
            run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM), height=Cm(IMAGE_HEIGHT_CM))
        else:
            set_cell_text(cell, "（未找到对应图片）")


def step2_generate_docx(full_data, docx_output_path):
    print(f"\n{'=' * 50}")
    print("  Step 2: 生成Word报告")
    print(f"{'=' * 50}\n")

    print(f"正在读取模板: {DOCX_TEMPLATE_PATH}")
    doc = Document(DOCX_TEMPLATE_PATH)

    step3_fill_word_data_table(doc, full_data)

    image_files = collect_image_files(IMAGE_DIR)
    atlas_codes = [row['atlas_code'] for row in full_data]
    row_to_image = match_images_to_rows(image_files, atlas_codes)
    step4_insert_images(doc, full_data, row_to_image)

    apply_font_to_whole_document(doc)

    doc.save(docx_output_path)
    print(f"\n测温报告已生成: {docx_output_path}")


# ===================== 主程序 =====================

def interactive_menu():
    """交互式菜单：PyCharm 直接运行时弹出选项。"""
    print()
    print("=" * 55)
    print("  红外测温报告自动生成程序 V5")
    print("=" * 55)
    print(f"  工作目录: {BASE_DIR}")
    print()

    excel_exists = os.path.isfile(EXCEL_PATH)
    latest_excel = find_latest_output(EXCEL_OUTPUT_BASE, "xlsx")
    excel_out_exists = latest_excel is not None
    image_dir_exists = os.path.isdir(IMAGE_DIR)
    template_exists = os.path.isfile(DOCX_TEMPLATE_PATH)
    existing_reports = sorted(glob.glob(f"{DOCX_OUTPUT_BASE}_*.docx"))

    print("  文件状态检查：")
    print(f"    模板 Excel:       {'✓ 存在' if excel_exists else '✗ 缺失！'}")
    print(f"    Step1 输出 Excel: {'✓ 已生成' if excel_out_exists else '○ 尚未生成'}")
    print(f"    图谱文件夹:       {'✓ 存在' if image_dir_exists else '✗ 缺失！'}")
    print(f"    Word 模板:        {'✓ 存在' if template_exists else '✗ 缺失！'}")
    if existing_reports:
        print(f"    已有报告文件:     {len(existing_reports)} 份")
    print()

    print("  请选择运行模式：")
    print("    [1] 仅 Step 1 — OCR识别 + 人工核对 + 写入Excel")
    print("    [2] 仅 Step 2 — 读取已有Excel，生成Word报告")
    print("    [3] 完整流程 — Step 1 + Step 2 串联执行")
    print("    [0] 退出")
    print()

    while True:
        choice = input("  请输入数字 (0/1/2/3): ").strip()
        if choice in ("0", "1", "2", "3"):
            break
        print("  ⚠️ 无效输入，请输入 0、1、2 或 3。")

    if choice == "0":
        print("  已退出。")
        sys.exit(0)

    print()
    return choice in ("1", "3"), choice in ("2", "3")


def main():
    # ---------- 命令行参数解析 ----------
    parser = argparse.ArgumentParser(
        description="红外测温报告自动生成程序 V5",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        add_help=False,
    )
    parser.add_argument("--step1", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--step2", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--cli", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--suffix", type=str, default=None, help=argparse.SUPPRESS)
    parser.add_argument("--help", action="store_true", help=argparse.SUPPRESS)

    has_args = len(sys.argv) > 1

    if has_args:
        args, _ = parser.parse_known_args()
        if args.help:
            print("""
使用方式：
  python3 脚本名.py                   PyCharm直接运行 → 交互式菜单
  python3 脚本名.py --step1            终端：仅 step1
  python3 脚本名.py --step2            终端：仅 step2
  python3 脚本名.py --step2 --suffix v2 终端：step2 + 版本后缀
  python3 脚本名.py --cli              终端：命令行核对模式
""")
            sys.exit(0)

        if args.cli:
            global USE_GUI
            USE_GUI = False
            print("已启用 --cli 模式：将使用命令行进行人工核对。\n")

        run_step1 = args.step1 or not (args.step1 or args.step2)
        run_step2 = args.step2 or not (args.step1 or args.step2)
        suffix = args.suffix
    else:
        run_step1, run_step2 = interactive_menu()
        suffix = None

    # ---------- 确定输出路径 ----------
    excel_output_path = resolve_output_path(EXCEL_OUTPUT_BASE, "xlsx", suffix)
    docx_output_path = resolve_output_path(DOCX_OUTPUT_BASE, "docx", suffix)

    # ---------- tkinter 预初始化（仅 macOS 需要，Windows 上反而会破坏 Tcl 解释器）----------
    if USE_GUI and run_step1 and platform.system() == "Darwin":
        try:
            import tkinter as tk
            _pre = tk.Tk()
            _pre.withdraw()
            _pre.destroy()
        except Exception:
            pass

    # ---------- 执行 Step 1 ----------
    if run_step1:
        if not os.path.isdir(IMAGE_DIR):
            print(f"错误：图谱文件夹不存在: {IMAGE_DIR}")
            sys.exit(1)
        if not os.path.isfile(EXCEL_PATH):
            print(f"错误：Excel文件不存在: {EXCEL_PATH}")
            sys.exit(1)
        if not os.path.isfile(DOCX_TEMPLATE_PATH):
            print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
            sys.exit(1)

        full_data = step1_ocr_and_review(excel_output_path)
        print(f"\n[Step 1 完成] 识别结果已写入: {excel_output_path}")

    # ---------- 执行 Step 2 ----------
    if run_step2:
        if not os.path.isfile(DOCX_TEMPLATE_PATH):
            print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
            sys.exit(1)

        if run_step1:
            print(f"\n{'=' * 50}")
            print("  进入 Step 2（使用 Step 1 的内存数据）")
            print(f"{'=' * 50}")
        else:
            full_data = read_excel_output()

        step2_generate_docx(full_data, docx_output_path)

        existing = sorted(glob.glob(f"{DOCX_OUTPUT_BASE}_*.docx"))
        if len(existing) > 1:
            print(f"\n当前目录下共有 {len(existing)} 份报告（均未被覆盖）：")
            for f in existing:
                print(f"  - {os.path.basename(f)}")

    # ---------- 汇总 ----------
    print(f"\n{'=' * 50}")
    print("  全部完成！")
    if run_step1:
        print(f"  [Step 1] Excel结果: {excel_output_path}")
    if run_step2:
        print(f"  [Step 2] Word报告:  {docx_output_path}")
    else:
        print(f"  （Step 2 未执行，再次运行本脚本选择 [2] 即可生成报告）")
    print(f"{'=' * 50}")


if __name__ == "__main__":
    # GUI 开关（43号文也保留 GUI 核对能力）
    USE_GUI = True
    main()
