import os
import re
import tkinter as tk
from tkinter import filedialog, messagebox
from PIL import Image, ImageTk
import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook
from docx import Document
from docx.shared import Cm
import sys

# ---------- 全局变量 ----------
ROI = None  # 存储框选区域 (x, y, w, h)
IMAGE_FOLDER = "图谱"
EXCEL_PATH = "必填信息.xlsx"
WORD_TEMPLATE = "测温报告.docx"
WORD_OUTPUT = "测温报告_生成.docx"


# ---------- GUI 区域选择 ----------
class ROISelector:
    def __init__(self, master, image_path):
        self.master = master
        self.image = Image.open(image_path)
        self.tk_image = ImageTk.PhotoImage(self.image)
        self.canvas = tk.Canvas(master, width=self.image.width, height=self.image.height)
        self.canvas.pack()
        self.canvas.create_image(0, 0, anchor=tk.NW, image=self.tk_image)
        self.rect = None
        self.start_x = None
        self.start_y = None
        self.roi = None
        self.canvas.bind("<ButtonPress-1>", self.on_press)
        self.canvas.bind("<B1-Motion>", self.on_drag)
        self.canvas.bind("<ButtonRelease-1>", self.on_release)
        tk.Button(master, text="确认区域", command=self.confirm).pack()
        master.title("框选温度数值区域")

    def on_press(self, event):
        self.start_x = event.x
        self.start_y = event.y
        if self.rect:
            self.canvas.delete(self.rect)
            self.rect = None

    def on_drag(self, event):
        if self.rect:
            self.canvas.delete(self.rect)
        self.rect = self.canvas.create_rectangle(self.start_x, self.start_y, event.x, event.y, outline='red')

    def on_release(self, event):
        pass

    def confirm(self):
        if self.rect:
            coords = self.canvas.coords(self.rect)
            x1, y1, x2, y2 = map(int, coords)
            self.roi = (min(x1, x2), min(y1, y2), abs(x2-x1), abs(y2-y1))
            self.master.quit()
        else:
            messagebox.showerror("错误", "请先框选一个区域！")


def select_roi(image_path):
    """弹出窗口让用户框选区域，返回 (x, y, w, h)"""
    root = tk.Tk()
    app = ROISelector(root, image_path)
    root.mainloop()
    root.destroy()
    if app.roi is None:
        sys.exit("未选择区域，程序退出")
    return app.roi


# ---------- OCR 识别温度 ----------
def recognize_temperature(image_crop):
    """对裁剪后的图像进行OCR，返回浮点数温度，若失败返回None"""
    # 转换为灰度，二值化增强
    gray = cv2.cvtColor(np.array(image_crop), cv2.COLOR_RGB2GRAY)
    _, thresh = cv2.threshold(gray, 150, 255, cv2.THRESH_BINARY)
    # 使用Tesseract识别数字
    custom_config = r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.'
    text = pytesseract.image_to_string(thresh, config=custom_config)
    # 提取第一个浮点数
    match = re.search(r'(\d+\.?\d*)', text)
    if match:
        return float(match.group(1))
    return None


# ---------- 处理图片提取温度 ----------
def extract_temperatures(image_folder, roi, excel_path):
    """遍历图片文件夹，识别温度并填入Excel的E列"""
    # 获取所有图片文件（按名称排序）
    ext = ('.jpg', '.jpeg', '.png', '.bmp')
    images = [f for f in os.listdir(image_folder) if f.lower().endswith(ext)]
    images.sort()  # 按名称排序

    # 加载Excel
    wb = load_workbook(excel_path)
    ws = wb['普测']  # 假设工作表名为'普测'

    # 从第2行开始（第1行是表头）
    row = 2
    for img_name in images:
        img_path = os.path.join(image_folder, img_name)
        try:
            img = Image.open(img_path)
            # 裁剪ROI
            x, y, w, h = roi
            crop = img.crop((x, y, x+w, y+h))
            temp = recognize_temperature(crop)
            if temp is not None:
                ws.cell(row=row, column=5, value=temp)  # E列
                print(f"{img_name} -> {temp}℃")
            else:
                print(f"警告：{img_name} 温度识别失败，留空")
        except Exception as e:
            print(f"处理 {img_name} 出错: {e}")
        row += 1

    wb.save(excel_path)
    print("温度提取完成，已保存Excel。")


# ---------- 计算正常温度 ----------
def calculate_normal_temp(excel_path):
    """根据规则计算正常温度并填入F列"""
    wb = load_workbook(excel_path)
    ws = wb['普测']

    # 先读取所有数据
    data = []
    for row in range(2, ws.max_row + 1):
        device = ws.cell(row=row, column=3).value  # C列 设备名称
        surface = ws.cell(row=row, column=5).value  # E列 表面温度
        data.append((row, device, surface))

    # 分组：提取前缀（去掉A/B/C相）
    groups = {}
    for row, device, surface in data:
        if not device or surface is None:
            continue
        # 判断是否包含A相/B相/C相
        match = re.search(r'(.*?)(?:A相|B相|C相)$', device)
        if match:
            prefix = match.group(1).strip()
            groups.setdefault(prefix, []).append((row, surface))
        else:
            # 单独一组，使用自身surface
            groups.setdefault(device, []).append((row, surface))

    # 计算每组正常温度
    for group, items in groups.items():
        if len(items) == 1:
            # 单独设备
            row, surface = items[0]
            normal = surface
        else:
            # 多个（A/B/C），取最小表面温度
            surfaces = [s for _, s in items]
            normal = min(surfaces)
        # 将正常温度填入该组所有行的F列
        for row, _ in items:
            ws.cell(row=row, column=6, value=normal)  # F列

    wb.save(excel_path)
    print("正常温度计算完成。")


# ---------- 生成Word报告 ----------
def generate_word_report(excel_path, word_template, output_path, image_folder):
    """填充Word报告"""
    # 读取Excel数据
    wb = load_workbook(excel_path)
    ws = wb['普测']
    rows_data = []
    for row in range(2, ws.max_row + 1):
        if ws.cell(row=row, column=1).value is None:  # 序号为空则停止
            break
        row_vals = []
        for col in range(1, 11):  # A~J
            row_vals.append(ws.cell(row=row, column=col).value)
        rows_data.append(row_vals)

    # 打开Word模板
    doc = Document(word_template)

    # ---------- 填充第一个表格（检测数据） ----------
    # 寻找包含'序号'的表格和行
    target_table = None
    header_row_index = None
    for table in doc.tables:
        for i, row in enumerate(table.rows):
            cells = row.cells
            if len(cells) >= 1 and '序号' in cells[0].text:
                target_table = table
                header_row_index = i
                break
        if target_table:
            break

    if not target_table:
        raise ValueError("未找到包含'序号'的表格！")

    # 删除模板中可能存在的示例数据行（从header_row_index+1开始到表格末尾）
    # 由于模板中已有空行，我们保留，但为了安全，我们删除所有非表头行
    while len(target_table.rows) > header_row_index + 1:
        target_table._element.remove(target_table.rows[header_row_index+1]._element)

    # 插入数据行
    for row_vals in rows_data:
        new_row = target_table.add_row()
        cells = new_row.cells
        for col_idx, val in enumerate(row_vals):
            if col_idx < len(cells):
                cells[col_idx].text = str(val) if val is not None else ''

    # ---------- 计算温差并填充结论 ----------
    # 计算温差（K列）并记录温差>15的设备
    wb = load_workbook(excel_path)
    ws = wb['普测']
    # 找到所有数据行
    diff_devices = []
    for row in range(2, ws.max_row + 1):
        surface = ws.cell(row=row, column=5).value
        normal = ws.cell(row=row, column=6).value
        if surface is not None and normal is not None:
            diff = surface - normal
            if diff > 15:
                device = ws.cell(row=row, column=3).value
                if device:
                    diff_devices.append(device)

    if diff_devices:
        # 去重并拼接
        unique = list(dict.fromkeys(diff_devices))
        conclusion = "、".join(unique) + "发热,达到一般缺陷标准,未达到严重缺陷标准（发热缺陷较前期无明显发展趋势），无其他异常"
    else:
        conclusion = "无异常"

    # 定位结论单元格（表格倒数第二行，4~10列合并）
    # 在第一个表格中查找包含'结论'的单元格
    conclusion_cell = None
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if '结论' in cell.text:
                    conclusion_cell = cell
                    break
            if conclusion_cell:
                break
        if conclusion_cell:
            break

    if conclusion_cell:
        # 清除原有内容，填入新结论
        conclusion_cell.paragraphs[0].clear()
        conclusion_cell.paragraphs[0].add_run(conclusion)
    else:
        print("警告：未找到'结论'单元格，请手动添加。")

    # ---------- 插入图片（第二个表格） ----------
    # 查找附件表格（通常包含'图谱编号'和'图谱'）
    attachment_table = None
    for table in doc.tables:
        for row in table.rows:
            cells = row.cells
            if len(cells) >= 2 and '图谱编号' in cells[0].text and '图谱' in cells[1].text:
                attachment_table = table
                break
        if attachment_table:
            break

    if not attachment_table:
        raise ValueError("未找到附件表格！")

    # 删除已有数据行（从第二行开始）
    while len(attachment_table.rows) > 1:
        attachment_table._element.remove(attachment_table.rows[1]._element)

    # 获取图片列表
    ext = ('.jpg', '.jpeg', '.png', '.bmp')
    images = [f for f in os.listdir(image_folder) if f.lower().endswith(ext)]
    images.sort()

    for idx, img_name in enumerate(images, start=1):
        img_path = os.path.join(image_folder, img_name)
        # 添加行
        new_row = attachment_table.add_row()
        cells = new_row.cells
        # 第一列：图编号
        cells[0].text = f"图{idx}"
        # 第二列：插入图片
        try:
            # 清除原有段落
            para = cells[1].paragraphs[0]
            para.clear()
            run = para.add_run()
            run.add_picture(img_path, width=Cm(12), height=Cm(9))
        except Exception as e:
            print(f"插入图片 {img_name} 失败: {e}")

    # 保存Word
    doc.save(output_path)
    print(f"Word报告已生成：{output_path}")


# ---------- 主流程 ----------
def main():
    # 检查文件夹是否存在
    if not os.path.exists(IMAGE_FOLDER):
        print(f"错误：图片文件夹 '{IMAGE_FOLDER}' 不存在！")
        return

    # 1. 选择示例图片
    img_files = [f for f in os.listdir(IMAGE_FOLDER) if f.lower().endswith(('.jpg','.jpeg','.png','.bmp'))]
    if not img_files:
        print("图谱文件夹中没有图片！")
        return
    # 默认选择第一个
    example_path = os.path.join(IMAGE_FOLDER, img_files[0])
    print(f"使用示例图片：{example_path}")

    # 2. 框选ROI
    roi = select_roi(example_path)
    print(f"框选区域：{roi}")

    # 3. 提取温度
    extract_temperatures(IMAGE_FOLDER, roi, EXCEL_PATH)

    # 4. 计算正常温度
    calculate_normal_temp(EXCEL_PATH)

    # 5. 生成Word报告
    generate_word_report(EXCEL_PATH, WORD_TEMPLATE, WORD_OUTPUT, IMAGE_FOLDER)

    print("全部流程完成！")


if __name__ == "__main__":
    main()