# -*- coding: utf-8 -*-
"""
91号文《红外精确测温报告》自动生成程序
======================================================

本脚本假设以下三样东西和本脚本放在【同一个文件夹】里（不需要修改任何路径配置）：
    ./图谱/                图片文件夹（红外精确测温图片，按文件名顺序对应Excel每一行）
    ./必填信息.xlsx         数据表（工作表名："91号文模板"，已提前填好B~I列，J/K/L列由本脚本识别）
    ./91号文报告.docx        Word报告模板（每份模板只装一个设备的信息，本脚本会自动复制成多份、
                             每个设备占一页，拼接成一份完整报告）

【整体流程】
    1. 识别：分别框选"温度(Ar1)"和"拍摄日期+时间"两个区域（各框一次，所有图片复用），
       对图谱文件夹里的每张图片（按文件名顺序，依次对应Excel第2、3、4...行）自动识别。
    2. 人工核对（两个独立窗口）：
       - 窗口1：Ar1温度 → 裁剪图+输入框，逐条核对，最后一条按"确认提交"
       - 窗口2：拍摄日期+时间 → 同一个裁剪图，日期和时间两个输入框，逐条核对，最后一条按"确认提交"
    3. 核对完成后，把J（Ar1温度）K（拍摄日期）L（拍摄时间）写入必填信息_已生成.xlsx。
    4. 根据必填信息_已生成.xlsx和图谱文件夹里的图片，生成91号文报告_已生成.docx：
       每个设备复制一份模板、填入对应信息和图片，每个设备占一页，全部设备依次拼接成一份文档。

依赖安装（与43号文脚本共用同一个虚拟环境即可，无需重新安装）：
    pip install opencv-python numpy pytesseract openpyxl pillow python-docx

还需要系统安装 Tesseract OCR 引擎本体（如果43号文脚本能跑，这里不需要重装）：
    brew install tesseract   或   conda install -c conda-forge tesseract

使用方法：
    python3 91号文报告自动生成.py
"""

import os
import re
import sys
import json
import glob
import copy
import time
import platform
import subprocess
import datetime

import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook

from docx import Document
from docx.shared import Cm, Pt
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ===================== 配置区域 =====================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

IMAGE_DIR = os.path.join(BASE_DIR, "图谱")
EXCEL_PATH = os.path.join(BASE_DIR, "必填信息.xlsx")
EXCEL_OUTPUT_PATH = os.path.join(BASE_DIR, "必填信息_已生成.xlsx")
EXCEL_SHEET_NAME = "91号文模板"

DOCX_TEMPLATE_PATH = os.path.join(BASE_DIR, "91号文报告.docx")
DOCX_OUTPUT_PATH = os.path.join(BASE_DIR, "91号文报告_已生成.docx")

AR1_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_ar1.json")
DATETIME_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_datetime.json")

DEBUG_DIR = os.path.join(BASE_DIR, "debug_crops_91")
SAVE_DEBUG_CROPS = True

NOT_RECOGNIZED_TEXT = "无法识别"

# GUI核对开关（--cli 参数可强制使用命令行模式，避免macOS上tkinter间歇性崩溃）
USE_GUI = True

# 插入图片尺寸
IMAGE_WIDTH_CM = 12
IMAGE_HEIGHT_CM = 9

# 新填入文字的字体（只针对脚本写入的值，模板原有文字不动）
CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
FONT_SIZE = Pt(8.5)

# =================================================================


# ===================== 工具函数 =====================

def natural_sort_key(s):
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def collect_image_files(image_dir):
    patterns = ["*.jpg", "*.JPG", "*.jpeg", "*.JPEG", "*.png", "*.PNG"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(image_dir, p)))
    files = sorted(set(files), key=lambda p: natural_sort_key(os.path.basename(p)))
    return files


def open_image_for_preview(image_path):
    try:
        system = platform.system()
        if system == "Darwin":
            subprocess.run(["open", image_path], check=False)
        elif system == "Windows":
            os.startfile(image_path)  # noqa
        else:
            subprocess.run(["xdg-open", image_path], check=False)
    except Exception as e:
        print(f"  （提示：自动打开图片失败：{e}，可手动查看：{image_path}）")


# ===================== ROI框选（相对比例，两个独立区域：Ar1温度 / 日期时间合并） =====================

def select_roi_interactively(sample_image_path, window_title):
    """自定义鼠标框选ROI区域（细线矩形），空格/回车确认，c键取消重选。"""
    img = cv2.imread(sample_image_path)
    if img is None:
        raise FileNotFoundError(f"无法读取示例图片: {sample_image_path}")
    img_h, img_w = img.shape[:2]

    state = {"drawing": False, "x0": -1, "y0": -1, "x1": -1, "y1": -1, "done": False}

    def on_mouse(event, x, y, flags, param):
        if state["done"]:
            return
        if event == cv2.EVENT_LBUTTONDOWN:
            state["drawing"] = True
            state["x0"], state["y0"] = x, y
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_MOUSEMOVE and state["drawing"]:
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_LBUTTONUP:
            state["drawing"] = False
            state["x1"], state["y1"] = x, y

    cv2.namedWindow(window_title)
    cv2.setMouseCallback(window_title, on_mouse)
    display = img.copy()

    print(f"\n请在弹出的窗口中用鼠标左键拖拽框选【{window_title}】区域。")
    print("  空格/回车=确认   c=取消重选   Esc=退出\n")

    while True:
        display = img.copy()
        x0, y0 = state["x0"], state["y0"]
        x1, y1 = state["x1"], state["y1"]
        if x0 >= 0 and y0 >= 0:
            cv2.rectangle(display, (x0, y0), (x1, y1), (0, 255, 0), 1)
        cv2.imshow(window_title, display)
        key = cv2.waitKey(20) & 0xFF
        if key == 13 or key == 32:  # Enter or Space
            break
        elif key == ord('c') or key == ord('C') or key == 27:  # c or Esc
            state["x0"] = state["y0"] = state["x1"] = state["y1"] = -1
            state["drawing"] = False
            if key == 27:
                break

    cv2.destroyAllWindows()
    cv2.waitKey(1)
    time.sleep(0.3)

    x0, y0 = state["x0"], state["y0"]
    x1, y1 = state["x1"], state["y1"]
    if x0 < 0 or y0 < 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    x = min(x0, x1)
    y = min(y0, y1)
    w = abs(x1 - x0)
    h = abs(y1 - y0)
    if w == 0 or h == 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    rx, ry, rw, rh = x / img_w, y / img_h, w / img_w, h / img_h
    print(f"框选区域（像素）: x={x}, y={y}, w={w}, h={h}  (示例图片尺寸: {img_w}x{img_h})")
    return rx, ry, rw, rh


def match_images_by_number(image_files, data_rows_info):
    """
    把图谱文件夹里的图片对应到Excel每一行。
    优先方式：从文件名里提取数字（比如"图5.jpg"->5），第N号图片对应第N行数据。
    这样即使中间缺了一张图（比如"图5.jpg"不存在），也只会让第5行没有图，
    不会导致第6行及以后的图全部错位往前顶。
    如果文件名根本不含数字规律，才退回成"按文件夹内文件名顺序直接依次对应"。
    返回：{Excel行号: 图片路径 或 None}
    """
    numbered = {}
    all_numbered = True
    for p in image_files:
        base = os.path.splitext(os.path.basename(p))[0]
        m = re.search(r'(\d+)', base)
        if m:
            numbered[int(m.group(1))] = p
        else:
            all_numbered = False

    row_to_image = {}
    if all_numbered and numbered:
        print("图片匹配方式：按文件名中的数字精确对应第几行（例如'图5.jpg'对应第5行数据）。")
        for i, info in enumerate(data_rows_info, start=1):
            row_to_image[info['row']] = numbered.get(i)
    else:
        print("图片匹配方式：文件名不含数字规律，改为按文件夹内文件名顺序直接依次对应。")
        for i, info in enumerate(data_rows_info):
            row_to_image[info['row']] = image_files[i] if i < len(image_files) else None

    missing = [info['row'] for info in data_rows_info if row_to_image.get(info['row']) is None]
    if missing:
        print(f"提示：以下Excel行号没有匹配到图片，需要人工确认：{missing}")

    return row_to_image


def get_roi(sample_image_path, config_file, prompt_text, window_title):
    if os.path.exists(config_file):
        with open(config_file, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        rx, ry, rw, rh = cfg["rx"], cfg["ry"], cfg["rw"], cfg["rh"]
        print(f"检测到已保存的【{prompt_text}】区域配置: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
        choice = input(f"是否复用该区域？(y=复用 / n=重新框选) [y]: ").strip().lower()
        if choice in ("", "y", "yes"):
            return rx, ry, rw, rh

    print(f"\n请在弹出的窗口中框选【{prompt_text}】区域，框选完成后按【空格/回车】确认，按 c 取消重选。")
    rx, ry, rw, rh = select_roi_interactively(sample_image_path, window_title)
    with open(config_file, "w", encoding="utf-8") as f:
        json.dump({"rx": rx, "ry": ry, "rw": rw, "rh": rh}, f, ensure_ascii=False, indent=2)
    print(f"识别区域已保存到: {config_file}")
    return rx, ry, rw, rh


def crop_by_relative_roi(img, roi_ratio):
    rx, ry, rw, rh = roi_ratio
    img_h, img_w = img.shape[:2]
    x = int(round(rx * img_w))
    y = int(round(ry * img_h))
    w = int(round(rw * img_w))
    h = int(round(rh * img_h))
    x2, y2 = min(x + w, img_w), min(y + h, img_h)
    x, y = max(x, 0), max(y, 0)
    return img[y:y2, x:x2]


def preprocess_variants(crop, scale=4):
    crop_big = cv2.resize(crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(crop_big, cv2.COLOR_BGR2GRAY)

    variants = []
    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh1 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh1)

    inv = cv2.bitwise_not(blur)
    _, thresh2 = cv2.threshold(inv, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh2)

    thresh3 = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
    )
    variants.append(thresh3)

    return variants


# ===================== OCR：Ar1温度（纯数字） =====================

def extract_number(text):
    if not text:
        return None
    match = re.search(r'-?\d+\.?\d*', text.strip())
    if not match:
        return None
    try:
        num = float(match.group())
        return str(int(num)) if num == int(num) else str(num)
    except ValueError:
        return None


def ocr_ar1_temperature(image_path, roi_ratio):
    img = cv2.imread(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, ["[裁剪区域为空]"]

    variants = preprocess_variants(crop)
    configs = [
        r'--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 11 -c tessedit_char_whitelist=0123456789.-',
    ]
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            value = extract_number(text)
            if value is not None:
                return value, crop, raw_texts
    return None, crop, raw_texts


# ===================== OCR：拍摄日期 / 拍摄时间（各自独立识别） =====================

DATE_PATTERN = re.compile(r'(\d{4})\D(\d{1,2})\D(\d{1,2})')
TIME_PATTERN = re.compile(r'(\d{1,2}):(\d{2}):(\d{2})')


def parse_date(text):
    if not text:
        return None
    dm = DATE_PATTERN.search(text)
    if not dm:
        return None
    try:
        y, mo, d = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
        return datetime.date(y, mo, d)
    except ValueError:
        return None


def parse_time(text):
    if not text:
        return None
    tm = TIME_PATTERN.search(text)
    if not tm:
        return None
    h, mi, s = tm.group(1), tm.group(2), tm.group(3)
    try:
        if 0 <= int(h) <= 23 and 0 <= int(mi) <= 59 and 0 <= int(s) <= 59:
            return f"{int(h):02d}:{mi}:{s}"
    except ValueError:
        pass
    return None


def _ocr_with_configs(crop, configs):
    variants = preprocess_variants(crop)
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            yield text, raw_texts


DATE_TIME_CHAR_WHITELIST = r'-c tessedit_char_whitelist=0123456789.:-年月日 '
_OCR_CONFIGS_DATE_TIME = [
    rf'--oem 3 --psm 7 {DATE_TIME_CHAR_WHITELIST}',
    rf'--oem 3 --psm 6 {DATE_TIME_CHAR_WHITELIST}',
]


def ocr_datetime(image_path, roi_ratio):
    """从图片同一个ROI区域同时提取日期和时间。"""
    img = cv2.imread(image_path)
    if img is None:
        return None, None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, crop, ["[裁剪区域为空]"]

    raw_texts = []
    for text, raw_texts in _ocr_with_configs(crop, _OCR_CONFIGS_DATE_TIME):
        date_val = parse_date(text)
        time_val = parse_time(text)
        if date_val is not None or time_val is not None:
            return date_val, time_val, crop, raw_texts
    return None, None, crop, raw_texts


# ===================== Excel读取 =====================

def find_header_row(ws, keyword="设备类型", max_scan_rows=5):
    for row in range(1, max_scan_rows + 1):
        for col in range(1, 5):
            if str(ws.cell(row=row, column=col).value).strip() == keyword:
                return row
    return 1


def read_excel_data_rows(ws, header_row_idx):
    """从表头下一行开始读取，直到设备名称(D列)为空为止。"""
    data_rows_info = []
    row = header_row_idx + 1
    while True:
        d_val = ws.cell(row=row, column=4).value  # D列 设备名称
        if d_val is None or str(d_val).strip() == "":
            break
        info = {
            'row': row,
            'device_type': ws.cell(row=row, column=2).value,     # B 设备类型
            'interval': ws.cell(row=row, column=3).value,         # C 间隔单元
            'device_name': d_val,                                    # D 设备名称
            'phase': ws.cell(row=row, column=5).value,             # E 相别
            'distance': ws.cell(row=row, column=6).value,          # F 测试距离
            'voltage': ws.cell(row=row, column=7).value,           # G 运行电压
            'current': ws.cell(row=row, column=8).value,           # H 负荷电流
            'power': ws.cell(row=row, column=9).value,             # I 有功功率
        }
        data_rows_info.append(info)
        row += 1
    return data_rows_info


# ===================== 人工核对（两个独立窗口：Ar1温度 / 日期时间） =====================

# --- 公共 GUI 工具 ---

def _load_image_into(label, path, ref_dict, ref_key, max_w=580, max_h=320):
    try:
        from PIL import Image, ImageTk
    except ImportError:
        return
    if path and os.path.exists(path):
        pil_img = Image.open(path)
        ratio = min(max_w / pil_img.width, max_h / pil_img.height, 1.0)
        new_size = (max(1, int(pil_img.width * ratio)), max(1, int(pil_img.height * ratio)))
        pil_img = pil_img.resize(new_size)
        photo = ImageTk.PhotoImage(pil_img)
        ref_dict[ref_key] = photo
        label.config(image=photo, text="")
    else:
        ref_dict[ref_key] = None
        label.config(image="", text="（未找到对应裁剪图片）")


# --- 核对窗口 1：Ar1温度 ---

def interactive_review_ar1_gui(data_rows_info, ar1_results, ar1_debug_map):
    """核对窗口1：只核对Ar1温度，裁剪图 + 输入框。"""
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（GUI不可用({e})，回退命令行。）")
        return interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map)

    reviewed = dict(ar1_results)
    total = len(data_rows_info)
    state = {'index': 0, 'confirmed': False}
    photo_refs = {}

    root = tk.Tk()
    root.title("91号文测温 - 核对 Ar1温度")
    root.geometry("620x520")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 14), justify="left", wraplength=580)
    info_label.pack(pady=(15, 5))

    img_label = tk.Label(root)
    img_label.pack(pady=8)

    entry_var = tk.StringVar()
    entry = tk.Entry(root, textvariable=entry_var, font=("PingFang SC", 22), justify="center", width=14)
    entry.pack(pady=8)

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(0, 5))

    btn_ref = {}

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']
        current = reviewed.get(r)
        status = current if current is not None else ""
        mark = "   ⚠️ 识别失败，需手动输入" if current is None else ""
        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}{mark}")
        _load_image_into(img_label, ar1_debug_map.get(r), photo_refs, 'ar1')
        entry_var.set(status)
        hint_label.config(text="↑ = 上一条    |    ↓ = 下一条    |    回车 = 确认提交", fg="gray")
        entry.focus_set()
        entry.select_range(0, tk.END)
        update_buttons()

    def update_buttons():
        idx = state['index']
        is_first = (idx == 0)
        is_last = (idx + 1 >= total)
        btn_ref['prev'].config(state=(tk.DISABLED if is_first else tk.NORMAL))
        btn_ref['next'].config(state=(tk.DISABLED if is_last else tk.NORMAL))

    def try_save():
        idx = state['index']
        r = data_rows_info[idx]['row']
        val = entry_var.get().strip()
        if val != "":
            try:
                float(val)
                reviewed[r] = val
            except ValueError:
                hint_label.config(text="⚠️ 不是有效数字，请重新输入！", fg="red")
                return False
        return True

    def go_previous(event=None):
        if not try_save():
            return
        idx = state['index']
        if idx == 0:
            return
        state['index'] -= 1
        load_row(state['index'])

    def go_next(event=None):
        if not try_save():
            return
        idx = state['index']
        if idx + 1 >= total:
            hint_label.config(text="已是最后一条，请按回车提交", fg="orange")
            return
        state['index'] += 1
        load_row(state['index'])

    def confirm_submit(event=None):
        if not try_save():
            return
        state['confirmed'] = True
        root.destroy()

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=15)
    btn_ref['prev'] = tk.Button(button_frame, text="↑ 上一条", command=go_previous, width=10)
    btn_ref['prev'].grid(row=0, column=0, padx=4)
    btn_ref['next'] = tk.Button(button_frame, text="↓ 下一条", command=go_next, width=10)
    btn_ref['next'].grid(row=0, column=1, padx=4)
    btn_ref['submit'] = tk.Button(button_frame, text="确认提交（回车）", command=confirm_submit, width=16, fg="blue")
    btn_ref['submit'].grid(row=0, column=2, padx=4)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=3, padx=4)

    root.bind('<Up>', go_previous)
    root.bind('<Down>', go_next)
    root.bind('<Return>', confirm_submit)

    print("\n已打开【Ar1温度】核对窗口...")
    load_row(0)
    root.lift()
    root.focus_force()
    entry.focus_set()
    root.mainloop()

    if state['confirmed']:
        print("【Ar1温度】核对完成，用户已确认提交。")
    else:
        print("【Ar1温度】核对已跳过。")
    return reviewed


def interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map):
    """命令行兜底：Ar1温度核对。"""
    reviewed = dict(ar1_results)
    total = len(data_rows_info)
    print(f"\n{'=' * 50}\n进入【Ar1温度】核对（命令行）\n{'=' * 50}")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        current = reviewed.get(r)
        print(f"\n[{i}/{total}] {info['device_name']}  当前值={current}")
        if ar1_debug_map.get(r):
            open_image_for_preview(ar1_debug_map[r])
        while True:
            val = input("  温度 (回车=保留): ").strip()
            if val == "" or _is_float(val):
                if val != "":
                    reviewed[r] = val
                break
            print("  不是有效数字。")
        if input("  s=跳过剩余: ").strip().lower() in ("s", "skip"):
            break
    return reviewed


# --- 核对窗口 2：拍摄日期+时间（同一个裁剪图，两个输入框） ---

def interactive_review_datetime_gui(data_rows_info, date_results, time_results, dt_debug_map):
    """核对窗口2：日期+时间共用一个裁剪图，两个输入框分别显示日期和时间。"""
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（GUI不可用({e})，回退命令行。）")
        return interactive_review_datetime_cli(data_rows_info, date_results, time_results, dt_debug_map)

    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)
    state = {'index': 0, 'confirmed': False}
    photo_refs = {}

    root = tk.Tk()
    root.title("91号文测温 - 核对 拍摄日期+时间")
    root.geometry("620x580")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 14), justify="left", wraplength=580)
    info_label.pack(pady=(15, 5))

    img_label = tk.Label(root)
    img_label.pack(pady=6)

    # 日期输入
    tk.Label(root, text="拍摄日期 (YYYY-MM-DD)", font=("PingFang SC", 11)).pack(pady=(8, 0))
    date_var = tk.StringVar()
    date_entry = tk.Entry(root, textvariable=date_var, font=("PingFang SC", 18), justify="center", width=18)
    date_entry.pack(pady=3)

    # 时间输入
    tk.Label(root, text="拍摄时间 (HH:MM:SS)", font=("PingFang SC", 11)).pack(pady=(6, 0))
    time_var = tk.StringVar()
    time_entry = tk.Entry(root, textvariable=time_var, font=("PingFang SC", 18), justify="center", width=18)
    time_entry.pack(pady=3)

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(6, 4))

    btn_ref = {}

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']
        dc = date_reviewed.get(r)
        tc = time_reviewed.get(r)
        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}")
        _load_image_into(img_label, dt_debug_map.get(r), photo_refs, 'dt')
        date_var.set(dc.isoformat() if isinstance(dc, datetime.date) else (dc or ""))
        time_var.set(tc if tc else "")
        hint_label.config(text="↑ = 上一条    |    ↓ = 下一条    |    回车 = 确认提交", fg="gray")
        date_entry.focus_set()
        date_entry.select_range(0, tk.END)
        update_buttons()

    def update_buttons():
        idx = state['index']
        is_first = (idx == 0)
        is_last = (idx + 1 >= total)
        btn_ref['prev'].config(state=(tk.DISABLED if is_first else tk.NORMAL))
        btn_ref['next'].config(state=(tk.DISABLED if is_last else tk.NORMAL))

    def try_save():
        idx = state['index']
        r = data_rows_info[idx]['row']
        ds = date_var.get().strip()
        ts = time_var.get().strip()
        if ds:
            try:
                date_reviewed[r] = datetime.date.fromisoformat(ds)
            except ValueError:
                hint_label.config(text="⚠️ 日期格式不对，请用 YYYY-MM-DD！", fg="red")
                return False
        if ts:
            if not re.match(r'^\d{1,2}:\d{2}:\d{2}$', ts):
                hint_label.config(text="⚠️ 时间格式不对，请用 HH:MM:SS！", fg="red")
                return False
            time_reviewed[r] = ts
        return True

    def go_previous(event=None):
        if not try_save():
            return
        idx = state['index']
        if idx == 0:
            return
        state['index'] -= 1
        load_row(state['index'])

    def go_next(event=None):
        if not try_save():
            return
        idx = state['index']
        if idx + 1 >= total:
            hint_label.config(text="已是最后一条，请按回车提交", fg="orange")
            return
        state['index'] += 1
        load_row(state['index'])

    def confirm_submit(event=None):
        if not try_save():
            return
        state['confirmed'] = True
        root.destroy()

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=12)
    btn_ref['prev'] = tk.Button(button_frame, text="↑ 上一条", command=go_previous, width=10)
    btn_ref['prev'].grid(row=0, column=0, padx=4)
    btn_ref['next'] = tk.Button(button_frame, text="↓ 下一条", command=go_next, width=10)
    btn_ref['next'].grid(row=0, column=1, padx=4)
    btn_ref['submit'] = tk.Button(button_frame, text="确认提交（回车）", command=confirm_submit, width=16, fg="blue")
    btn_ref['submit'].grid(row=0, column=2, padx=4)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=3, padx=4)

    root.bind('<Up>', go_previous)
    root.bind('<Down>', go_next)
    root.bind('<Return>', confirm_submit)

    print("\n已打开【拍摄日期+时间】核对窗口...")
    load_row(0)
    root.lift()
    root.focus_force()
    date_entry.focus_set()
    root.mainloop()

    if state['confirmed']:
        print("【拍摄日期+时间】核对完成，用户已确认提交。")
    else:
        print("【拍摄日期+时间】核对已跳过。")
    return date_reviewed, time_reviewed


def interactive_review_datetime_cli(data_rows_info, date_results, time_results, dt_debug_map):
    """命令行兜底：日期时间核对。"""
    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)
    print(f"\n{'=' * 50}\n进入【拍摄日期+时间】核对（命令行）\n{'=' * 50}")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        dc = date_reviewed.get(r)
        tc = time_reviewed.get(r)
        print(f"\n[{i}/{total}] {info['device_name']}  日期={dc}  时间={tc}")
        if dt_debug_map.get(r):
            open_image_for_preview(dt_debug_map[r])
        while True:
            val = input("  日期 YYYY-MM-DD (回车=保留): ").strip()
            if val == "":
                break
            try:
                date_reviewed[r] = datetime.date.fromisoformat(val)
                break
            except ValueError:
                print("  格式不对。")
        while True:
            val = input("  时间 HH:MM:SS (回车=保留): ").strip()
            if val == "":
                break
            if re.match(r'^\d{1,2}:\d{2}:\d{2}$', val):
                time_reviewed[r] = val
                break
            print("  格式不对。")
        if input("  s=跳过剩余: ").strip().lower() in ("s", "skip"):
            break
    return date_reviewed, time_reviewed


def _is_float(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


# ===================== 主流程第一部分：识别 + 核对 + 写入Excel =====================

def step1_recognize_and_review():
    print(f"正在读取Excel模板: {EXCEL_PATH}  (工作表: {EXCEL_SHEET_NAME})")
    wb = load_workbook(EXCEL_PATH)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        print(f"错误：Excel中找不到名为'{EXCEL_SHEET_NAME}'的工作表。")
        sys.exit(1)
    ws = wb[EXCEL_SHEET_NAME]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行，请检查表格内容。")
        sys.exit(1)
    print(f"共读取到 {len(data_rows_info)} 行设备数据。")

    image_files = collect_image_files(IMAGE_DIR)
    if len(image_files) < len(data_rows_info):
        print(f"警告：图谱文件夹只有 {len(image_files)} 张图片，少于Excel的 {len(data_rows_info)} 行数据，"
              f"部分行可能无法识别。")

    row_to_image = match_images_by_number(image_files, data_rows_info)

    sample_path = next((p for p in row_to_image.values() if p), None)
    if sample_path is None:
        print("错误：没有可用的示例图片用于框选区域。")
        sys.exit(1)

    ar1_roi = get_roi(sample_path, AR1_ROI_CONFIG_FILE, "温度（Ar1）数值",
                       "框选温度(Ar1)区域 (Space/Enter确认, C取消)")
    datetime_roi = get_roi(sample_path, DATETIME_ROI_CONFIG_FILE, "拍摄日期+时间",
                            "框选拍摄日期+时间区域 (Space/Enter确认, C取消)")

    if SAVE_DEBUG_CROPS:
        os.makedirs(DEBUG_DIR, exist_ok=True)

    ar1_results = {}
    date_results = {}
    time_results = {}
    ar1_debug_map = {}
    dt_debug_map = {}

    print("\n开始批量识别，请稍候...\n")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        img_path = row_to_image.get(r)
        if img_path is None:
            ar1_results[r] = None
            date_results[r] = None
            time_results[r] = None
            ar1_debug_map[r] = None
            dt_debug_map[r] = None
            print(f"[{i}/{len(data_rows_info)}] 设备: {info['device_name']}  ->  无对应图片")
            continue

        ar1_val, ar1_crop, ar1_raw = ocr_ar1_temperature(img_path, ar1_roi)
        date_val, time_val, dt_crop, dt_raw = ocr_datetime(img_path, datetime_roi)

        ar1_results[r] = ar1_val
        date_results[r] = date_val
        time_results[r] = time_val

        status = f"温度={ar1_val if ar1_val is not None else NOT_RECOGNIZED_TEXT}  " \
                 f"日期={date_val if date_val else NOT_RECOGNIZED_TEXT}  " \
                 f"时间={time_val if time_val else NOT_RECOGNIZED_TEXT}"
        print(f"[{i}/{len(data_rows_info)}] {os.path.basename(img_path)}  ->  {status}")

        ar1_debug_map[r] = None
        dt_debug_map[r] = None
        if SAVE_DEBUG_CROPS:
            if ar1_crop is not None:
                name = re.sub(r'[\\/:*?"<>|]', "_", f"row{r}_ar1.png")
                save_img = cv2.resize(ar1_crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                path = os.path.join(DEBUG_DIR, name)
                cv2.imwrite(path, save_img)
                ar1_debug_map[r] = path
            if dt_crop is not None:
                name = re.sub(r'[\\/:*?"<>|]', "_", f"row{r}_datetime.png")
                save_img = cv2.resize(dt_crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                path = os.path.join(DEBUG_DIR, name)
                cv2.imwrite(path, save_img)
                dt_debug_map[r] = path

    fail_count = sum(
        1 for info in data_rows_info
        if ar1_results.get(info['row']) is None
        or date_results.get(info['row']) is None
        or time_results.get(info['row']) is None
    )

    do_review = True
    if fail_count == 0:
        print(f"\n本次 {len(data_rows_info)} 行全部识别成功。")
        choice = input("是否仍要逐张人工核对确认？(y=核对 / n=跳过，直接使用识别结果) [n]: ").strip().lower()
        do_review = choice in ("y", "yes")
    else:
        print(f"\n共有 {fail_count} / {len(data_rows_info)} 行存在识别失败的内容，需要人工核对补充。")

    if do_review:
        if USE_GUI:
            ar1_results = interactive_review_ar1_gui(data_rows_info, ar1_results, ar1_debug_map)
            date_results, time_results = interactive_review_datetime_gui(
                data_rows_info, date_results, time_results, dt_debug_map
            )
        else:
            print("\n（已启用 --cli 模式，使用命令行核对）")
            ar1_results = interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map)
            date_results, time_results = interactive_review_datetime_cli(
            data_rows_info, date_results, time_results, dt_debug_map
        )

    # 写回Excel
    for info in data_rows_info:
        r = info['row']
        ar1_val = ar1_results.get(r)
        date_val = date_results.get(r)
        time_val = time_results.get(r)

        ws.cell(row=r, column=10).value = float(ar1_val) if ar1_val is not None else NOT_RECOGNIZED_TEXT
        ws.cell(row=r, column=11).value = date_val if date_val is not None else NOT_RECOGNIZED_TEXT
        ws.cell(row=r, column=12).value = time_val if time_val is not None else NOT_RECOGNIZED_TEXT

    wb.save(EXCEL_OUTPUT_PATH)
    print(f"\nExcel已生成: {EXCEL_OUTPUT_PATH}")

    full_data = []
    for info in data_rows_info:
        r = info['row']
        row_data = dict(info)
        row_data['ar1'] = ar1_results.get(r)
        row_data['date'] = date_results.get(r)
        row_data['time'] = time_results.get(r)
        row_data['image'] = row_to_image.get(r)
        full_data.append(row_data)

    return full_data


# ===================== 主流程第二部分：生成 91号文报告_已生成.docx =====================

def apply_font_to_run(run):
    """对单个run设置字体：中文宋体、英文Times New Roman、字号8.5pt。"""
    run.font.name = WESTERN_FONT
    run.font.size = FONT_SIZE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rFonts.set(qn('w:ascii'), WESTERN_FONT)
    rFonts.set(qn('w:hAnsi'), WESTERN_FONT)


def set_cell_text(cell, text):
    """清空单元格已有内容，写入新文本，并对新文本应用字体（宋体/TNR/8.5pt）。"""
    text = "" if text is None else str(text)
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    apply_font_to_run(run)
    for extra_p in cell.paragraphs[1:]:
        extra_p.text = ""


def visual_cells(row):
    """按底层XML元素身份去重（正确识别真实的合并单元格），返回去重后的Cell对象列表。"""
    result = []
    last_tc = None
    for cell in row.cells:
        if cell._tc is not last_tc:
            result.append(cell)
            last_tc = cell._tc
    return result


def format_value(val, suffix=""):
    if val is None or str(val).strip() == "":
        return ""
    return f"{val}{suffix}"


def fill_table0(table, row_data):
    rows = table.rows
    # row1: 设备类型 -> raw idx5~7 (合并)
    set_cell_text(rows[0].cells[5], row_data['device_type'])
    # row2: 间隔单元 -> raw idx1~3；相别 -> raw idx5~7
    set_cell_text(rows[1].cells[1], row_data['interval'])
    set_cell_text(rows[1].cells[5], row_data['phase'])
    # row3: 设备名称 -> raw idx1~3；测试距离 -> raw idx5~7
    set_cell_text(rows[2].cells[1], row_data['device_name'])
    set_cell_text(rows[2].cells[5], row_data['distance'])
    # row4: 运行电压 -> raw idx1~3；负荷电流 -> raw idx5；有功功率 -> raw idx7
    set_cell_text(rows[3].cells[1], row_data['voltage'])
    set_cell_text(rows[3].cells[5], row_data['current'])
    set_cell_text(rows[3].cells[7], row_data['power'])
    # row5: 拍摄日期 -> raw idx1~3；拍摄时间 -> raw idx5~7
    date_str = row_data['date'].strftime('%Y-%m-%d') if isinstance(row_data['date'], datetime.date) else row_data['date']
    set_cell_text(rows[4].cells[1], date_str if date_str else NOT_RECOGNIZED_TEXT)
    set_cell_text(rows[4].cells[5], row_data['time'] if row_data['time'] else NOT_RECOGNIZED_TEXT)


def fill_table1(table, row_data):
    # 唯一一行: 红外分析 | Ar1(标签) | 值(raw idx2) | Ar2 | | Ar3 | |
    row = table.rows[0]
    ar1_display = f"{row_data['ar1']}℃" if row_data['ar1'] not in (None, NOT_RECOGNIZED_TEXT) else NOT_RECOGNIZED_TEXT
    set_cell_text(row.cells[2], ar1_display)


def fill_table2(table, row_data):
    # row2(0索引row1) 是一个横跨全部raw列的合并单元格，用于插入图片
    row = table.rows[1]
    cell = row.cells[0]
    for p in cell.paragraphs:
        for run in list(p.runs):
            run.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    img_path = row_data.get('image')
    if img_path and os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM), height=Cm(IMAGE_HEIGHT_CM))
    else:
        run = p.add_run("（未找到对应图片）")


def add_page_break(body, before_element):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    body.insert(list(body).index(before_element), p)


def step2_generate_docx(full_data):
    print(f"\n正在生成91号文报告: {DOCX_TEMPLATE_PATH}")
    doc = Document(DOCX_TEMPLATE_PATH)
    body = doc.element.body

    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        raise RuntimeError("模板文档缺少页面设置(sectPr)，无法继续。")

    # 捕获模板原有的内容块（标题、3个表格、结尾段落及它们之间的空段落），深拷贝一份留作反复使用的"模板块"
    template_children = [child for child in body if child is not sectPr]
    template_block = [copy.deepcopy(el) for el in template_children]

    for el in template_children:
        body.remove(el)

    for idx, row_data in enumerate(full_data):
        if idx > 0:
            add_page_break(body, sectPr)

        block_copy = [copy.deepcopy(el) for el in template_block]
        for el in block_copy:
            body.insert(list(body).index(sectPr), el)

        tbls = [Table(el, doc) for el in block_copy if el.tag == qn('w:tbl')]
        if len(tbls) != 3:
            raise RuntimeError(f"模板表格数量异常，期望3个，实际{len(tbls)}个。")
        table0, table1, table2 = tbls

        fill_table0(table0, row_data)
        fill_table1(table1, row_data)
        fill_table2(table2, row_data)

        print(f"  [{idx + 1}/{len(full_data)}] 已生成: {row_data['device_name']}")

    doc.save(DOCX_OUTPUT_PATH)
    print(f"\n91号文报告已生成: {DOCX_OUTPUT_PATH}")


# ===================== 主程序 =====================

def main():
    global USE_GUI
    if "--cli" in sys.argv:
        USE_GUI = False
        print("已启用 --cli 模式：将使用命令行进行人工核对（避免macOS上tkinter间歇性崩溃）。")

    # 预初始化tkinter（在OpenCV弹出任何窗口之前），避免macOS上Cocoa资源冲突
    if USE_GUI:
        try:
            import tkinter as tk
            _pre = tk.Tk()
            _pre.withdraw()
            _pre.destroy()
        except Exception:
            pass

    if not os.path.isdir(IMAGE_DIR):
        print(f"错误：图谱文件夹不存在: {IMAGE_DIR}")
        sys.exit(1)
    if not os.path.isfile(EXCEL_PATH):
        print(f"错误：Excel文件不存在: {EXCEL_PATH}")
        sys.exit(1)
    if not os.path.isfile(DOCX_TEMPLATE_PATH):
        print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
        sys.exit(1)

    full_data = step1_recognize_and_review()
    step2_generate_docx(full_data)

    print("\n全部完成！")
    print(f"  - Excel结果: {EXCEL_OUTPUT_PATH}")
    print(f"  - Word报告: {DOCX_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
