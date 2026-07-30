# -*- coding: utf-8 -*-
"""
91号文《红外精确测温报告》自动生成程序
======================================================

本脚本假设以下三样东西和本脚本放在【同一个文件夹】里（不需要修改任何路径配置）：
    ./图谱/                图片文件夹（红外精确测温图片，按文件名顺序对应Excel每一行）
    ./必填信息.xlsx         数据表（工作表名："91号文模板"，已提前填好B~I列，J/K/L列由本脚本识别）
    ./91号文报告.docx        Word报告模板（每份模板只装一个设备的信息，本脚本会自动复制成多份、
                             每个设备占一页，拼接成一份完整报告）

【整体流程】
    1. 识别：分别框选"温度(Ar1)"、"拍摄日期"、"拍摄时间"三个独立区域（各框一次，所有图片复用），
       对图谱文件夹里的每张图片（按文件名顺序，依次对应Excel第2、3、4...行）自动识别。
    2. 人工核对：弹出一个界面，每一行同时显示"温度裁剪图+输入框""日期裁剪图+输入框"
       "时间裁剪图+输入框"，回车确认并下一条，↑回到上一条，可随时跳过剩余核对。
    3. 核对完成后，把J（Ar1温度）K（拍摄日期）L（拍摄时间）写入必填信息_已生成.xlsx。
    4. 根据必填信息_已生成.xlsx和图谱文件夹里的图片，生成91号文报告_已生成.docx：
       每个设备复制一份模板、填入对应信息和图片，每个设备占一页，全部设备依次拼接成一份文档。

依赖安装（与43号文脚本共用同一个虚拟环境即可，无需重新安装）：
    pip install opencv-python numpy pytesseract openpyxl pillow python-docx

还需要系统安装 Tesseract OCR 引擎本体（如果43号文脚本能跑，这里不需要重装）：
    brew install tesseract   或   conda install -c conda-forge tesseract

使用方法：
    python3 91号文报告自动生成.py
"""

import os
import re
import sys
import json
import glob
import copy
import platform
import subprocess
import datetime

import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook

from docx import Document
from docx.shared import Cm
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ===================== 配置区域 =====================

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

IMAGE_DIR = os.path.join(BASE_DIR, "图谱")
EXCEL_PATH = os.path.join(BASE_DIR, "必填信息.xlsx")
EXCEL_OUTPUT_PATH = os.path.join(BASE_DIR, "必填信息_已生成.xlsx")
EXCEL_SHEET_NAME = "91号文模板"

DOCX_TEMPLATE_PATH = os.path.join(BASE_DIR, "91号文报告.docx")
DOCX_OUTPUT_PATH = os.path.join(BASE_DIR, "91号文报告_已生成.docx")

AR1_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_ar1.json")
DATE_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_date.json")
TIME_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_time.json")

DEBUG_DIR = os.path.join(BASE_DIR, "debug_crops_91")
SAVE_DEBUG_CROPS = True

NOT_RECOGNIZED_TEXT = "无法识别"

# 插入图片尺寸
IMAGE_WIDTH_CM = 15
IMAGE_HEIGHT_CM = 12

# =================================================================


# ===================== 工具函数 =====================

def natural_sort_key(s):
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def collect_image_files(image_dir):
    patterns = ["*.jpg", "*.JPG", "*.jpeg", "*.JPEG", "*.png", "*.PNG"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(image_dir, p)))
    files = sorted(set(files), key=lambda p: natural_sort_key(os.path.basename(p)))
    return files


def open_image_for_preview(image_path):
    try:
        system = platform.system()
        if system == "Darwin":
            subprocess.run(["open", image_path], check=False)
        elif system == "Windows":
            os.startfile(image_path)  # noqa
        else:
            subprocess.run(["xdg-open", image_path], check=False)
    except Exception as e:
        print(f"  （提示：自动打开图片失败：{e}，可手动查看：{image_path}）")


# ===================== ROI框选（相对比例，两个独立区域） =====================

def select_roi_interactively(sample_image_path, window_title):
    img = cv2.imread(sample_image_path)
    if img is None:
        raise FileNotFoundError(f"无法读取示例图片: {sample_image_path}")
    img_h, img_w = img.shape[:2]

    roi = cv2.selectROI(window_title, img, showCrosshair=True, fromCenter=False)
    cv2.destroyAllWindows()

    x, y, w, h = roi
    if w == 0 or h == 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    rx, ry, rw, rh = x / img_w, y / img_h, w / img_w, h / img_h
    print(f"框选区域（像素）: x={x}, y={y}, w={w}, h={h}  (示例图片尺寸: {img_w}x{img_h})")
    return rx, ry, rw, rh


def match_images_by_number(image_files, data_rows_info):
    """
    把图谱文件夹里的图片对应到Excel每一行。
    优先方式：从文件名里提取数字（比如"图5.jpg"->5），第N号图片对应第N行数据。
    这样即使中间缺了一张图（比如"图5.jpg"不存在），也只会让第5行没有图，
    不会导致第6行及以后的图全部错位往前顶。
    如果文件名根本不含数字规律，才退回成"按文件夹内文件名顺序直接依次对应"。
    返回：{Excel行号: 图片路径 或 None}
    """
    numbered = {}
    all_numbered = True
    for p in image_files:
        base = os.path.splitext(os.path.basename(p))[0]
        m = re.search(r'(\d+)', base)
        if m:
            numbered[int(m.group(1))] = p
        else:
            all_numbered = False

    row_to_image = {}
    if all_numbered and numbered:
        print("图片匹配方式：按文件名中的数字精确对应第几行（例如'图5.jpg'对应第5行数据）。")
        for i, info in enumerate(data_rows_info, start=1):
            row_to_image[info['row']] = numbered.get(i)
    else:
        print("图片匹配方式：文件名不含数字规律，改为按文件夹内文件名顺序直接依次对应。")
        for i, info in enumerate(data_rows_info):
            row_to_image[info['row']] = image_files[i] if i < len(image_files) else None

    missing = [info['row'] for info in data_rows_info if row_to_image.get(info['row']) is None]
    if missing:
        print(f"提示：以下Excel行号没有匹配到图片，需要人工确认：{missing}")

    return row_to_image


def get_roi(sample_image_path, config_file, prompt_text, window_title):
    if os.path.exists(config_file):
        with open(config_file, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        rx, ry, rw, rh = cfg["rx"], cfg["ry"], cfg["rw"], cfg["rh"]
        print(f"检测到已保存的【{prompt_text}】区域配置: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
        choice = input(f"是否复用该区域？(y=复用 / n=重新框选) [y]: ").strip().lower()
        if choice in ("", "y", "yes"):
            return rx, ry, rw, rh

    print(f"\n请在弹出的窗口中框选【{prompt_text}】区域，框选完成后按【空格/回车】确认，按 c 取消重选。")
    rx, ry, rw, rh = select_roi_interactively(sample_image_path, window_title)
    with open(config_file, "w", encoding="utf-8") as f:
        json.dump({"rx": rx, "ry": ry, "rw": rw, "rh": rh}, f, ensure_ascii=False, indent=2)
    print(f"识别区域已保存到: {config_file}")
    return rx, ry, rw, rh


def crop_by_relative_roi(img, roi_ratio):
    rx, ry, rw, rh = roi_ratio
    img_h, img_w = img.shape[:2]
    x = int(round(rx * img_w))
    y = int(round(ry * img_h))
    w = int(round(rw * img_w))
    h = int(round(rh * img_h))
    x2, y2 = min(x + w, img_w), min(y + h, img_h)
    x, y = max(x, 0), max(y, 0)
    return img[y:y2, x:x2]


def preprocess_variants(crop, scale=4):
    crop_big = cv2.resize(crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(crop_big, cv2.COLOR_BGR2GRAY)

    variants = []
    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh1 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh1)

    inv = cv2.bitwise_not(blur)
    _, thresh2 = cv2.threshold(inv, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh2)

    thresh3 = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
    )
    variants.append(thresh3)

    return variants


# ===================== OCR：Ar1温度（纯数字） =====================

def extract_number(text):
    if not text:
        return None
    match = re.search(r'-?\d+\.?\d*', text.strip())
    if not match:
        return None
    try:
        num = float(match.group())
        return str(int(num)) if num == int(num) else str(num)
    except ValueError:
        return None


def ocr_ar1_temperature(image_path, roi_ratio):
    img = cv2.imread(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, ["[裁剪区域为空]"]

    variants = preprocess_variants(crop)
    configs = [
        r'--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 11 -c tessedit_char_whitelist=0123456789.-',
    ]
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            value = extract_number(text)
            if value is not None:
                return value, crop, raw_texts
    return None, crop, raw_texts


# ===================== OCR：拍摄日期 / 拍摄时间（各自独立识别） =====================

DATE_PATTERN = re.compile(r'(\d{4})\D(\d{1,2})\D(\d{1,2})')
TIME_PATTERN = re.compile(r'(\d{1,2}):(\d{2}):(\d{2})')


def parse_date(text):
    if not text:
        return None
    dm = DATE_PATTERN.search(text)
    if not dm:
        return None
    try:
        y, mo, d = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
        return datetime.date(y, mo, d)
    except ValueError:
        return None


def parse_time(text):
    if not text:
        return None
    tm = TIME_PATTERN.search(text)
    if not tm:
        return None
    h, mi, s = tm.group(1), tm.group(2), tm.group(3)
    try:
        if 0 <= int(h) <= 23 and 0 <= int(mi) <= 59 and 0 <= int(s) <= 59:
            return f"{int(h):02d}:{mi}:{s}"
    except ValueError:
        pass
    return None


def _ocr_with_configs(crop, configs):
    variants = preprocess_variants(crop)
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            yield text, raw_texts


DATE_TIME_CHAR_WHITELIST = r'-c tessedit_char_whitelist=0123456789.:-年月日 '
_OCR_CONFIGS_DATE_TIME = [
    rf'--oem 3 --psm 7 {DATE_TIME_CHAR_WHITELIST}',
    rf'--oem 3 --psm 6 {DATE_TIME_CHAR_WHITELIST}',
]


def ocr_date(image_path, roi_ratio):
    img = cv2.imread(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, crop, ["[裁剪区域为空]"]

    raw_texts = []
    for text, raw_texts in _ocr_with_configs(crop, _OCR_CONFIGS_DATE_TIME):
        date_val = parse_date(text)
        if date_val is not None:
            return date_val, crop, raw_texts
    return None, crop, raw_texts


def ocr_time(image_path, roi_ratio):
    img = cv2.imread(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, crop, ["[裁剪区域为空]"]

    raw_texts = []
    for text, raw_texts in _ocr_with_configs(crop, _OCR_CONFIGS_DATE_TIME):
        time_val = parse_time(text)
        if time_val is not None:
            return time_val, crop, raw_texts
    return None, crop, raw_texts


# ===================== Excel读取 =====================

def find_header_row(ws, keyword="设备类型", max_scan_rows=5):
    for row in range(1, max_scan_rows + 1):
        for col in range(1, 5):
            if str(ws.cell(row=row, column=col).value).strip() == keyword:
                return row
    return 1


def read_excel_data_rows(ws, header_row_idx):
    """从表头下一行开始读取，直到设备名称(D列)为空为止。"""
    data_rows_info = []
    row = header_row_idx + 1
    while True:
        d_val = ws.cell(row=row, column=4).value  # D列 设备名称
        if d_val is None or str(d_val).strip() == "":
            break
        info = {
            'row': row,
            'device_type': ws.cell(row=row, column=2).value,     # B 设备类型
            'interval': ws.cell(row=row, column=3).value,         # C 间隔单元
            'device_name': d_val,                                    # D 设备名称
            'phase': ws.cell(row=row, column=5).value,             # E 相别
            'distance': ws.cell(row=row, column=6).value,          # F 测试距离
            'voltage': ws.cell(row=row, column=7).value,           # G 运行电压
            'current': ws.cell(row=row, column=8).value,           # H 负荷电流
            'power': ws.cell(row=row, column=9).value,             # I 有功功率
        }
        data_rows_info.append(info)
        row += 1
    return data_rows_info


# ===================== 人工核对（图形界面，Ar1+日期时间合并在一个窗口） =====================

def interactive_review_gui(data_rows_info, ar1_results, date_results, time_results,
                            ar1_debug_map, date_debug_map, time_debug_map):
    """
    单窗口人工核对：每一行同时显示三个独立区块——
    温度(Ar1)裁剪图+输入框、日期裁剪图+输入框、时间裁剪图+输入框。
    回车=确认并下一条，↑=上一条，可点击按钮跳过剩余核对。
    如果tkinter/Pillow不可用，自动回退到命令行版本。
    """
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（提示：图形界面核对工具不可用（{e}），自动改用命令行方式核对。）")
        print("  - Homebrew Python: brew install python-tk@3.10 （版本号按你的Python调整）")
        print("  - conda环境通常自带tkinter")
        return interactive_review_cli(data_rows_info, ar1_results, date_results, time_results,
                                       ar1_debug_map, date_debug_map, time_debug_map)

    ar1_reviewed = dict(ar1_results)
    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)
    state = {'index': 0}
    photo_refs = {'ar1': None, 'date': None, 'time': None}
    prev_btn_ref = {}

    root = tk.Tk()
    root.title("91号文测温 - 人工核对")
    root.geometry("680x880")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 13), justify="left", wraplength=640)
    info_label.pack(pady=(12, 6))

    def build_section(title):
        tk.Label(root, text=title, font=("PingFang SC", 11, "bold")).pack(pady=(8, 0))
        img_label = tk.Label(root)
        img_label.pack(pady=4)
        var = tk.StringVar()
        entry = tk.Entry(root, textvariable=var, font=("PingFang SC", 16), justify="center", width=16)
        entry.pack(pady=4)
        return img_label, var, entry

    ar1_img_label, ar1_var, ar1_entry = build_section("温度（Ar1）")
    date_img_label, date_var, date_entry = build_section("拍摄日期（YYYY-MM-DD）")
    time_img_label, time_var, time_entry = build_section("拍摄时间（HH:MM:SS）")

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(10, 4))

    DEFAULT_HINT = "↑ = 上一条    |    回车 = 确认并下一条"

    def load_image_into(label, path, ref_key, max_w=560, max_h=180):
        if path and os.path.exists(path):
            pil_img = Image.open(path)
            ratio = min(max_w / pil_img.width, max_h / pil_img.height, 1.0)
            new_size = (max(1, int(pil_img.width * ratio)), max(1, int(pil_img.height * ratio)))
            pil_img = pil_img.resize(new_size)
            photo = ImageTk.PhotoImage(pil_img)
            photo_refs[ref_key] = photo
            label.config(image=photo, text="")
        else:
            photo_refs[ref_key] = None
            label.config(image="", text="（未找到对应裁剪图片）")

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']

        ar1_current = ar1_reviewed.get(r)
        date_current = date_reviewed.get(r)
        time_current = time_reviewed.get(r)
        mark = ""
        if ar1_current is None or date_current is None or time_current is None:
            mark = "   ⚠️ 有内容识别失败，需手动输入"

        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}{mark}")

        load_image_into(ar1_img_label, ar1_debug_map.get(r), 'ar1')
        load_image_into(date_img_label, date_debug_map.get(r), 'date')
        load_image_into(time_img_label, time_debug_map.get(r), 'time')

        ar1_var.set(ar1_current if ar1_current is not None else "")
        date_var.set(date_current.isoformat() if date_current else "")
        time_var.set(time_current if time_current else "")

        hint_label.config(text=DEFAULT_HINT, fg="gray")
        ar1_entry.focus_set()
        ar1_entry.select_range(0, tk.END)

        if 'prev' in prev_btn_ref:
            prev_btn_ref['prev'].config(state=(tk.DISABLED if idx == 0 else tk.NORMAL))

    def try_save_current():
        idx = state['index']
        r = data_rows_info[idx]['row']

        ar1_val = ar1_var.get().strip()
        if ar1_val != "":
            try:
                float(ar1_val)
                ar1_reviewed[r] = ar1_val
            except ValueError:
                hint_label.config(text="⚠️ 温度不是有效数字，请重新输入！", fg="red")
                return False

        date_val_str = date_var.get().strip()
        if date_val_str != "":
            try:
                date_reviewed[r] = datetime.date.fromisoformat(date_val_str)
            except ValueError:
                hint_label.config(text="⚠️ 日期格式不对，请用 YYYY-MM-DD 格式！", fg="red")
                return False

        time_val_str = time_var.get().strip()
        if time_val_str != "":
            if not re.match(r'^\d{1,2}:\d{2}:\d{2}$', time_val_str):
                hint_label.config(text="⚠️ 时间格式不对，请用 HH:MM:SS 格式！", fg="red")
                return False
            time_reviewed[r] = time_val_str

        return True

    def confirm_and_next(event=None):
        if not try_save_current():
            return
        idx = state['index']
        if idx + 1 >= total:
            root.destroy()
        else:
            state['index'] += 1
            load_row(state['index'])

    def go_previous(event=None):
        if not try_save_current():
            return
        idx = state['index']
        if idx == 0:
            hint_label.config(text="已经是第一条了，无法再往前。", fg="orange")
            return
        state['index'] -= 1
        load_row(state['index'])

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=14)
    prev_btn_ref['prev'] = tk.Button(button_frame, text="⬆ 上一条", command=go_previous, width=12)
    prev_btn_ref['prev'].grid(row=0, column=0, padx=6)
    tk.Button(button_frame, text="确认，下一条（回车）", command=confirm_and_next, width=20).grid(row=0, column=1, padx=6)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=2, padx=6)

    root.bind('<Return>', confirm_and_next)
    root.bind('<Up>', go_previous)

    print("\n已打开图形界面核对窗口（↑=上一条，回车=确认下一条）...")
    load_row(0)
    root.lift()
    root.focus_force()
    ar1_entry.focus_set()
    root.mainloop()

    print("人工核对完成（或已跳过剩余行）。")
    return ar1_reviewed, date_reviewed, time_reviewed


def interactive_review_cli(data_rows_info, ar1_results, date_results, time_results,
                            ar1_debug_map, date_debug_map, time_debug_map):
    """命令行版兜底核对方案。"""
    ar1_reviewed = dict(ar1_results)
    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)

    print("\n" + "=" * 60)
    print("进入人工核对环节（命令行模式）")
    print("=" * 60)

    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        ar1_current = ar1_reviewed.get(r)
        date_current = date_reviewed.get(r)
        time_current = time_reviewed.get(r)

        print(f"\n[{i}/{total}] 设备: {info['device_name']}")
        print(f"  当前 温度={ar1_current}  拍摄日期={date_current}  拍摄时间={time_current}")

        if ar1_debug_map.get(r):
            open_image_for_preview(ar1_debug_map[r])
        if date_debug_map.get(r):
            open_image_for_preview(date_debug_map[r])
        if time_debug_map.get(r):
            open_image_for_preview(time_debug_map[r])

        while True:
            val = input("  温度 (回车=保留当前值): ").strip()
            if val == "" or _is_float(val):
                if val != "":
                    ar1_reviewed[r] = val
                break
            print("  不是有效数字，请重新输入。")

        while True:
            val = input("  拍摄日期 YYYY-MM-DD (回车=保留当前值): ").strip()
            if val == "":
                break
            try:
                date_reviewed[r] = datetime.date.fromisoformat(val)
                break
            except ValueError:
                print("  日期格式不对，请用 YYYY-MM-DD。")

        while True:
            val = input("  拍摄时间 HH:MM:SS (回车=保留当前值): ").strip()
            if val == "":
                break
            if re.match(r'^\d{1,2}:\d{2}:\d{2}$', val):
                time_reviewed[r] = val
                break
            print("  时间格式不对，请用 HH:MM:SS。")

        skip = input("  输入 s 跳过剩余核对，其他任意键继续下一条: ").strip().lower()
        if skip in ("s", "skip"):
            print("已跳过剩余行核对。")
            break

    return ar1_reviewed, date_reviewed, time_reviewed


def _is_float(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


# ===================== 主流程第一部分：识别 + 核对 + 写入Excel =====================

def step1_recognize_and_review():
    print(f"正在读取Excel模板: {EXCEL_PATH}  (工作表: {EXCEL_SHEET_NAME})")
    wb = load_workbook(EXCEL_PATH)
    if EXCEL_SHEET_NAME not in wb.sheetnames:
        print(f"错误：Excel中找不到名为'{EXCEL_SHEET_NAME}'的工作表。")
        sys.exit(1)
    ws = wb[EXCEL_SHEET_NAME]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行，请检查表格内容。")
        sys.exit(1)
    print(f"共读取到 {len(data_rows_info)} 行设备数据。")

    image_files = collect_image_files(IMAGE_DIR)
    if len(image_files) < len(data_rows_info):
        print(f"警告：图谱文件夹只有 {len(image_files)} 张图片，少于Excel的 {len(data_rows_info)} 行数据，"
              f"部分行可能无法识别。")

    row_to_image = match_images_by_number(image_files, data_rows_info)

    sample_path = next((p for p in row_to_image.values() if p), None)
    if sample_path is None:
        print("错误：没有可用的示例图片用于框选区域。")
        sys.exit(1)

    ar1_roi = get_roi(sample_path, AR1_ROI_CONFIG_FILE, "温度（Ar1）数值",
                       "框选温度(Ar1)区域 (Space/Enter确认, C取消)")
    date_roi = get_roi(sample_path, DATE_ROI_CONFIG_FILE, "拍摄日期",
                        "框选拍摄日期区域 (Space/Enter确认, C取消)")
    time_roi = get_roi(sample_path, TIME_ROI_CONFIG_FILE, "拍摄时间",
                        "框选拍摄时间区域 (Space/Enter确认, C取消)")

    if SAVE_DEBUG_CROPS:
        os.makedirs(DEBUG_DIR, exist_ok=True)

    ar1_results = {}
    date_results = {}
    time_results = {}
    ar1_debug_map = {}
    date_debug_map = {}
    time_debug_map = {}

    print("\n开始批量识别，请稍候...\n")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        img_path = row_to_image.get(r)
        if img_path is None:
            ar1_results[r] = None
            date_results[r] = None
            time_results[r] = None
            ar1_debug_map[r] = None
            date_debug_map[r] = None
            time_debug_map[r] = None
            print(f"[{i}/{len(data_rows_info)}] 设备: {info['device_name']}  ->  无对应图片")
            continue

        ar1_val, ar1_crop, ar1_raw = ocr_ar1_temperature(img_path, ar1_roi)
        date_val, date_crop, date_raw = ocr_date(img_path, date_roi)
        time_val, time_crop, time_raw = ocr_time(img_path, time_roi)

        ar1_results[r] = ar1_val
        date_results[r] = date_val
        time_results[r] = time_val

        status = f"温度={ar1_val if ar1_val is not None else NOT_RECOGNIZED_TEXT}  " \
                 f"日期={date_val if date_val else NOT_RECOGNIZED_TEXT}  " \
                 f"时间={time_val if time_val else NOT_RECOGNIZED_TEXT}"
        print(f"[{i}/{len(data_rows_info)}] {os.path.basename(img_path)}  ->  {status}")

        ar1_debug_map[r] = None
        date_debug_map[r] = None
        time_debug_map[r] = None
        if SAVE_DEBUG_CROPS:
            for crop, tag, target_map in (
                (ar1_crop, "ar1", ar1_debug_map),
                (date_crop, "date", date_debug_map),
                (time_crop, "time", time_debug_map),
            ):
                if crop is not None:
                    name = re.sub(r'[\\/:*?"<>|]', "_", f"row{r}_{tag}.png")
                    save_img = cv2.resize(crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                    path = os.path.join(DEBUG_DIR, name)
                    cv2.imwrite(path, save_img)
                    target_map[r] = path

    fail_count = sum(
        1 for info in data_rows_info
        if ar1_results.get(info['row']) is None
        or date_results.get(info['row']) is None
        or time_results.get(info['row']) is None
    )

    do_review = True
    if fail_count == 0:
        print(f"\n本次 {len(data_rows_info)} 行全部识别成功。")
        choice = input("是否仍要逐张人工核对确认？(y=核对 / n=跳过，直接使用识别结果) [n]: ").strip().lower()
        do_review = choice in ("y", "yes")
    else:
        print(f"\n共有 {fail_count} / {len(data_rows_info)} 行存在识别失败的内容，需要人工核对补充。")

    if do_review:
        ar1_results, date_results, time_results = interactive_review_gui(
            data_rows_info, ar1_results, date_results, time_results,
            ar1_debug_map, date_debug_map, time_debug_map
        )

    # 写回Excel
    for info in data_rows_info:
        r = info['row']
        ar1_val = ar1_results.get(r)
        date_val = date_results.get(r)
        time_val = time_results.get(r)

        ws.cell(row=r, column=10).value = float(ar1_val) if ar1_val is not None else NOT_RECOGNIZED_TEXT
        ws.cell(row=r, column=11).value = date_val if date_val is not None else NOT_RECOGNIZED_TEXT
        ws.cell(row=r, column=12).value = time_val if time_val is not None else NOT_RECOGNIZED_TEXT

    wb.save(EXCEL_OUTPUT_PATH)
    print(f"\nExcel已生成: {EXCEL_OUTPUT_PATH}")

    full_data = []
    for info in data_rows_info:
        r = info['row']
        row_data = dict(info)
        row_data['ar1'] = ar1_results.get(r)
        row_data['date'] = date_results.get(r)
        row_data['time'] = time_results.get(r)
        row_data['image'] = row_to_image.get(r)
        full_data.append(row_data)

    return full_data


# ===================== 主流程第二部分：生成 91号文报告_已生成.docx =====================

def set_cell_text(cell, text):
    """清空单元格已有内容，写入新文本，保留单元格原有的格式和字体。"""
    text = "" if text is None else str(text)
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    for extra_p in cell.paragraphs[1:]:
        extra_p.text = ""


def visual_cells(row):
    """按底层XML元素身份去重（正确识别真实的合并单元格），返回去重后的Cell对象列表。"""
    result = []
    last_tc = None
    for cell in row.cells:
        if cell._tc is not last_tc:
            result.append(cell)
            last_tc = cell._tc
    return result


def format_value(val, suffix=""):
    if val is None or str(val).strip() == "":
        return ""
    return f"{val}{suffix}"


def fill_table0(table, row_data):
    rows = table.rows
    # row1: 设备类型 -> raw idx5~7 (合并)
    set_cell_text(rows[0].cells[5], row_data['device_type'])
    # row2: 间隔单元 -> raw idx1~3；相别 -> raw idx5~7
    set_cell_text(rows[1].cells[1], row_data['interval'])
    set_cell_text(rows[1].cells[5], row_data['phase'])
    # row3: 设备名称 -> raw idx1~3；测试距离 -> raw idx5~7
    set_cell_text(rows[2].cells[1], row_data['device_name'])
    set_cell_text(rows[2].cells[5], row_data['distance'])
    # row4: 运行电压 -> raw idx1~3；负荷电流 -> raw idx5；有功功率 -> raw idx7
    set_cell_text(rows[3].cells[1], row_data['voltage'])
    set_cell_text(rows[3].cells[5], row_data['current'])
    set_cell_text(rows[3].cells[7], row_data['power'])
    # row5: 拍摄日期 -> raw idx1~3；拍摄时间 -> raw idx5~7
    date_str = row_data['date'].strftime('%Y-%m-%d') if isinstance(row_data['date'], datetime.date) else row_data['date']
    set_cell_text(rows[4].cells[1], date_str if date_str else NOT_RECOGNIZED_TEXT)
    set_cell_text(rows[4].cells[5], row_data['time'] if row_data['time'] else NOT_RECOGNIZED_TEXT)


def fill_table1(table, row_data):
    # 唯一一行: 红外分析 | Ar1(标签) | 值(raw idx2) | Ar2 | | Ar3 | |
    row = table.rows[0]
    ar1_display = f"{row_data['ar1']}℃" if row_data['ar1'] not in (None, NOT_RECOGNIZED_TEXT) else NOT_RECOGNIZED_TEXT
    set_cell_text(row.cells[2], ar1_display)


def fill_table2(table, row_data):
    # row2(0索引row1) 是一个横跨全部raw列的合并单元格，用于插入图片
    row = table.rows[1]
    cell = row.cells[0]
    for p in cell.paragraphs:
        for run in list(p.runs):
            run.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    img_path = row_data.get('image')
    if img_path and os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM), height=Cm(IMAGE_HEIGHT_CM))
    else:
        run = p.add_run("（未找到对应图片）")


def add_page_break(body, before_element):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    body.insert(list(body).index(before_element), p)


def step2_generate_docx(full_data):
    print(f"\n正在生成91号文报告: {DOCX_TEMPLATE_PATH}")
    doc = Document(DOCX_TEMPLATE_PATH)
    body = doc.element.body

    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        raise RuntimeError("模板文档缺少页面设置(sectPr)，无法继续。")

    # 捕获模板原有的内容块（标题、3个表格、结尾段落及它们之间的空段落），深拷贝一份留作反复使用的"模板块"
    template_children = [child for child in body if child is not sectPr]
    template_block = [copy.deepcopy(el) for el in template_children]

    for el in template_children:
        body.remove(el)

    for idx, row_data in enumerate(full_data):
        if idx > 0:
            add_page_break(body, sectPr)

        block_copy = [copy.deepcopy(el) for el in template_block]
        for el in block_copy:
            body.insert(list(body).index(sectPr), el)

        tbls = [Table(el, doc) for el in block_copy if el.tag == qn('w:tbl')]
        if len(tbls) != 3:
            raise RuntimeError(f"模板表格数量异常，期望3个，实际{len(tbls)}个。")
        table0, table1, table2 = tbls

        fill_table0(table0, row_data)
        fill_table1(table1, row_data)
        fill_table2(table2, row_data)

        print(f"  [{idx + 1}/{len(full_data)}] 已生成: {row_data['device_name']}")

    doc.save(DOCX_OUTPUT_PATH)
    print(f"\n91号文报告已生成: {DOCX_OUTPUT_PATH}")


# ===================== 主程序 =====================

def main():
    if not os.path.isdir(IMAGE_DIR):
        print(f"错误：图谱文件夹不存在: {IMAGE_DIR}")
        sys.exit(1)
    if not os.path.isfile(EXCEL_PATH):
        print(f"错误：Excel文件不存在: {EXCEL_PATH}")
        sys.exit(1)
    if not os.path.isfile(DOCX_TEMPLATE_PATH):
        print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
        sys.exit(1)

    full_data = step1_recognize_and_review()
    step2_generate_docx(full_data)

    print("\n全部完成！")
    print(f"  - Excel结果: {EXCEL_OUTPUT_PATH}")
    print(f"  - Word报告: {DOCX_OUTPUT_PATH}")


if __name__ == "__main__":
    main()
