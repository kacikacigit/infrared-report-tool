# -*- coding: utf-8 -*-
"""
91号文《红外精确测温报告》自动生成程序 V6
======================================================

V6 新增功能：
  1. 温度框不固定自动定位：运行时询问"是否温度框不固定？"，选 y 则
     OCR 整图找到"最高温"的"温"字位置，裁剪其右侧数值区域再局部 OCR，
     Ar1 温度位置每张图不同也能提取；选 n 则沿用之前的手动框选流程。

V5 新增功能：
  1. PyInstaller 打包支持：exe 内嵌便携版 Tesseract OCR，用户无需安装任何依赖

V4 已有功能（全部保留）：
  1. Excel和Word输出全部自动递增编号，不再交互询问版本后缀
     - 必填信息_已生成_1.xlsx, 必填信息_已生成_2.xlsx ...
     - 测温报告_已生成_1.docx, 测温报告_已生成_2.docx ...
  2. step2 自动读取编号最大的 Excel 文件，无需手动指定

V3 已有功能（全部保留）：
  - Step1 时间标框可选（部分图片无时间信息可跳过）
  - 分步执行：--step1 / --step2 / 交互式菜单
  - 命令行参数 --suffix 仍可用于终端模式手动指定版本名

使用方法：
  python3 91号文报告自动生成-终端版V4.py              # 完整运行（step1 + step2）
  python3 91号文报告自动生成-终端版V4.py --step1       # 仅运行 step1
  python3 91号文报告自动生成-终端版V4.py --step2       # 仅运行 step2（自动读最新Excel）
  python3 91号文报告自动生成-终端版V4.py --cli         # 命令行核对模式

文件结构（与脚本同目录）：
  ./图谱/                      红外精确测温图片
  ./必填信息.xlsx               数据表（模板）
  ./91号文报告.docx              Word报告模板
  ./必填信息_已生成_N.xlsx       step1 输出（自动编号，step2 自动读最新）
  ./测温报告_已生成_N.docx     step2 输出（自动编号）

依赖：与 V1 相同（opencv-python numpy pytesseract openpyxl pillow python-docx）
"""

import os
import re
import sys
import json
import glob
import copy
import time
import platform
import subprocess
import datetime
import argparse

import cv2
import numpy as np
import pytesseract
from openpyxl import load_workbook

from docx import Document
from docx.shared import Cm, Pt
from docx.table import Table
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL

# ===================== 配置区域（PlanB 分路径结构） =====================

BASE_DIR = os.path.dirname(sys.executable) if getattr(sys, 'frozen', False) else os.path.dirname(os.path.abspath(__file__))

REPORT_DIR = os.path.abspath(os.path.join(BASE_DIR, "..", "2_report"))  # PlanB: 模板和输出在上级

IMAGE_DIR = os.path.join(REPORT_DIR, "图谱")
EXCEL_PATH = os.path.join(REPORT_DIR, "91号文必填信息.xlsx")
EXCEL_OUTPUT_BASE = os.path.join(REPORT_DIR, "必填信息_已生成")
# 自动读取第一个 sheet（不按名称，按位置）

DOCX_TEMPLATE_PATH = os.path.join(REPORT_DIR, "91号文测温报告.docx")
DOCX_OUTPUT_BASE = os.path.join(REPORT_DIR, "测温报告_已生成")

AR1_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_ar1.json")
DATETIME_ROI_CONFIG_FILE = os.path.join(BASE_DIR, "roi_config_datetime.json")

DEBUG_DIR = os.path.join(BASE_DIR, "debug_crops_91")
SAVE_DEBUG_CROPS = True

NOT_RECOGNIZED_TEXT = "无法识别"
NO_TIME_TEXT = "无时间信息"  # 图片上无时间时用此标识，区别于OCR识别失败

# 自动定位"最高温"的标记字符（OCR 可能把"最高温"拆成多个片段，按优先级找）
MARKER_KEYWORDS = ("温", "最高", "高")

# GUI核对开关（--cli 参数可强制使用命令行模式）
USE_GUI = True

# 插入图片尺寸
IMAGE_WIDTH_CM = 12
IMAGE_HEIGHT_CM = 9

# 新填入文字的字体（只针对脚本写入的值，模板原有文字不动）
CHINESE_FONT = "宋体"
WESTERN_FONT = "Times New Roman"
FONT_SIZE = Pt(8.5)

# =================================================================

# ===================== Tesseract 路径自适应 =====================
import sys as _sys
if getattr(_sys, 'frozen', False):
    _tesseract_dir = os.path.join(os.path.dirname(_sys.executable), 'tesseract')
else:
    _tesseract_dir = os.path.join(BASE_DIR, 'tesseract')
_tesseract_exe = os.path.join(_tesseract_dir, 'tesseract.exe')
if os.path.exists(_tesseract_exe):
    if hasattr(os, 'add_dll_directory'):
        os.add_dll_directory(_tesseract_dir)
    os.environ['PATH'] = _tesseract_dir + os.pathsep + os.environ.get('PATH', '')
    os.environ['TESSDATA_PREFIX'] = os.path.join(_tesseract_dir, 'tessdata')
    os.environ['TMP'] = _tesseract_dir
    os.environ['TEMP'] = _tesseract_dir
    pytesseract.pytesseract.tesseract_cmd = _tesseract_exe


def resolve_output_path(base_path, ext, suffix=None):
    """生成不冲突的输出路径（Excel / Word 通用）。

    规则：
      - 指定 suffix xxx → base_xxx.ext
      - 未指定 → base_1.ext, base_2.ext ... 自动递增（从 1 开始）
    """
    if suffix:
        return f"{base_path}_{suffix}.{ext}"

    counter = 1
    while True:
        path = f"{base_path}_{counter}.{ext}"
        if not os.path.exists(path):
            return path
        counter += 1


def find_latest_output(base_path, ext):
    """找到编号最大的已有输出文件路径，供 step2 自动读取最新 Excel。"""
    pattern = f"{base_path}_*.{ext}"
    files = sorted(glob.glob(pattern))
    if not files:
        return None
    # 按编号排序，返回最新的
    return sorted(files, key=lambda p: natural_sort_key(p))[-1]


# ===================== 工具函数 =====================

def natural_sort_key(s):
    return [int(t) if t.isdigit() else t for t in re.split(r'(\d+)', s)]


def collect_image_files(image_dir):
    patterns = ["*.jpg", "*.JPG", "*.jpeg", "*.JPEG", "*.png", "*.PNG"]
    files = []
    for p in patterns:
        files.extend(glob.glob(os.path.join(image_dir, p)))
    files = sorted(set(files), key=lambda p: natural_sort_key(os.path.basename(p)))
    return files


def open_image_for_preview(image_path):
    try:
        system = platform.system()
        if system == "Darwin":
            subprocess.run(["open", image_path], check=False)
        elif system == "Windows":
            os.startfile(image_path)
        else:
            subprocess.run(["xdg-open", image_path], check=False)
    except Exception as e:
        print(f"  （提示：自动打开图片失败：{e}，可手动查看：{image_path}）")


# ===================== ROI框选 =====================

def imread_unicode(path):
    """cv2.imread 的 Unicode 安全替代（Windows 中文路径兼容）。"""
    img_array = np.fromfile(path, dtype=np.uint8)
    return cv2.imdecode(img_array, cv2.IMREAD_COLOR)


def select_roi_interactively(sample_image_path, window_title):
    img = imread_unicode(sample_image_path)
    if img is None:
        raise FileNotFoundError(f"无法读取示例图片: {sample_image_path}")
    img_h, img_w = img.shape[:2]

    state = {"drawing": False, "x0": -1, "y0": -1, "x1": -1, "y1": -1, "done": False}

    def on_mouse(event, x, y, flags, param):
        if state["done"]:
            return
        if event == cv2.EVENT_LBUTTONDOWN:
            state["drawing"] = True
            state["x0"], state["y0"] = x, y
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_MOUSEMOVE and state["drawing"]:
            state["x1"], state["y1"] = x, y
        elif event == cv2.EVENT_LBUTTONUP:
            state["drawing"] = False
            state["x1"], state["y1"] = x, y

    cv2.namedWindow(window_title)
    cv2.setMouseCallback(window_title, on_mouse)
    display = img.copy()

    print(f"\n请在弹出的窗口中用鼠标左键拖拽框选【{window_title}】区域。")
    print("  空格/回车=确认   c=取消重选   Esc=退出\n")

    while True:
        display = img.copy()
        x0, y0 = state["x0"], state["y0"]
        x1, y1 = state["x1"], state["y1"]
        if x0 >= 0 and y0 >= 0:
            cv2.rectangle(display, (x0, y0), (x1, y1), (0, 255, 0), 1)
        cv2.imshow(window_title, display)
        key = cv2.waitKey(20) & 0xFF
        if key == 13 or key == 32:
            break
        elif key == ord('c') or key == ord('C') or key == 27:
            state["x0"] = state["y0"] = state["x1"] = state["y1"] = -1
            state["drawing"] = False
            if key == 27:
                break

    cv2.destroyAllWindows()
    cv2.waitKey(1)
    time.sleep(0.3)

    x0, y0 = state["x0"], state["y0"]
    x1, y1 = state["x1"], state["y1"]
    if x0 < 0 or y0 < 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    x = min(x0, x1)
    y = min(y0, y1)
    w = abs(x1 - x0)
    h = abs(y1 - y0)
    if w == 0 or h == 0:
        raise ValueError("未框选有效区域，程序退出，请重新运行脚本。")

    rx, ry, rw, rh = x / img_w, y / img_h, w / img_w, h / img_h
    print(f"框选区域（像素）: x={x}, y={y}, w={w}, h={h}  (示例图片尺寸: {img_w}x{img_h})")
    return rx, ry, rw, rh


def match_images_by_number(image_files, data_rows_info):
    numbered = {}
    all_numbered = True
    for p in image_files:
        base = os.path.splitext(os.path.basename(p))[0]
        m = re.search(r'(\d+)', base)
        if m:
            numbered[int(m.group(1))] = p
        else:
            all_numbered = False

    row_to_image = {}
    if all_numbered and numbered:
        print("图片匹配方式：按文件名中的数字精确对应第几行（例如'图5.jpg'对应第5行数据）。")
        for i, info in enumerate(data_rows_info, start=1):
            row_to_image[info['row']] = numbered.get(i)
    else:
        print("图片匹配方式：文件名不含数字规律，改为按文件夹内文件名顺序直接依次对应。")
        for i, info in enumerate(data_rows_info):
            row_to_image[info['row']] = image_files[i] if i < len(image_files) else None

    missing = [info['row'] for info in data_rows_info if row_to_image.get(info['row']) is None]
    if missing:
        print(f"提示：以下Excel行号没有匹配到图片，需要人工确认：{missing}")

    return row_to_image


def get_roi(sample_image_path, config_file, prompt_text, window_title):
    if os.path.exists(config_file):
        with open(config_file, "r", encoding="utf-8") as f:
            cfg = json.load(f)
        rx, ry, rw, rh = cfg["rx"], cfg["ry"], cfg["rw"], cfg["rh"]
        print(f"检测到已保存的【{prompt_text}】区域配置: rx={rx:.4f}, ry={ry:.4f}, rw={rw:.4f}, rh={rh:.4f}")
        choice = input(f"是否复用该区域？(y=复用 / n=重新框选) [y]: ").strip().lower()
        if choice in ("", "y", "yes"):
            return rx, ry, rw, rh

    print(f"\n请在弹出的窗口中框选【{prompt_text}】区域，框选完成后按【空格/回车】确认，按 c 取消重选。")
    rx, ry, rw, rh = select_roi_interactively(sample_image_path, window_title)
    with open(config_file, "w", encoding="utf-8") as f:
        json.dump({"rx": rx, "ry": ry, "rw": rw, "rh": rh}, f, ensure_ascii=False, indent=2)
    print(f"识别区域已保存到: {config_file}")
    return rx, ry, rw, rh


def crop_by_relative_roi(img, roi_ratio):
    rx, ry, rw, rh = roi_ratio
    img_h, img_w = img.shape[:2]
    x = int(round(rx * img_w))
    y = int(round(ry * img_h))
    w = int(round(rw * img_w))
    h = int(round(rh * img_h))
    x2, y2 = min(x + w, img_w), min(y + h, img_h)
    x, y = max(x, 0), max(y, 0)
    return img[y:y2, x:x2]


def preprocess_variants(crop, scale=4):
    crop_big = cv2.resize(crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)
    gray = cv2.cvtColor(crop_big, cv2.COLOR_BGR2GRAY)

    variants = []
    blur = cv2.GaussianBlur(gray, (3, 3), 0)
    _, thresh1 = cv2.threshold(blur, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh1)

    inv = cv2.bitwise_not(blur)
    _, thresh2 = cv2.threshold(inv, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(thresh2)

    thresh3 = cv2.adaptiveThreshold(
        gray, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10
    )
    variants.append(thresh3)

    return variants


# ===================== OCR：Ar1温度 =====================

def extract_number(text):
    if not text:
        return None
    match = re.search(r'-?\d+\.?\d*', text.strip())
    if not match:
        return None
    try:
        num = float(match.group())
        return str(int(num)) if num == int(num) else str(num)
    except ValueError:
        return None


def ocr_ar1_temperature(image_path, roi_ratio):
    img = imread_unicode(image_path)
    if img is None:
        return None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, ["[裁剪区域为空]"]

    variants = preprocess_variants(crop)
    configs = [
        r'--oem 3 --psm 7 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 6 -c tessedit_char_whitelist=0123456789.-',
        r'--oem 3 --psm 11 -c tessedit_char_whitelist=0123456789.-',
    ]
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            value = extract_number(text)
            if value is not None:
                return value, crop, raw_texts
    return None, crop, raw_texts


# ===================== OCR：拍摄日期 / 拍摄时间 =====================

DATE_PATTERN = re.compile(r'(\d{4})\D(\d{1,2})\D(\d{1,2})')
TIME_PATTERN = re.compile(r'(\d{1,2}):(\d{2}):(\d{2})')


def parse_date(text):
    if not text:
        return None
    dm = DATE_PATTERN.search(text)
    if not dm:
        return None
    try:
        y, mo, d = int(dm.group(1)), int(dm.group(2)), int(dm.group(3))
        return datetime.date(y, mo, d)
    except ValueError:
        return None


def parse_time(text):
    if not text:
        return None
    tm = TIME_PATTERN.search(text)
    if not tm:
        return None
    h, mi, s = tm.group(1), tm.group(2), tm.group(3)
    try:
        if 0 <= int(h) <= 23 and 0 <= int(mi) <= 59 and 0 <= int(s) <= 59:
            return f"{int(h):02d}:{mi}:{s}"
    except ValueError:
        pass
    return None


def _ocr_with_configs(crop, configs):
    variants = preprocess_variants(crop)
    raw_texts = []
    for variant in variants:
        for cfg in configs:
            text = pytesseract.image_to_string(variant, config=cfg)
            raw_texts.append(text.strip())
            yield text, raw_texts


DATE_TIME_CHAR_WHITELIST = r'-c tessedit_char_whitelist=0123456789.:-年月日 '
_OCR_CONFIGS_DATE_TIME = [
    rf'--oem 3 --psm 7 {DATE_TIME_CHAR_WHITELIST}',
    rf'--oem 3 --psm 6 {DATE_TIME_CHAR_WHITELIST}',
]


def ocr_datetime(image_path, roi_ratio):
    img = imread_unicode(image_path)
    if img is None:
        return None, None, None, ["[无法读取图片]"]
    crop = crop_by_relative_roi(img, roi_ratio)
    if crop.size == 0:
        return None, None, crop, ["[裁剪区域为空]"]

    raw_texts = []
    for text, raw_texts in _ocr_with_configs(crop, _OCR_CONFIGS_DATE_TIME):
        date_val = parse_date(text)
        time_val = parse_time(text)
        if date_val is not None or time_val is not None:
            return date_val, time_val, crop, raw_texts
    return None, None, crop, raw_texts


# ===================== Excel读取 =====================

def find_header_row(ws, keyword="设备类型", max_scan_rows=5):
    for row in range(1, max_scan_rows + 1):
        for col in range(1, 5):
            if str(ws.cell(row=row, column=col).value).strip() == keyword:
                return row
    return 1


def read_excel_data_rows(ws, header_row_idx):
    data_rows_info = []
    row = header_row_idx + 1
    while True:
        d_val = ws.cell(row=row, column=4).value
        if d_val is None or str(d_val).strip() == "":
            break
        info = {
            'row': row,
            'device_type': ws.cell(row=row, column=2).value,
            'interval': ws.cell(row=row, column=3).value,
            'device_name': d_val,
            'phase': ws.cell(row=row, column=5).value,
            'distance': ws.cell(row=row, column=6).value,
            'voltage': ws.cell(row=row, column=7).value,
            'current': ws.cell(row=row, column=8).value,
            'power': ws.cell(row=row, column=9).value,
        }
        data_rows_info.append(info)
        row += 1
    return data_rows_info


# ===================== 从已生成的Excel读回数据（step2 复用） =====================

def read_excel_output(excel_output_path=None):
    """从 step1 生成的最新 Excel 中读回完整数据（含 OCR 结果），供 step2 使用。

    如果不指定路径，自动找到编号最大的 必填信息_已生成_*.xlsx。
    返回格式与 step1_recognize_and_review() 一致，可直接传给 step2_generate_docx()。
    """
    if excel_output_path is None:
        excel_output_path = find_latest_output(EXCEL_OUTPUT_BASE, "xlsx")
    if not os.path.isfile(excel_output_path):
        print(f"错误：找不到 step1 生成的 Excel 文件: {excel_output_path}")
        print("  请先运行 --step1 完成识别与核对。")
        sys.exit(1)

    print(f"正在从 {excel_output_path} 读取已核对的识别结果...")
    wb = load_workbook(excel_output_path)
    sheet_name = wb.sheetnames[0]
    ws = wb[sheet_name]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行。")
        sys.exit(1)

    image_files = collect_image_files(IMAGE_DIR)
    row_to_image = match_images_by_number(image_files, data_rows_info)

    full_data = []
    for info in data_rows_info:
        r = info['row']
        # 从 J/K/L 列读取 step1 的识别结果
        j_val = ws.cell(row=r, column=10).value  # Ar1温度
        k_val = ws.cell(row=r, column=11).value  # 拍摄日期
        l_val = ws.cell(row=r, column=12).value  # 拍摄时间

        # 解析日期（Excel中可能是date对象、字符串或NOT_RECOGNIZED_TEXT）
        date_val = k_val
        if isinstance(k_val, datetime.datetime):
            date_val = k_val.date()
        elif isinstance(k_val, str) and k_val != NOT_RECOGNIZED_TEXT:
            parsed = parse_date(k_val)
            if parsed:
                date_val = parsed

        # 解析温度
        ar1_val = j_val
        if isinstance(j_val, (int, float)):
            ar1_val = str(int(j_val)) if j_val == int(j_val) else str(j_val)
        elif isinstance(j_val, str) and j_val == NOT_RECOGNIZED_TEXT:
            ar1_val = NOT_RECOGNIZED_TEXT

        # 解析时间
        time_val = l_val
        if isinstance(l_val, datetime.time):
            time_val = l_val.strftime("%H:%M:%S")

        row_data = dict(info)
        row_data['ar1'] = ar1_val
        row_data['date'] = date_val
        row_data['time'] = time_val
        row_data['image'] = row_to_image.get(r)
        full_data.append(row_data)

    print(f"  共读取 {len(full_data)} 行设备数据。")
    return full_data


# ===================== 人工核对GUI（两个独立窗口） =====================

def _load_image_into(label, path, ref_dict, ref_key, max_w=580, max_h=320):
    try:
        from PIL import Image, ImageTk
    except ImportError:
        return
    if path and os.path.exists(path):
        pil_img = Image.open(path)
        ratio = min(max_w / pil_img.width, max_h / pil_img.height, 1.0)
        new_size = (max(1, int(pil_img.width * ratio)), max(1, int(pil_img.height * ratio)))
        pil_img = pil_img.resize(new_size)
        photo = ImageTk.PhotoImage(pil_img)
        ref_dict[ref_key] = photo
        label.config(image=photo, text="")
    else:
        ref_dict[ref_key] = None
        label.config(image="", text="（未找到对应裁剪图片）")


# --- 核对窗口 1：Ar1温度 ---

def interactive_review_ar1_gui(data_rows_info, ar1_results, ar1_debug_map):
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（GUI不可用({e})，回退命令行。）")
        return interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map)

    reviewed = dict(ar1_results)
    total = len(data_rows_info)
    state = {'index': 0, 'confirmed': False}
    photo_refs = {}

    root = tk.Tk()
    root.title("91号文测温 - 核对 Ar1温度")
    root.geometry("620x520")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 14), justify="left", wraplength=580)
    info_label.pack(pady=(15, 5))

    img_label = tk.Label(root)
    img_label.pack(pady=8)

    entry_var = tk.StringVar()
    entry = tk.Entry(root, textvariable=entry_var, font=("PingFang SC", 22), justify="center", width=14)
    entry.pack(pady=8)

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(0, 5))

    btn_ref = {}

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']
        current = reviewed.get(r)
        status = current if current is not None else ""
        mark = "   ⚠️ 识别失败，需手动输入" if current is None else ""
        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}{mark}")
        _load_image_into(img_label, ar1_debug_map.get(r), photo_refs, 'ar1')
        entry_var.set(status)
        hint_label.config(text="↑ = 上一条    |    ↓ = 下一条    |    回车 = 确认提交", fg="gray")
        entry.focus_set()
        entry.select_range(0, tk.END)
        update_buttons()

    def update_buttons():
        idx = state['index']
        btn_ref['prev'].config(state=(tk.DISABLED if idx == 0 else tk.NORMAL))
        btn_ref['next'].config(state=(tk.DISABLED if idx + 1 >= total else tk.NORMAL))

    def try_save():
        idx = state['index']
        r = data_rows_info[idx]['row']
        val = entry_var.get().strip()
        if val != "":
            try:
                float(val)
                reviewed[r] = val
            except ValueError:
                hint_label.config(text="⚠️ 不是有效数字，请重新输入！", fg="red")
                return False
        return True

    def go_previous(event=None):
        if not try_save():
            return
        if state['index'] == 0:
            return
        state['index'] -= 1
        load_row(state['index'])

    def go_next(event=None):
        if not try_save():
            return
        if state['index'] + 1 >= total:
            hint_label.config(text="已是最后一条，请按回车提交", fg="orange")
            return
        state['index'] += 1
        load_row(state['index'])

    def confirm_submit(event=None):
        if not try_save():
            return
        state['confirmed'] = True
        root.destroy()

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=15)
    btn_ref['prev'] = tk.Button(button_frame, text="↑ 上一条", command=go_previous, width=10)
    btn_ref['prev'].grid(row=0, column=0, padx=4)
    btn_ref['next'] = tk.Button(button_frame, text="↓ 下一条", command=go_next, width=10)
    btn_ref['next'].grid(row=0, column=1, padx=4)
    btn_ref['submit'] = tk.Button(button_frame, text="确认提交（回车）", command=confirm_submit, width=16, fg="blue")
    btn_ref['submit'].grid(row=0, column=2, padx=4)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=3, padx=4)

    root.bind('<Up>', go_previous)
    root.bind('<Down>', go_next)
    root.bind('<Return>', confirm_submit)

    print("\n已打开【Ar1温度】核对窗口...")
    load_row(0)
    root.lift()
    root.focus_force()
    entry.focus_set()
    root.mainloop()

    if state['confirmed']:
        print("【Ar1温度】核对完成，用户已确认提交。")
    else:
        print("【Ar1温度】核对已跳过。")
    return reviewed


def interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map):
    reviewed = dict(ar1_results)
    total = len(data_rows_info)
    print(f"\n{'=' * 50}\n进入【Ar1温度】核对（命令行）\n{'=' * 50}")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        current = reviewed.get(r)
        print(f"\n[{i}/{total}] {info['device_name']}  当前值={current}")
        if ar1_debug_map.get(r):
            open_image_for_preview(ar1_debug_map[r])
        while True:
            val = input("  温度 (回车=保留): ").strip()
            if val == "" or _is_float(val):
                if val != "":
                    reviewed[r] = val
                break
            print("  不是有效数字。")
        if input("  s=跳过剩余: ").strip().lower() in ("s", "skip"):
            break
    return reviewed


# --- 核对窗口 2：拍摄日期+时间 ---

def interactive_review_datetime_gui(data_rows_info, date_results, time_results, dt_debug_map):
    try:
        import tkinter as tk
        from PIL import Image, ImageTk
    except ImportError as e:
        print(f"\n（GUI不可用({e})，回退命令行。）")
        return interactive_review_datetime_cli(data_rows_info, date_results, time_results, dt_debug_map)

    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)
    state = {'index': 0, 'confirmed': False}
    photo_refs = {}

    root = tk.Tk()
    root.title("91号文测温 - 核对 拍摄日期+时间")
    root.geometry("620x580")
    root.attributes('-topmost', True)

    info_label = tk.Label(root, text="", font=("PingFang SC", 14), justify="left", wraplength=580)
    info_label.pack(pady=(15, 5))

    img_label = tk.Label(root)
    img_label.pack(pady=6)

    tk.Label(root, text="拍摄日期 (YYYY-MM-DD)", font=("PingFang SC", 11)).pack(pady=(8, 0))
    date_var = tk.StringVar()
    date_entry = tk.Entry(root, textvariable=date_var, font=("PingFang SC", 18), justify="center", width=18)
    date_entry.pack(pady=3)

    tk.Label(root, text="拍摄时间 (HH:MM:SS)", font=("PingFang SC", 11)).pack(pady=(6, 0))
    time_var = tk.StringVar()
    time_entry = tk.Entry(root, textvariable=time_var, font=("PingFang SC", 18), justify="center", width=18)
    time_entry.pack(pady=3)

    hint_label = tk.Label(root, text="", font=("PingFang SC", 11), fg="gray")
    hint_label.pack(pady=(6, 4))

    btn_ref = {}

    def load_row(idx):
        info = data_rows_info[idx]
        r = info['row']
        dc = date_reviewed.get(r)
        tc = time_reviewed.get(r)
        info_label.config(text=f"[{idx + 1}/{total}]  设备: {info['device_name']}")
        _load_image_into(img_label, dt_debug_map.get(r), photo_refs, 'dt')
        date_var.set(dc.isoformat() if isinstance(dc, datetime.date) else (dc or ""))
        time_var.set(tc if tc else "")
        hint_label.config(text="↑ = 上一条    |    ↓ = 下一条    |    回车 = 确认提交", fg="gray")
        date_entry.focus_set()
        date_entry.select_range(0, tk.END)
        update_buttons()

    def update_buttons():
        idx = state['index']
        btn_ref['prev'].config(state=(tk.DISABLED if idx == 0 else tk.NORMAL))
        btn_ref['next'].config(state=(tk.DISABLED if idx + 1 >= total else tk.NORMAL))

    def try_save():
        idx = state['index']
        r = data_rows_info[idx]['row']
        ds = date_var.get().strip()
        ts = time_var.get().strip()
        if ds:
            try:
                date_reviewed[r] = datetime.date.fromisoformat(ds)
            except ValueError:
                hint_label.config(text="⚠️ 日期格式不对，请用 YYYY-MM-DD！", fg="red")
                return False
        if ts:
            if not re.match(r'^\d{1,2}:\d{2}:\d{2}$', ts):
                hint_label.config(text="⚠️ 时间格式不对，请用 HH:MM:SS！", fg="red")
                return False
            time_reviewed[r] = ts
        return True

    def go_previous(event=None):
        if not try_save():
            return
        if state['index'] == 0:
            return
        state['index'] -= 1
        load_row(state['index'])

    def go_next(event=None):
        if not try_save():
            return
        if state['index'] + 1 >= total:
            hint_label.config(text="已是最后一条，请按回车提交", fg="orange")
            return
        state['index'] += 1
        load_row(state['index'])

    def confirm_submit(event=None):
        if not try_save():
            return
        state['confirmed'] = True
        root.destroy()

    def skip_rest():
        root.destroy()

    button_frame = tk.Frame(root)
    button_frame.pack(pady=12)
    btn_ref['prev'] = tk.Button(button_frame, text="↑ 上一条", command=go_previous, width=10)
    btn_ref['prev'].grid(row=0, column=0, padx=4)
    btn_ref['next'] = tk.Button(button_frame, text="↓ 下一条", command=go_next, width=10)
    btn_ref['next'].grid(row=0, column=1, padx=4)
    btn_ref['submit'] = tk.Button(button_frame, text="确认提交（回车）", command=confirm_submit, width=16, fg="blue")
    btn_ref['submit'].grid(row=0, column=2, padx=4)
    tk.Button(button_frame, text="跳过剩余核对", command=skip_rest, width=14).grid(row=0, column=3, padx=4)

    root.bind('<Up>', go_previous)
    root.bind('<Down>', go_next)
    root.bind('<Return>', confirm_submit)

    print("\n已打开【拍摄日期+时间】核对窗口...")
    load_row(0)
    root.lift()
    root.focus_force()
    date_entry.focus_set()
    root.mainloop()

    if state['confirmed']:
        print("【拍摄日期+时间】核对完成，用户已确认提交。")
    else:
        print("【拍摄日期+时间】核对已跳过。")
    return date_reviewed, time_reviewed


def interactive_review_datetime_cli(data_rows_info, date_results, time_results, dt_debug_map):
    date_reviewed = dict(date_results)
    time_reviewed = dict(time_results)
    total = len(data_rows_info)
    print(f"\n{'=' * 50}\n进入【拍摄日期+时间】核对（命令行）\n{'=' * 50}")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        dc = date_reviewed.get(r)
        tc = time_reviewed.get(r)
        print(f"\n[{i}/{total}] {info['device_name']}  日期={dc}  时间={tc}")
        if dt_debug_map.get(r):
            open_image_for_preview(dt_debug_map[r])
        while True:
            val = input("  日期 YYYY-MM-DD (回车=保留): ").strip()
            if val == "":
                break
            try:
                date_reviewed[r] = datetime.date.fromisoformat(val)
                break
            except ValueError:
                print("  格式不对。")
        while True:
            val = input("  时间 HH:MM:SS (回车=保留): ").strip()
            if val == "":
                break
            if re.match(r'^\d{1,2}:\d{2}:\d{2}$', val):
                time_reviewed[r] = val
                break
            print("  格式不对。")
        if input("  s=跳过剩余: ").strip().lower() in ("s", "skip"):
            break
    return date_reviewed, time_reviewed


def _is_float(s):
    try:
        float(s)
        return True
    except ValueError:
        return False


# ===================== Step 1：识别 + 核对 + 写入Excel =====================

# ===================== 自动定位（温度框不固定） =====================

def find_wen_marker(img):
    """在整图中定位"最高温"的"温"字，返回 (x, y, w, h) 或 None。

    先做白色文字分割（热像图上文字是白色），把白字变黑、其余变白后再 OCR，
    能大幅提高"最高温"的识别率——直接灰度 OCR 常把"最高温"读乱（如读成"Sia:"）。
    """
    b, g, r = img[:, :, 0], img[:, :, 1], img[:, :, 2]
    white = (b > 180) & (g > 180) & (r > 180)
    mask = np.where(white, 0, 255).astype(np.uint8)
    data = pytesseract.image_to_data(
        mask, lang="chi_sim+eng", config="--psm 11", output_type=pytesseract.Output.DICT
    )
    for keyword in MARKER_KEYWORDS:
        for i, t in enumerate(data["text"]):
            t = t.strip()
            if not t:
                continue
            if keyword in t:
                return (data["left"][i], data["top"][i], data["width"][i], data["height"][i])
    return None


def extract_number_from_text(text):
    """从OCR文本提取数值，优先带小数点（温度格式 XX.X，取 1 位小数）。"""
    text = text.strip().replace("，", ".").replace(":", "").replace("：", "")
    m = re.search(r"\d+\.\d", text)
    if m:
        return m.group(0)
    m = re.search(r"\d+", text)
    if m:
        return m.group(0)
    return None


def ocr_number(crop):
    """从数值小图（BGR）识别数字。优先白色文字分割，回退灰度多阈值。"""
    b, g, r = crop[:, :, 0], crop[:, :, 1], crop[:, :, 2]
    white = (b > 180) & (g > 180) & (r > 180)
    if white.sum() > 0:
        mask = np.where(white, 0, 255).astype(np.uint8)
        big = cv2.resize(mask, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
        _, th = cv2.threshold(big, 127, 255, cv2.THRESH_BINARY)
        txt = pytesseract.image_to_string(th, config="--psm 7 -c tessedit_char_whitelist=0123456789.")
        val = extract_number_from_text(txt)
        if val is not None:
            return val

    candidates = []
    big = cv2.resize(crop, None, fx=4, fy=4, interpolation=cv2.INTER_CUBIC)
    g = cv2.cvtColor(big, cv2.COLOR_BGR2GRAY)
    variants = []
    _, th1 = cv2.threshold(g, 0, 255, cv2.THRESH_BINARY + cv2.THRESH_OTSU)
    variants.append(th1)
    variants.append(cv2.bitwise_not(th1))
    variants.append(cv2.adaptiveThreshold(g, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, cv2.THRESH_BINARY, 31, 10))

    for v in variants:
        txt = pytesseract.image_to_string(v, config="--psm 7 -c tessedit_char_whitelist=0123456789.")
        val = extract_number_from_text(txt)
        if val is not None:
            if "." in val:
                return val
            candidates.append(val)

    return candidates[0] if candidates else None


def find_number_box(crop):
    """在数值区域内定位数字的精确边界（白色连通域），去掉冒号、空白和右侧的"°C"等。"""
    b, g, r = crop[:, :, 0], crop[:, :, 1], crop[:, :, 2]
    white = ((b > 180) & (g > 180) & (r > 180)).astype(np.uint8)
    if white.sum() == 0:
        return None
    n, labels, stats, _ = cv2.connectedComponentsWithStats(white, connectivity=8)

    comps = []
    for i in range(1, n):
        if stats[i, cv2.CC_STAT_AREA] < 15:
            continue
        l = stats[i, cv2.CC_STAT_LEFT]
        t = stats[i, cv2.CC_STAT_TOP]
        r = l + stats[i, cv2.CC_STAT_WIDTH]
        btm = t + stats[i, cv2.CC_STAT_HEIGHT]
        comps.append((l, t, r, btm))
    if not comps:
        return None

    # 按 x 排序，从左往右扩展，遇到大间隙（>20px）停止
    comps.sort(key=lambda c: c[0])
    x0, y0, x1, y1 = comps[0]
    for (l, t, r, btm) in comps[1:]:
        if l - x1 > 20:
            break
        x0 = min(x0, l)
        y0 = min(y0, t)
        x1 = max(x1, r)
        y1 = max(y1, btm)

    m = 3
    x0 = max(0, x0 - m)
    y0 = max(0, y0 - m)
    x1 = min(crop.shape[1], x1 + m)
    y1 = min(crop.shape[0], y1 + m)
    return (x0, y0, x1, y1)


def find_label_left(img, wen_box):
    """找到"温"字左侧标签（RO1/最高温）的最左边界，用于显示完整标签。"""
    wx, wy, ww, wh = wen_box
    b, g, r = img[:, :, 0], img[:, :, 1], img[:, :, 2]
    white = (b > 180) & (g > 180) & (r > 180)
    mask = np.where(white, 0, 255).astype(np.uint8)
    data = pytesseract.image_to_data(
        mask, lang="chi_sim+eng", config="--psm 11", output_type=pytesseract.Output.DICT
    )
    label_left = wx
    for i, t in enumerate(data["text"]):
        t = t.strip()
        if not t:
            continue
        x = data["left"][i]
        y = data["top"][i]
        hh = data["height"][i]
        if x < wx and wx - x < 250 and abs((y + hh / 2) - (wy + wh / 2)) < wh * 2:
            label_left = min(label_left, x)
    return label_left


def extract_max_temp(image_path):
    """自动定位并提取最高温数值，返回 (数值字符串或None, 显示用裁剪图或None)。

    识别用紧致数字框（只框"XX.X"）；显示用完整标签框（"RO1 最高温：XX.X"）。
    """
    img = imread_unicode(image_path)
    if img is None:
        return None, None

    marker = find_wen_marker(img)
    if marker is None:
        return None, None

    x, y, w, h = marker

    # --- 识别：裁剪"温"右侧数值区，收紧到只含数字 ---
    x0 = max(0, x + w)
    x1 = min(img.shape[1], x + w + 150)
    y0 = max(0, y - 10)
    y1 = min(img.shape[0], y + h + 10)
    crop = img[y0:y1, x0:x1]
    if crop.size == 0:
        return None, None

    box = find_number_box(crop)
    if box is not None:
        bx0, by0, bx1, by1 = box
        crop = crop[by0:by1, bx0:bx1]
        val_right = x0 + bx1
    else:
        val_right = x1

    if crop.size == 0:
        return None, None

    value = ocr_number(crop)

    # --- 显示：裁剪完整标签"RO1 最高温：XX.X" ---
    label_left = find_label_left(img, marker)
    dx0 = max(0, label_left - 3)
    dx1 = min(img.shape[1], val_right + 3)
    dy0 = max(0, y - 12)
    dy1 = min(img.shape[0], y + h + 12)
    display_crop = img[dy0:dy1, dx0:dx1]
    scale = max(3.0, 160.0 / max(1, display_crop.shape[0]))
    display = cv2.resize(display_crop, None, fx=scale, fy=scale, interpolation=cv2.INTER_CUBIC)

    return value, display


def step1_recognize_and_review(excel_output_path):
    print(f"\n{'=' * 50}")
    print("  Step 1: OCR识别 + 人工核对 + 写入Excel")
    print(f"{'=' * 50}\n")

    wb = load_workbook(EXCEL_PATH)
    sheet_name = wb.sheetnames[0]
    print(f"正在读取Excel模板: {EXCEL_PATH}  (工作表: {sheet_name})")
    ws = wb[sheet_name]

    header_row_idx = find_header_row(ws)
    data_rows_info = read_excel_data_rows(ws, header_row_idx)
    if not data_rows_info:
        print("错误：Excel中没有读取到任何数据行，请检查表格内容。")
        sys.exit(1)
    print(f"共读取到 {len(data_rows_info)} 行设备数据。")

    image_files = collect_image_files(IMAGE_DIR)
    if len(image_files) < len(data_rows_info):
        print(f"警告：图谱文件夹只有 {len(image_files)} 张图片，少于Excel的 {len(data_rows_info)} 行数据，"
              f"部分行可能无法识别。")

    row_to_image = match_images_by_number(image_files, data_rows_info)

    sample_path = next((p for p in row_to_image.values() if p), None)
    if sample_path is None:
        print("错误：没有可用的示例图片用于框选区域。")
        sys.exit(1)

    # 询问温度框是否不固定：y=自动定位，n=固定框手动框选
    choice = input("是否温度框不固定？(y=不固定，自动定位 / n=固定，手动框选) [n]: ").strip().lower()
    use_auto = choice in ("y", "yes")

    ar1_roi = None
    if not use_auto:
        ar1_roi = get_roi(sample_path, AR1_ROI_CONFIG_FILE, "温度（Ar1）数值",
                           "框选温度(Ar1)区域 (Space/Enter确认, C取消)")

    # ===== V3：时间标框改为可选 =====
    print()
    print("=" * 50)
    print("  部分测温图片上没有拍摄日期和时间。")
    choice = input("  是否需要标框识别拍摄日期+时间？(y=需要 / n=跳过，图片上无时间) [y]: ").strip().lower()
    has_datetime = choice in ("", "y", "yes")

    datetime_roi = None
    if has_datetime:
        datetime_roi = get_roi(sample_path, DATETIME_ROI_CONFIG_FILE, "拍摄日期+时间",
                                "框选拍摄日期+时间区域 (Space/Enter确认, C取消)")
    else:
        print("已跳过日期时间标框，日期和时间列将填入「无时间信息」。")

    if SAVE_DEBUG_CROPS:
        os.makedirs(DEBUG_DIR, exist_ok=True)

    ar1_results = {}
    date_results = {}
    time_results = {}
    ar1_debug_map = {}
    dt_debug_map = {}

    print("\n开始批量识别，请稍候...\n")
    for i, info in enumerate(data_rows_info, start=1):
        r = info['row']
        img_path = row_to_image.get(r)
        if img_path is None:
            ar1_results[r] = None
            date_results[r] = NO_TIME_TEXT if not has_datetime else None
            time_results[r] = NO_TIME_TEXT if not has_datetime else None
            ar1_debug_map[r] = None
            dt_debug_map[r] = None
            print(f"[{i}/{len(data_rows_info)}] 设备: {info['device_name']}  ->  无对应图片")
            continue

        if use_auto:
            ar1_val, ar1_display = extract_max_temp(img_path)
            ar1_crop = ar1_display
            ar1_raw = []
        else:
            ar1_val, ar1_crop, ar1_raw = ocr_ar1_temperature(img_path, ar1_roi)

        if has_datetime:
            date_val, time_val, dt_crop, dt_raw = ocr_datetime(img_path, datetime_roi)
            date_results[r] = date_val
            time_results[r] = time_val
        else:
            date_val, time_val, dt_crop = NO_TIME_TEXT, NO_TIME_TEXT, None
            date_results[r] = NO_TIME_TEXT
            time_results[r] = NO_TIME_TEXT

        ar1_results[r] = ar1_val

        if has_datetime:
            status = f"温度={ar1_val if ar1_val is not None else NOT_RECOGNIZED_TEXT}  " \
                     f"日期={date_val if date_val else NOT_RECOGNIZED_TEXT}  " \
                     f"时间={time_val if time_val else NOT_RECOGNIZED_TEXT}"
        else:
            status = f"温度={ar1_val if ar1_val is not None else NOT_RECOGNIZED_TEXT}  " \
                     f"日期/时间=已跳过"

        print(f"[{i}/{len(data_rows_info)}] {os.path.basename(img_path)}  ->  {status}")

        ar1_debug_map[r] = None
        dt_debug_map[r] = None
        if SAVE_DEBUG_CROPS:
            if ar1_crop is not None:
                name = re.sub(r'[\\/:*?"<>|]', "_", f"row{r}_ar1.png")
                save_img = cv2.resize(ar1_crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                path = os.path.join(DEBUG_DIR, name)
                cv2.imwrite(path, save_img)
                ar1_debug_map[r] = path
            if dt_crop is not None:
                name = re.sub(r'[\\/:*?"<>|]', "_", f"row{r}_datetime.png")
                save_img = cv2.resize(dt_crop, None, fx=2, fy=2, interpolation=cv2.INTER_CUBIC)
                path = os.path.join(DEBUG_DIR, name)
                cv2.imwrite(path, save_img)
                dt_debug_map[r] = path

    if has_datetime:
        fail_count = sum(
            1 for info in data_rows_info
            if ar1_results.get(info['row']) is None
            or date_results.get(info['row']) is None
            or time_results.get(info['row']) is None
        )
    else:
        fail_count = sum(
            1 for info in data_rows_info
            if ar1_results.get(info['row']) is None
        )
    do_review = True
    if fail_count == 0:
        print(f"\n本次 {len(data_rows_info)} 行全部识别成功。")
        choice = input("是否仍要逐张人工核对确认？(y=核对 / n=跳过，直接使用识别结果) [n]: ").strip().lower()
        do_review = choice in ("y", "yes")
    else:
        print(f"\n共有 {fail_count} / {len(data_rows_info)} 行存在识别失败的内容，需要人工核对补充。")

    if do_review:
        if USE_GUI:
            ar1_results = interactive_review_ar1_gui(data_rows_info, ar1_results, ar1_debug_map)
            if has_datetime:
                date_results, time_results = interactive_review_datetime_gui(
                    data_rows_info, date_results, time_results, dt_debug_map
                )
        else:
            print("\n（已启用 --cli 模式，使用命令行核对）")
            ar1_results = interactive_review_ar1_cli(data_rows_info, ar1_results, ar1_debug_map)
            if has_datetime:
                date_results, time_results = interactive_review_datetime_cli(
                    data_rows_info, date_results, time_results, dt_debug_map
                )

    # 写回Excel
    for info in data_rows_info:
        r = info['row']
        ar1_val = ar1_results.get(r)
        date_val = date_results.get(r)
        time_val = time_results.get(r)

        ws.cell(row=r, column=10).value = float(ar1_val) if ar1_val is not None else NOT_RECOGNIZED_TEXT
        if has_datetime:
            ws.cell(row=r, column=11).value = date_val if date_val is not None else NOT_RECOGNIZED_TEXT
            ws.cell(row=r, column=12).value = time_val if time_val is not None else NOT_RECOGNIZED_TEXT
        else:
            ws.cell(row=r, column=11).value = NO_TIME_TEXT
            ws.cell(row=r, column=12).value = NO_TIME_TEXT

    wb.save(excel_output_path)
    print(f"\nExcel已生成: {excel_output_path}")

    full_data = []
    for info in data_rows_info:
        r = info['row']
        row_data = dict(info)
        row_data['ar1'] = ar1_results.get(r)
        row_data['date'] = date_results.get(r)
        row_data['time'] = time_results.get(r)
        row_data['image'] = row_to_image.get(r)
        full_data.append(row_data)

    return full_data


# ===================== Step 2：生成 Word报告 =====================

def apply_font_to_run(run):
    run.font.name = WESTERN_FONT
    run.font.size = FONT_SIZE
    rPr = run._element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is None:
        rFonts = OxmlElement('w:rFonts')
        rPr.insert(0, rFonts)
    rFonts.set(qn('w:eastAsia'), CHINESE_FONT)
    rFonts.set(qn('w:ascii'), WESTERN_FONT)
    rFonts.set(qn('w:hAnsi'), WESTERN_FONT)


def set_cell_text(cell, text):
    text = "" if text is None else str(text)
    p = cell.paragraphs[0]
    for run in list(p.runs):
        run.text = ""
    if p.runs:
        run = p.runs[0]
        run.text = text
    else:
        run = p.add_run(text)
    apply_font_to_run(run)
    for extra_p in cell.paragraphs[1:]:
        extra_p.text = ""


def fill_table0(table, row_data):
    rows = table.rows
    set_cell_text(rows[0].cells[5], row_data['device_type'])
    set_cell_text(rows[1].cells[1], row_data['interval'])
    set_cell_text(rows[1].cells[5], row_data['phase'])
    set_cell_text(rows[2].cells[1], row_data['device_name'])
    set_cell_text(rows[2].cells[5], row_data['distance'])
    set_cell_text(rows[3].cells[1], row_data['voltage'])
    set_cell_text(rows[3].cells[5], row_data['current'])
    set_cell_text(rows[3].cells[7], row_data['power'])
    date_str = row_data['date'].strftime('%Y-%m-%d') if isinstance(row_data['date'], datetime.date) else row_data['date']
    set_cell_text(rows[4].cells[1], date_str if date_str else NOT_RECOGNIZED_TEXT)
    set_cell_text(rows[4].cells[5], row_data['time'] if row_data['time'] else NOT_RECOGNIZED_TEXT)


def fill_table1(table, row_data):
    row = table.rows[0]
    ar1_display = f"{row_data['ar1']}℃" if row_data['ar1'] not in (None, NOT_RECOGNIZED_TEXT) else NOT_RECOGNIZED_TEXT
    set_cell_text(row.cells[2], ar1_display)


def fill_table2(table, row_data):
    row = table.rows[1]
    cell = row.cells[0]
    for p in cell.paragraphs:
        for run in list(p.runs):
            run.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    img_path = row_data.get('image')
    if img_path and os.path.exists(img_path):
        run = p.add_run()
        run.add_picture(img_path, width=Cm(IMAGE_WIDTH_CM), height=Cm(IMAGE_HEIGHT_CM))
    else:
        run = p.add_run("（未找到对应图片）")


def add_page_break(body, before_element):
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    br.set(qn('w:type'), 'page')
    r.append(br)
    p.append(r)
    body.insert(list(body).index(before_element), p)


def step2_generate_docx(full_data, output_path):
    print(f"\n{'=' * 50}")
    print("  Step 2: 生成Word报告")
    print(f"{'=' * 50}\n")

    print(f"正在读取模板: {DOCX_TEMPLATE_PATH}")
    doc = Document(DOCX_TEMPLATE_PATH)
    body = doc.element.body

    sectPr = body.find(qn('w:sectPr'))
    if sectPr is None:
        raise RuntimeError("模板文档缺少页面设置(sectPr)，无法继续。")

    template_children = [child for child in body if child is not sectPr]
    template_block = [copy.deepcopy(el) for el in template_children]

    for el in template_children:
        body.remove(el)

    for idx, row_data in enumerate(full_data):
        if idx > 0:
            add_page_break(body, sectPr)

        block_copy = [copy.deepcopy(el) for el in template_block]
        for el in block_copy:
            body.insert(list(body).index(sectPr), el)

        tbls = [Table(el, doc) for el in block_copy if el.tag == qn('w:tbl')]
        if len(tbls) != 3:
            raise RuntimeError(f"模板表格数量异常，期望3个，实际{len(tbls)}个。")
        table0, table1, table2 = tbls

        fill_table0(table0, row_data)
        fill_table1(table1, row_data)
        fill_table2(table2, row_data)

        print(f"  [{idx + 1}/{len(full_data)}] 已生成: {row_data['device_name']}")

    doc.save(output_path)
    print(f"\nWord报告已生成: {output_path}")


# ===================== 主程序 =====================

def interactive_menu():
    """交互式菜单：PyCharm 直接运行时弹出选项，无需记命令。"""
    print()
    print("=" * 55)
    print("  91号文《红外精确测温报告》自动生成程序 V4")
    print("=" * 55)
    print(f"  工作目录: {BASE_DIR}")
    print()

    # 检查文件状态
    excel_exists = os.path.isfile(EXCEL_PATH)
    latest_excel = find_latest_output(EXCEL_OUTPUT_BASE, "xlsx")
    excel_out_exists = latest_excel is not None
    image_dir_exists = os.path.isdir(IMAGE_DIR)
    template_exists = os.path.isfile(DOCX_TEMPLATE_PATH)
    existing_reports = sorted(glob.glob(f"{DOCX_OUTPUT_BASE}_*.docx"))

    print("  文件状态检查：")
    print(f"    模板 Excel:       {'✓ 存在' if excel_exists else '✗ 缺失！'}")
    print(f"    Step1 输出 Excel: {'✓ 已生成' if excel_out_exists else '○ 尚未生成'}")
    print(f"    图谱文件夹:       {'✓ 存在' if image_dir_exists else '✗ 缺失！'}")
    print(f"    Word 模板:        {'✓ 存在' if template_exists else '✗ 缺失！'}")
    print(f"    已有报告文件:     {len(existing_reports)} 份" +
          (f" ({', '.join(os.path.basename(r) for r in existing_reports[-3:])})" if existing_reports else ""))
    print()

    print("  请选择运行模式：")
    print("    [1] 仅 Step 1 — OCR识别 + 人工核对 + 写入Excel")
    print("    [2] 仅 Step 2 — 读取已有Excel，生成Word报告")
    print("    [3] 完整流程 — Step 1 + Step 2 串联执行")
    print("    [0] 退出")
    print()

    while True:
        choice = input("  请输入数字 (0/1/2/3): ").strip()
        if choice in ("0", "1", "2", "3"):
            break
        print("  ⚠️ 无效输入，请输入 0、1、2 或 3。")

    if choice == "0":
        print("  已退出。")
        sys.exit(0)

    run_step1 = choice in ("1", "3")
    run_step2 = choice in ("2", "3")
    print()

    return run_step1, run_step2


def main():
    global USE_GUI

    # ---------- 命令行参数解析（终端用户保留） ----------
    parser = argparse.ArgumentParser(
        description="91号文《红外精确测温报告》自动生成程序 V4",
        formatter_class=argparse.RawDescriptionHelpFormatter,
        add_help=False,  # 先关掉默认help，避免和交互模式冲突
    )
    parser.add_argument("--step1", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--step2", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--cli", action="store_true", help=argparse.SUPPRESS)
    parser.add_argument("--suffix", type=str, default=None, help=argparse.SUPPRESS)
    parser.add_argument("--help", action="store_true", help=argparse.SUPPRESS)

    # 判断用户是否在终端传了参数
    has_args = len(sys.argv) > 1

    if has_args:
        # ========= 终端模式：走命令行参数 =========
        args, _ = parser.parse_known_args()

        if args.help:
            print("""
使用方式：
  python3 脚本名.py                        PyCharm直接运行 → 交互式菜单（无需记命令）
  python3 脚本名.py --step1                 终端：仅 step1
  python3 脚本名.py --step2                 终端：仅 step2
  python3 脚本名.py --step2 --suffix v2     终端：step2 + 版本后缀
  python3 脚本名.py --cli                   终端：命令行核对模式
""")
            sys.exit(0)

        run_step1 = args.step1 or not (args.step1 or args.step2)
        run_step2 = args.step2 or not (args.step1 or args.step2)

        if args.cli:
            USE_GUI = False
            print("已启用 --cli 模式：将使用命令行进行人工核对。\n")

        suffix = args.suffix
    else:
        # ========= PyCharm 模式：交互式菜单 =========
        run_step1, run_step2 = interactive_menu()
        suffix = None

    # ---------- GUI 预初始化（仅 macOS 需要，Windows 上反而会破坏 Tcl 解释器）----------
    if USE_GUI and run_step1 and platform.system() == "Darwin":
        try:
            import tkinter as tk
            _pre = tk.Tk()
            _pre.withdraw()
            _pre.destroy()
        except Exception:
            pass

    # ---------- 确定输出路径（自动编号） ----------
    excel_output_path = resolve_output_path(EXCEL_OUTPUT_BASE, "xlsx", suffix)
    docx_output_path = resolve_output_path(DOCX_OUTPUT_BASE, "docx", suffix)

    # ---------- 执行 Step 1 ----------
    if run_step1:
        if not os.path.isdir(IMAGE_DIR):
            print(f"错误：图谱文件夹不存在: {IMAGE_DIR}")
            sys.exit(1)
        if not os.path.isfile(EXCEL_PATH):
            print(f"错误：Excel文件不存在: {EXCEL_PATH}")
            sys.exit(1)
        if not os.path.isfile(DOCX_TEMPLATE_PATH):
            print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
            sys.exit(1)

        full_data = step1_recognize_and_review(excel_output_path)
        print(f"\n[Step 1 完成] 识别结果已写入: {excel_output_path}")

    # ---------- 执行 Step 2 ----------
    if run_step2:
        if not os.path.isfile(DOCX_TEMPLATE_PATH):
            print(f"错误：Word模板文件不存在: {DOCX_TEMPLATE_PATH}")
            sys.exit(1)

        if run_step1:
            # step1 刚跑完，full_data 在内存中，直接复用
            print(f"\n{'=' * 50}")
            print("  进入 Step 2（使用 Step 1 的内存数据）")
            print(f"{'=' * 50}")
        else:
            # 只跑 step2：自动找最新的 Excel 文件
            latest_excel = find_latest_output(EXCEL_OUTPUT_BASE, "xlsx")
            if latest_excel is None:
                print(f"错误：找不到任何 必填信息_已生成_*.xlsx 文件。")
                print("  请先运行 Step 1 生成识别结果。")
                sys.exit(1)
            print(f"\n自动读取最新 Excel: {os.path.basename(latest_excel)}")
            full_data = read_excel_output(latest_excel)

        step2_generate_docx(full_data, docx_output_path)

        # 显示已有报告列表
        existing = sorted(glob.glob(f"{DOCX_OUTPUT_BASE}_*.docx"))
        if len(existing) > 1:
            print(f"\n当前目录下共有 {len(existing)} 份报告（均未被覆盖）：")
            for f in existing:
                print(f"  - {os.path.basename(f)}")

    # ---------- 汇总 ----------
    print(f"\n{'=' * 50}")
    print("  全部完成！")
    if run_step1:
        print(f"  [Step 1] Excel结果: {excel_output_path}")
    if run_step2:
        print(f"  [Step 2] Word报告:  {docx_output_path}")
    else:
        print(f"  （Step 2 未执行，再次运行本脚本选择 [2] 即可生成报告）")
    print(f"{'=' * 50}")


if __name__ == "__main__":
    main()
